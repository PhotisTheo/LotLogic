import base64
import csv
import hashlib
import json
import logging
import math
import re
import textwrap
from collections import defaultdict
from urllib.parse import urlencode, urlsplit, urlunsplit, parse_qsl
from concurrent.futures import ThreadPoolExecutor, as_completed
from decimal import ROUND_HALF_UP, Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple, NamedTuple
from urllib.parse import quote, urlencode, urljoin
from datetime import datetime, timedelta

try:
    from PIL import Image
except ImportError:
    Image = None


try:  # Stripe is optional; endpoints guard against missing dependency
    import stripe  # type: ignore
except ImportError:  # pragma: no cover - handled gracefully at runtime
    stripe = None

import pandas as pd
import requests
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.db.utils import NotSupportedError
from django.http import Http404, HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils import timezone
from django.utils.text import slugify
from django.core.cache import cache

# Core Django helpers + service imports for lead CRM views.
# Request method guards for the various endpoints.
from django.views.decorators.http import require_GET, require_POST, require_http_methods

from accounts.models import get_workspace_owner

from .forms import (
    LeadForm,
    ParcelListSaveForm,
    ParcelSearchForm,
    ScheduleCallRequestForm,
    UploadFileForm,
)
from .mailers import (
    MailerAgentProfile,
    MailerFallbackContext,
    collect_property_facts,
    get_mailer_script_options,
    guess_property_sector,
    render_mailer_script,
)
from .models import Lead, SavedParcelList, ScheduleCallRequest, SkipTraceRecord, AttomData
from .services import (
    PARCEL_SEARCH_MAX_RESULTS,
    MassGISDataError,
    SkipTraceError,
    calculate_equity_metrics,
    fetch_parcel_shape_for_lead,
    geocode_address,
    get_massgis_parcel_detail,
    get_massgis_parcel_shape,
    get_massgis_property_type_choices,
    get_precomputed_parcels_in_bbox,
    has_precomputed_parcels,
    load_massgis_parcels_by_ids,
    massgis_stateplane_to_wgs84,
    preload_massgis_dataset,
    search_massgis_parcels,
    search_precomputed_parcels,
    skiptrace_property,
)
from .attom_service import (
    build_attom_cache_key,
    ensure_attom_cache_record,
    get_or_fetch_attom_data,
    update_attom_data_for_parcel,
)
from .background_lien_search import (
    ensure_legal_action_record,
    search_parcel_background,
    should_search_parcel,
)


logger = logging.getLogger(__name__)

SKIPTRACE_CACHE_TTL_DAYS = getattr(settings, "SKIPTRACE_CACHE_TTL_DAYS", 90)
LIEN_SEARCH_AUTO_THRESHOLD = getattr(settings, "LIEN_SEARCH_AUTO_THRESHOLD", 1000)


def _parse_boundary_shape(data) -> Optional[Dict[str, object]]:
    shape_type = data.get("boundary_shape_type")
    if not shape_type:
        return None
    shape_type = shape_type.strip().lower()
    if shape_type == "circle":
        try:
            lat = float(data.get("boundary_circle_lat"))
            lng = float(data.get("boundary_circle_lng"))
            radius = float(data.get("boundary_circle_radius_miles"))
        except (TypeError, ValueError):
            return None
        return {
            "type": "circle",
            "center_lat": lat,
            "center_lng": lng,
            "radius_miles": radius,
            "source": "boundary",
        }
    if shape_type == "polygon":
        raw = data.get("boundary_polygon_coords")
        if not raw:
            return None
        try:
            coords = json.loads(raw)
        except json.JSONDecodeError:
            return None
        if not isinstance(coords, list) or not coords:
            return None
        return {
            "type": "polygon",
            "coordinates": coords,
        }
    return None


def _serialize_shape_filter(shape_filter: Optional[Dict[str, object]]) -> Optional[Dict[str, object]]:
    if not shape_filter:
        return None
    if shape_filter.get("type") == "circle":
        return {
            "type": "circle",
            "center_lat": shape_filter.get("center_lat"),
            "center_lng": shape_filter.get("center_lng"),
            "radius_miles": shape_filter.get("radius_miles"),
        }
    if shape_filter.get("type") == "polygon":
        return {
            "type": "polygon",
            "coordinates": shape_filter.get("coordinates"),
        }
    return None


def _calculate_mortgage_balance_from_attom(attom_data, current_value):
    """
    Calculate current mortgage balance using ATTOM data.
    Uses actual mortgage info if available, otherwise returns None.

    Args:
        attom_data: AttomData model instance
        current_value: Current property value

    Returns:
        Tuple of (mortgage_balance, equity_value, equity_percent, roi_percent, monthly_payment)
    """
    from datetime import datetime
    from decimal import Decimal
    from .services import MORTGAGE_RATE_BY_YEAR

    if not attom_data or not attom_data.mortgage_loan_amount:
        return None, None, None, None, None

    original_loan = float(attom_data.mortgage_loan_amount)
    interest_rate = float(attom_data.mortgage_interest_rate) if attom_data.mortgage_interest_rate else None
    term_years = attom_data.mortgage_term_years or 30  # Default to 30 years if unknown
    recording_date_str = attom_data.mortgage_recording_date

    # If no interest rate from ATTOM, use historical average based on mortgage year
    if not interest_rate and recording_date_str:
        try:
            mortgage_year = datetime.strptime(recording_date_str, "%Y-%m-%d").year
            interest_rate = MORTGAGE_RATE_BY_YEAR.get(mortgage_year)
            if interest_rate:
                print(f"[MORTGAGE CALC] Using historical rate {interest_rate}% for year {mortgage_year}")
            else:
                # Use most recent year's rate as fallback
                latest_year = max(MORTGAGE_RATE_BY_YEAR.keys())
                interest_rate = MORTGAGE_RATE_BY_YEAR[latest_year]
                print(f"[MORTGAGE CALC] Year {mortgage_year} not in historical data, using {latest_year} rate: {interest_rate}%")
        except (ValueError, TypeError) as e:
            print(f"[MORTGAGE CALC] Could not parse mortgage date {recording_date_str}: {e}")

    # Calculate age of mortgage in months
    if recording_date_str:
        try:
            recording_date = datetime.strptime(recording_date_str, "%Y-%m-%d")
            months_elapsed = (datetime.now() - recording_date).days / 30.44
            months_elapsed = max(0, months_elapsed)  # Can't be negative
        except (ValueError, TypeError):
            months_elapsed = 0
    else:
        months_elapsed = 0

    # Calculate current balance
    if interest_rate and interest_rate > 0:
        # Use amortization formula
        monthly_rate = interest_rate / 100 / 12
        total_months = term_years * 12

        # Monthly payment formula: M = P * [r(1+r)^n] / [(1+r)^n - 1]
        if monthly_rate > 0:
            monthly_payment = original_loan * (monthly_rate * (1 + monthly_rate) ** total_months) / ((1 + monthly_rate) ** total_months - 1)
        else:
            monthly_payment = original_loan / total_months

        # Current balance formula: B = P * [(1+r)^n - (1+r)^p] / [(1+r)^n - 1]
        if months_elapsed < total_months:
            current_balance = original_loan * ((1 + monthly_rate) ** total_months - (1 + monthly_rate) ** months_elapsed) / ((1 + monthly_rate) ** total_months - 1)
        else:
            current_balance = 0  # Loan paid off
    else:
        # No interest rate available, use straight-line depreciation
        monthly_payment = original_loan / (term_years * 12)
        payments_made = min(months_elapsed, term_years * 12)
        current_balance = original_loan - (monthly_payment * payments_made)
        current_balance = max(0, current_balance)

    # Calculate equity metrics
    if current_value:
        equity_value = current_value - current_balance
        equity_percent = (equity_value / current_value * 100) if current_value > 0 else 0

        # ROI calculation: (Current Equity - Down Payment) / Down Payment
        # Assume 20% down payment
        estimated_down_payment = original_loan / 0.8 * 0.2  # If loan was 80%, down was 20%
        roi_percent = ((equity_value - estimated_down_payment) / estimated_down_payment * 100) if estimated_down_payment > 0 else 0
    else:
        equity_value = None
        equity_percent = None
        roi_percent = None

    return (
        round(current_balance, 2) if current_balance else None,
        round(equity_value, 2) if equity_value else None,
        round(equity_percent, 2) if equity_percent is not None else None,
        round(roi_percent, 2) if roi_percent is not None else None,
        round(monthly_payment, 2) if monthly_payment else None
    )

CITY_FALLBACK_LABEL = "Other / Unknown City"
ZONING_FALLBACK_LABEL = "Other / Unknown Zoning"


def _saved_list_queryset_for_user(user):
    owner = get_workspace_owner(user) if getattr(user, "is_authenticated", False) else None
    if not owner:
        return SavedParcelList.objects.none()
    return SavedParcelList.objects.filter(created_by=owner)


def _lead_queryset_for_user(user):
    owner = get_workspace_owner(user) if getattr(user, "is_authenticated", False) else None
    if not owner:
        return Lead.objects.none()
    return Lead.objects.filter(created_by=owner)


# ---------------------------------------------------------------------------
# Shared helper utilities (map embeds, formatting, data normalisation)
# ---------------------------------------------------------------------------
def _build_mapillary_embed_url(
    lat: Optional[float], lng: Optional[float], *, zoom: int = 17
) -> Optional[str]:
    if lat is None or lng is None:
        return None
    client_id = getattr(settings, "MAPILLARY_CLIENT_ID", "")
    if not client_id:
        return None
    return (
        "https://www.mapillary.com/embed"
        f"?lat={lat:.6f}&lng={lng:.6f}&z={zoom}&bearing=0&pitch=0&fov=70&layer=street&mapillaryClientId={client_id}"
    )


def _build_google_street_view_embed(
    lat: Optional[float], lng: Optional[float]
) -> Optional[str]:
    if lat is None or lng is None:
        return None
    base = "https://maps.google.com/maps"
    params = {
        "q": f"{lat:.6f},{lng:.6f}",
        "layer": "c",
        "cbll": f"{lat:.6f},{lng:.6f}",
        "ll": f"{lat:.6f},{lng:.6f}",
        "cbp": "11,0,0,0,0",
        "output": "svembed",
    }
    return f"{base}?{urlencode(params)}"


def _format_bed_bath(value: Optional[object]) -> str:
    if value in (None, "", " "):
        return "—"
    try:
        number = float(value)
    except (TypeError, ValueError):
        text = str(value).strip()
        return text or "—"
    if not math.isfinite(number):
        return "—"
    if abs(number) < 0.01:
        return "—"
    if abs(number - round(number)) < 0.01:
        return str(int(round(number)))
    formatted = f"{number:.1f}".rstrip("0").rstrip(".")
    return formatted or str(number)


def _serialize_skiptrace_record(record: SkipTraceRecord) -> Optional[dict]:
    if not record:
        return None
    phones = record.phones or []
    return {
        "ownerName": record.owner_name or "",
        "email": record.email or "",
        "phones": phones[:3],
        "lastUpdated": record.updated_at.isoformat() if record.updated_at else None,
    }


def _normalize_loc_id(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


class SavedListParcelRef(NamedTuple):
    town_id: int
    loc_id: str
    normalized_loc_id: str


def _iter_saved_list_parcel_refs(saved_list: SavedParcelList) -> Iterable[SavedListParcelRef]:
    """
    Yield normalized parcel references stored on a SavedParcelList while preserving order.
    Supports both legacy string lists and the new dict-based storage that includes town ids.
    """
    loc_entries = saved_list.loc_ids if isinstance(saved_list.loc_ids, list) else []
    default_town_id = saved_list.town_id

    for entry in loc_entries:
        entry_town_id = default_town_id
        loc_value = entry

        if isinstance(entry, dict):
            entry_town_id = entry.get("town_id", default_town_id)
            loc_value = (
                entry.get("loc_id")
                or entry.get("locId")
                or entry.get("id")
            )

        if loc_value is None:
            continue

        normalized_loc = _normalize_loc_id(loc_value)
        if not normalized_loc:
            continue

        try:
            town_int = int(entry_town_id) if entry_town_id is not None else None
        except (TypeError, ValueError):
            town_int = None

        if town_int is None:
            continue

        loc_text = str(loc_value).strip()
        if not loc_text:
            loc_text = normalized_loc

        yield SavedListParcelRef(town_int, loc_text, normalized_loc)


def _skiptrace_cache_ttl_days_value() -> Optional[int]:
    if SKIPTRACE_CACHE_TTL_DAYS is None:
        return None
    try:
        days = int(SKIPTRACE_CACHE_TTL_DAYS)
    except (TypeError, ValueError):
        return None
    return days if days > 0 else None


def _skiptrace_cache_cutoff():
    days = _skiptrace_cache_ttl_days_value()
    if not days:
        return None
    return timezone.now() - timedelta(days=days)


def _skiptrace_record_is_fresh(record: SkipTraceRecord) -> bool:
    cutoff = _skiptrace_cache_cutoff()
    if cutoff is None:
        return True
    timestamp = record.updated_at or record.created_at
    if timestamp is None:
        return False
    return timestamp >= cutoff


def _append_query_param(url: str, params: dict[str, str]) -> str:
    scheme, netloc, path, query, fragment = urlsplit(url)
    query_params = dict(parse_qsl(query, keep_blank_values=True))
    query_params.update(params)
    return urlunsplit((scheme, netloc, path, urlencode(query_params, doseq=True), fragment))


def _get_skiptrace_record_for_loc_id(
    town_id: Optional[int],
    loc_id: Optional[str],
    *,
    user=None,
    fresh_only: bool = False,
) -> Optional[SkipTraceRecord]:
    normalized = _normalize_loc_id(loc_id)
    if not normalized:
        return None

    owner = get_workspace_owner(user) if user else None
    if owner is None:
        return None

    base_qs = SkipTraceRecord.objects.filter(loc_id__iexact=normalized, created_by=owner)

    candidates: List[Optional[SkipTraceRecord]] = []
    if town_id is not None:
        candidates.append(base_qs.filter(town_id=town_id).order_by("-updated_at").first())
    candidates.append(base_qs.filter(town_id__isnull=True).order_by("-updated_at").first())
    candidates.append(base_qs.order_by("-updated_at").first())

    for record in candidates:
        if record and (not fresh_only or _skiptrace_record_is_fresh(record)):
            return record

    if fresh_only:
        return None

    for record in candidates:
        if record:
            return record
    return None


def _get_skiptrace_record_for_loc_ids(
    loc_ids: Iterable[str],
    *,
    town_id: Optional[int] = None,
    user=None,
    fresh_only: bool = False,
) -> Optional[SkipTraceRecord]:
    candidates = []
    for raw_loc_id in loc_ids:
        record = _get_skiptrace_record_for_loc_id(
            town_id, raw_loc_id, user=user, fresh_only=fresh_only
        )
        if record:
            candidates.append(record)
    if not candidates:
        return None
    return max(
        candidates,
        key=lambda rec: rec.updated_at or rec.created_at,
    )


def _get_shared_skiptrace_record(
    town_id: Optional[int],
    loc_id: Optional[str],
    *,
    exclude_owner=None,
    fresh_only: bool = True,
) -> Optional[SkipTraceRecord]:
    normalized = _normalize_loc_id(loc_id)
    if not normalized:
        return None

    queryset = SkipTraceRecord.objects.filter(loc_id__iexact=normalized)
    if town_id is not None:
        queryset = queryset.filter(Q(town_id=town_id) | Q(town_id__isnull=True))
    if exclude_owner:
        queryset = queryset.exclude(created_by=exclude_owner)

    queryset = queryset.order_by("-updated_at")
    for record in queryset:
        if not fresh_only or _skiptrace_record_is_fresh(record):
            return record
    return None


def _clone_skiptrace_record_for_owner(
    source_record: SkipTraceRecord,
    *,
    owner,
    town_id: Optional[int],
    loc_id: Optional[str],
) -> Optional[SkipTraceRecord]:
    normalized = _normalize_loc_id(loc_id)
    if not (source_record and owner and normalized):
        return None

    target_town_id = town_id if town_id is not None else source_record.town_id

    defaults = {
        "owner_name": source_record.owner_name or "",
        "email": source_record.email or "",
        "phones": source_record.phones or [],
        "raw_payload": source_record.raw_payload or {},
    }

    record, _ = SkipTraceRecord.objects.update_or_create(
        created_by=owner,
        town_id=target_town_id,
        loc_id=normalized,
        defaults=defaults,
    )

    # Preserve original freshness timestamp when possible so future TTL checks align.
    if source_record.updated_at and record.updated_at < source_record.updated_at:
        SkipTraceRecord.objects.filter(pk=record.pk).update(updated_at=source_record.updated_at)
        record.refresh_from_db()

    return record


def _bulk_skiptrace_record_map(
    loc_ids: Iterable[str],
    *,
    town_id: Optional[int] = None,
    user=None,
) -> dict[str, SkipTraceRecord]:
    normalized_ids = {
        _normalize_loc_id(loc_id) for loc_id in loc_ids if _normalize_loc_id(loc_id)
    }
    if not normalized_ids:
        return {}

    owner = get_workspace_owner(user) if user else None
    if owner is None:
        return {}

    queryset = SkipTraceRecord.objects.filter(loc_id__in=normalized_ids, created_by=owner)
    if town_id is not None:
        queryset = queryset.filter(Q(town_id=town_id) | Q(town_id__isnull=True))
    queryset = queryset.order_by("-updated_at")

    record_map: dict[str, SkipTraceRecord] = {}
    for record in queryset:
        if not _skiptrace_record_is_fresh(record):
            continue
        key = _normalize_loc_id(record.loc_id)
        if key is None:
            continue
        existing = record_map.get(key)
        if existing is None:
            record_map[key] = record
            continue
        if town_id is not None:
            if record.town_id == town_id and existing.town_id != town_id:
                record_map[key] = record
                continue
        if (record.updated_at or record.created_at) and (
            existing.updated_at or existing.created_at
        ):
            if (record.updated_at or record.created_at) > (
                existing.updated_at or existing.created_at
            ):
                record_map[key] = record
    return record_map


def _iter_saved_lists_with_loc_id(
    loc_id: Optional[str],
    *,
    town_id: Optional[int] = None,
    user=None,
):
    normalized_target = _normalize_loc_id(loc_id)
    if not normalized_target:
        return

    base_qs = _saved_list_queryset_for_user(user) if user else SavedParcelList.objects.none()
    seen_ids: set[int] = set()
    found_match = False

    try:
        for saved_list in base_qs.filter(loc_ids__contains=[loc_id]):
            if saved_list.pk in seen_ids:
                continue
            for ref in _iter_saved_list_parcel_refs(saved_list):
                if ref.normalized_loc_id == normalized_target or ref.loc_id == loc_id:
                    if town_id is None or ref.town_id == town_id:
                        seen_ids.add(saved_list.pk)
                        found_match = True
                        yield saved_list
                        break
        if normalized_target != loc_id:
            for saved_list in base_qs.filter(loc_ids__contains=[normalized_target]):
                if saved_list.pk in seen_ids:
                    continue
                for ref in _iter_saved_list_parcel_refs(saved_list):
                    if ref.normalized_loc_id == normalized_target:
                        if town_id is None or ref.town_id == town_id:
                            seen_ids.add(saved_list.pk)
                            found_match = True
                            yield saved_list
                            break
        if found_match:
            return
    except NotSupportedError:
        pass

    for saved_list in base_qs.iterator():
        if saved_list.pk in seen_ids:
            continue
        for ref in _iter_saved_list_parcel_refs(saved_list):
            if ref.normalized_loc_id != normalized_target and ref.loc_id != loc_id:
                continue
            if town_id is not None and ref.town_id != town_id:
                continue
            seen_ids.add(saved_list.pk)
            yield saved_list
            break


def _saved_list_contains_loc_id(
    town_id: Optional[int], loc_id: str, *, user=None
) -> bool:
    for _ in _iter_saved_lists_with_loc_id(loc_id, town_id=town_id, user=user):
        return True
    return False


def _first_saved_list_town_id_for_loc_id(
    loc_id: Optional[str], *, user=None
) -> Optional[int]:
    for saved_list in _iter_saved_lists_with_loc_id(loc_id, user=user):
        return saved_list.town_id
    return None


def _resolve_owner_for_loc_id(
    loc_id: Optional[str], *, town_id: Optional[int] = None
):
    normalized = _normalize_loc_id(loc_id)
    if not normalized:
        return None

    lead_owner = (
        Lead.objects.filter(loc_id__iexact=normalized, created_by__isnull=False)
        .order_by("created_at")
        .first()
    )
    if lead_owner and lead_owner.created_by:
        return lead_owner.created_by

    candidate_loc_ids = {normalized}
    if loc_id and loc_id != normalized:
        candidate_loc_ids.add(loc_id)

    base_qs = SavedParcelList.objects.exclude(created_by__isnull=True)

    try:
        for candidate in candidate_loc_ids:
            match = base_qs.filter(loc_ids__contains=[candidate]).order_by("created_at").first()
            if match and match.created_by:
                return match.created_by
    except NotSupportedError:
        pass

    for saved_list in base_qs.iterator():
        if not saved_list.created_by:
            continue
        for ref in _iter_saved_list_parcel_refs(saved_list):
            if ref.normalized_loc_id in candidate_loc_ids or ref.loc_id in candidate_loc_ids:
                if town_id is None or ref.town_id == town_id:
                    return saved_list.created_by
                break

    return None


def _parse_skiptrace_payload(request) -> dict:
    if not getattr(request, "body", None):
        return {}
    try:
        payload = json.loads(
            request.body.decode("utf-8")
            if isinstance(request.body, bytes)
            else request.body
        )
    except (json.JSONDecodeError, TypeError, UnicodeDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def _is_truthy(value) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "y"}
    return False


def _skiptrace_allowed_for_parcel(town_id: int, loc_id: str, *, user=None) -> bool:
    # Superusers bypass all restrictions
    if user and getattr(user, 'is_superuser', False):
        return True

    parcel_saved = _saved_list_contains_loc_id(town_id, loc_id, user=user)
    lead_saved = (
        _lead_queryset_for_user(user).filter(loc_id__iexact=loc_id).exists()
        if user
        else False
    )
    return parcel_saved or lead_saved


def _normalize_dnc_value(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    upper = text.upper()
    if upper in {"Y", "YES", "TRUE"}:
        return "TRUE"
    if upper in {"N", "NO", "FALSE"}:
        return "FALSE"
    return upper


def _store_skiptrace_result(
    *, town_id: Optional[int], loc_id: str, result, user=None
) -> Optional[SkipTraceRecord]:
    normalized_loc_id = _normalize_loc_id(loc_id)
    if not normalized_loc_id:
        return None

    owner = get_workspace_owner(user) if user else None
    if owner is None:
        return None

    phone_payload = [
        {
            "number": phone.number,
            "type": phone.type,
            "score": phone.score,
            "dnc": phone.dnc,
        }
        for phone in (result.phones or [])
    ]
    record, _ = SkipTraceRecord.objects.update_or_create(
        created_by=owner,
        town_id=town_id,
        loc_id=normalized_loc_id,
        defaults={
            "owner_name": result.owner_name or "",
            "email": result.email or "",
            "phones": phone_payload,
            "raw_payload": result.raw_payload or {},
        },
    )
    return record


def _update_lead_contact_from_skiptrace(lead: Lead, result) -> None:
    updated_fields: list[str] = []

    if result.email:
        lead.email = result.email
        updated_fields.append("email")

    phones = result.phones or []
    field_mapping = [
        ("phone_1", "dnc_1"),
        ("phone_2", "dnc_2"),
        ("phone_3", "dnc_3"),
    ]

    for index, (phone_field, dnc_field) in enumerate(field_mapping):
        phone_meta = phones[index] if index < len(phones) else None
        if phone_meta:
            lead_value = phone_meta.number or None
            if getattr(lead, phone_field) != lead_value:
                setattr(lead, phone_field, lead_value)
                updated_fields.append(phone_field)

            dnc_value = _normalize_dnc_value(phone_meta.dnc)
            if getattr(lead, dnc_field) != dnc_value:
                setattr(lead, dnc_field, dnc_value)
                updated_fields.append(dnc_field)
        else:
            if getattr(lead, phone_field):
                setattr(lead, phone_field, None)
                updated_fields.append(phone_field)
            if getattr(lead, dnc_field):
                setattr(lead, dnc_field, None)
                updated_fields.append(dnc_field)

    if updated_fields:
        # Preserve field order but drop duplicates.
        seen = set()
        unique_fields = [
            field for field in updated_fields if not (field in seen or seen.add(field))
        ]
        lead.save(update_fields=unique_fields)


def _pending_parcels_for_saved_list(saved_list: SavedParcelList, *, user=None):
    parcel_refs = list(_iter_saved_list_parcel_refs(saved_list))
    if not parcel_refs:
        return [], {}, []

    parcels_by_key: dict[tuple[int, str], object] = {}
    grouped_loc_ids: dict[int, List[str]] = defaultdict(list)
    for ref in parcel_refs:
        grouped_loc_ids[ref.town_id].append(ref.loc_id)

    for town_id, loc_list in grouped_loc_ids.items():
        for parcel in load_massgis_parcels_by_ids(town_id, loc_list, saved_list=saved_list):
            normalized = _normalize_loc_id(parcel.loc_id)
            if not normalized:
                continue
            parcels_by_key[(town_id, normalized)] = parcel

    # Preserve original ordering from the saved list definition.
    parcels: list = []
    for ref in parcel_refs:
        parcel = parcels_by_key.get((ref.town_id, ref.normalized_loc_id))
        if parcel:
            parcels.append(parcel)

    loc_ids = [ref.normalized_loc_id for ref in parcel_refs]
    unique_town_ids = {ref.town_id for ref in parcel_refs}
    skiptrace_town_scope = unique_town_ids.pop() if len(unique_town_ids) == 1 else None
    skiptrace_records = _bulk_skiptrace_record_map(
        loc_ids, town_id=skiptrace_town_scope, user=user
    )

    pending: list = []
    seen = set()
    for parcel in parcels:
        normalized = _normalize_loc_id(parcel.loc_id)
        parcel_town_id = parcel.town.town_id if getattr(parcel, "town", None) else None
        if parcel_town_id is None:
            parcel_town_id = saved_list.town_id
        if not normalized or normalized in seen:
            continue
        if skiptrace_records.get(normalized):
            continue
        if not _skiptrace_allowed_for_parcel(parcel_town_id, parcel.loc_id, user=user):
            continue
        pending.append(parcel)
        seen.add(normalized)

    return parcels, skiptrace_records, pending


CSV_HEADER_LABELS = [
    "Parcel Address",
    "Owner",
    "Owner Mailing Address",
    "Email",
    "Phone Number 1",
    "Phone 1 DNC",
    "Phone Number 2",
    "Phone 2 DNC",
    "Phone Number 3",
    "Phone 3 DNC",
    "Total Value",
    "Est Mortgage Balance",
    "Est Equity",
    "Est ROI",
    "Sale Price",
    "Sale Date",
]


def _format_decimal_for_csv(value: Optional[object], *, digits: int) -> str:
    if value in (None, "", " "):
        return ""
    text = str(value).replace(",", "").strip()
    if not text:
        return ""
    try:
        decimal_value = Decimal(text)
    except (InvalidOperation, ValueError):
        return text
    quantizer = Decimal("1") if digits <= 0 else Decimal("1").scaleb(-digits)
    try:
        decimal_value = decimal_value.quantize(quantizer, rounding=ROUND_HALF_UP)
    except InvalidOperation:
        return text
    if digits <= 0:
        return f"{decimal_value:.0f}"
    return f"{decimal_value:.{digits}f}"


def _format_dnc_flag(value: Optional[object]) -> str:
    if value in (None, "", " "):
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    text = str(value).strip()
    if not text:
        return ""
    upper = text.upper()
    if upper in {"Y", "YES", "TRUE", "1"}:
        return "Yes"
    if upper in {"N", "NO", "FALSE", "0"}:
        return "No"
    return text


def _parcel_to_csv_row(
    parcel,
    skiptrace_record: Optional[SkipTraceRecord],
    *,
    default_city: Optional[str] = None,
) -> list[str]:
    city = (
        parcel.site_city
        or default_city
        or getattr(getattr(parcel, "town", None), "name", "")
        or ""
    )
    zip_code = parcel.site_zip or ""
    address_parts = [part for part in [parcel.site_address, city, zip_code] if part]
    parcel_address = ", ".join(address_parts)

    owner = parcel.owner_name or ""
    owner_address = parcel.owner_address or ""
    email = ""
    phones = []
    if skiptrace_record:
        email = skiptrace_record.email or ""
        phones = list(skiptrace_record.phones or [])

    phone_cells: list[str] = []
    for index in range(3):
        number = ""
        dnc_flag = ""
        phone_meta = phones[index] if index < len(phones) else None
        if isinstance(phone_meta, dict):
            number = phone_meta.get("number") or ""
            dnc_flag = _format_dnc_flag(phone_meta.get("dnc"))
        elif phone_meta:
            number = str(phone_meta)
        phone_cells.extend([number, dnc_flag])

    attrs = getattr(parcel, "attributes", {}) or {}
    total_value = _format_decimal_for_csv(
        getattr(parcel, "total_value", None), digits=0
    )
    est_mortgage_balance = _format_decimal_for_csv(
        getattr(parcel, "estimated_mortgage_balance", None), digits=0
    )
    est_equity = _format_decimal_for_csv(
        getattr(parcel, "estimated_equity_value", None), digits=0
    )
    est_roi = _format_decimal_for_csv(
        getattr(parcel, "estimated_roi_percent", None), digits=2
    )
    sale_price = _format_decimal_for_csv(attrs.get("LS_PRICE"), digits=0)
    sale_date = str(attrs.get("LS_DATE") or "")

    return [
        parcel_address,
        owner,
        owner_address,
        email,
        *phone_cells,
        total_value,
        est_mortgage_balance,
        est_equity,
        est_roi,
        sale_price,
        sale_date,
    ]


def _build_parcel_csv_response(
    parcels: Iterable,
    skiptrace_records: Optional[dict[str, SkipTraceRecord]],
    *,
    filename: str,
    default_city: Optional[str] = None,
) -> HttpResponse:
    safe_slug = slugify(filename) or "parcel-list"
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = f'attachment; filename="{safe_slug}.csv"'

    writer = csv.writer(response)
    writer.writerow(CSV_HEADER_LABELS)

    record_map = skiptrace_records or {}
    for parcel in parcels:
        normalized = _normalize_loc_id(getattr(parcel, "loc_id", None))
        skiptrace_record = record_map.get(normalized) if normalized else None
        row = _parcel_to_csv_row(parcel, skiptrace_record, default_city=default_city)
        writer.writerow(row)

    return response


def _extract_skiptrace_address_from_parcel(
    parcel,
) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
    attrs = getattr(parcel, "attributes", {}) or {}
    street = (
        attrs.get("OWN_ADDR")
        or getattr(parcel, "owner_address", None)
        or getattr(parcel, "site_address", None)
        or attrs.get("SITE_ADDR")
        or attrs.get("LOCATION")
        or attrs.get("FULL_STR")
    )
    city = (
        attrs.get("OWN_CITY") or getattr(parcel, "site_city", None) or attrs.get("CITY")
    )
    state = attrs.get("OWN_STATE") or getattr(parcel, "site_state", None) or "MA"
    zip_code = (
        attrs.get("OWN_ZIP")
        or getattr(parcel, "site_zip", None)
        or attrs.get("ZIP")
        or ""
    )
    return street, city, state, zip_code


def _build_zillow_url(full_address: Optional[str]) -> Optional[str]:
    """Build a Zillow deep link that lands on the target address when possible."""
    if not full_address:
        return None

    cleaned = " ".join(str(full_address).split())
    if not cleaned:
        return None

    slug = slugify(cleaned)
    if slug:
        return f"https://www.zillow.com/homes/{slug}_rb/"

    encoded = quote(cleaned, safe="")
    return f"https://www.zillow.com/homes/{encoded}"


def _build_google_street_view_image_url(
    lat: Optional[float], lng: Optional[float]
) -> Optional[str]:
    if lat is None or lng is None:
        return None

    endpoint = getattr(
        settings,
        "GOOGLE_STREET_VIEW_IMAGE_ENDPOINT",
        "https://maps.googleapis.com/maps/api/streetview",
    ).rstrip("?&")
    size = getattr(settings, "GOOGLE_STREET_VIEW_IMAGE_SIZE", "1200x600")
    key = getattr(settings, "GOOGLE_STREET_VIEW_API_KEY", "")
    heading = getattr(settings, "GOOGLE_STREET_VIEW_HEADING", None)
    pitch = getattr(settings, "GOOGLE_STREET_VIEW_PITCH", None)
    fov = getattr(settings, "GOOGLE_STREET_VIEW_FOV", None)

    params: dict[str, object] = {
        "size": size,
        "location": f"{lat},{lng}",
    }
    if heading is not None:
        params["heading"] = heading
    if pitch is not None:
        params["pitch"] = pitch
    if fov is not None:
        params["fov"] = fov
    if key:
        params["key"] = key

    query = urlencode({k: v for k, v in params.items() if v not in (None, "")})
    separator = "&" if "?" in endpoint else "?"
    return f"{endpoint}{separator}{query}"


def _load_agent_photo_data() -> Optional[str]:
    base_dir = Path(getattr(settings, "BASE_DIR", Path(__file__).resolve().parents[2]))
    configured_path = getattr(
        settings, "MAILER_AGENT_PHOTO_PATH", "leadcrm/photos/Home.png"
    )
    candidate_paths = [
        Path(configured_path),
        base_dir / configured_path,
        base_dir / "leadcrm" / "photos" / "Home.png",
    ]
    for path in candidate_paths:
        if not path.exists():
            continue
        try:
            data = path.read_bytes()
        except OSError:
            continue
        mime = "image/png"
        if path.suffix.lower() in {".jpg", ".jpeg"}:
            mime = "image/jpeg"
        if Image is not None:
            try:
                with Image.open(path) as img:
                    img = img.convert("RGB")
                    img.thumbnail((600, 600))
                    buffer = BytesIO()
                    img.save(buffer, format="JPEG", quality=85)
                    data = buffer.getvalue()
                    mime = "image/jpeg"
            except Exception:
                pass
        encoded = base64.b64encode(data).decode("ascii")
        return f"data:{mime};base64,{encoded}"
    return None


def _extract_owner_first_name(owner_name: Optional[str]) -> Optional[str]:
    if not owner_name:
        return None
    text = str(owner_name).strip()
    if not text:
        return None

    # Handle comma-delimited "Last, First" format.
    if "," in text:
        parts = [part.strip() for part in text.split(",") if part.strip()]
        if len(parts) >= 2:
            first = parts[1].split()
            if first:
                candidate = first[1] if len(first) >= 2 else first[0]
                if candidate:
                    return candidate.title()

    separators = [" and ", "&", "/", ";"]
    for sep in separators:
        if sep in text:
            text = text.split(sep, 1)[0].strip()
            break

    cleaned = text.replace(".", " ")
    tokens = [token for token in cleaned.split() if token]
    if not tokens:
        return None
    honorifics = {
        "mr",
        "mrs",
        "ms",
        "miss",
        "dr",
        "doctor",
        "estate",
        "est",
        "rev",
        "sir",
        "madam",
    }
    for token in tokens:
        normalized = token.lower()
        if normalized in honorifics:
            continue
        return token.title()
    return tokens[0].title()


def _build_zillow_screenshot_url(zillow_url: Optional[str]) -> Optional[str]:
    if not zillow_url:
        return None

    template = getattr(
        settings,
        "ZILLOW_SCREENSHOT_ENDPOINT",
        "https://image.thum.io/get/width/1200/crop/700/__URL__",
    )
    encoded = quote(zillow_url, safe="")

    if "__URL_ENCODED__" in template:
        return template.replace("__URL_ENCODED__", encoded)
    if "__URL__" in template:
        return template.replace("__URL__", zillow_url)
    if "{url}" in template:
        return template.replace("{url}", encoded)

    # Heuristic fallback when no placeholder is provided.
    if (
        template.endswith("/")
        or template.endswith("=")
        or template.endswith("-")
        or template.endswith("?")
        or template.endswith("&")
        or template.endswith("?url=")
    ):
        return f"{template}{zillow_url}"
    if "?" in template:
        return f"{template}&url={encoded}"
    return f"{template}/{zillow_url}"


def _build_qr_code_image_url(target: Optional[str]) -> Optional[str]:
    if not target:
        return None
    encoded = quote(target, safe="")

    template = getattr(
        settings,
        "MAILER_QR_IMAGE_ENDPOINT",
        "https://api.qrserver.com/v1/create-qr-code/?size=320x320&data=__DATA__",
    )
    if "__DATA__" in template:
        return template.replace("__DATA__", encoded)
    if "__URL__" in template:
        return template.replace("__URL__", encoded)

    if template.endswith("/") or template.endswith("="):
        return f"{template}{encoded}"
    joiner = "&" if "?" in template else "?"
    return f"{template}{joiner}data={encoded}"


def _normalize_capitalization(value: Optional[str]) -> Optional[str]:
    if not value:
        return value
    text = str(value)
    if not text:
        return text
    if text.isupper():
        return text.title()
    return text


# --- Mailer context assembly: build scripts, QR codes, and personalization.
def _generate_mailer_bundle(
    parcel,
    *,
    full_address: Optional[str],
    zillow_url: Optional[str],
    hero_image_url: Optional[str] = None,
    agent_photo_url: Optional[str] = None,
    request=None,
    town_id_override: Optional[int] = None,
    skiptrace_record=None,
    user=None,
) -> dict:
    property_address_parts = [
        full_address,
        (
            ", ".join(
                filter(None, [getattr(parcel, "site_address", None), getattr(parcel, "site_city", None), getattr(parcel, "site_zip", None)])
            )
            if parcel
            else None
        ),
        getattr(parcel, "site_address", None),
        getattr(parcel, "loc_id", None),
    ]
    property_address = next((part for part in property_address_parts if part), "")
    property_address = property_address.strip(", ")
    property_address = _normalize_capitalization(property_address) or property_address

    skiptrace_owner_name = (
        getattr(skiptrace_record, "owner_name", None) if skiptrace_record else None
    )
    parcel_owner_name = getattr(parcel, "owner_name", None)

    if skiptrace_owner_name:
        recipient_full_name = (
            _extract_owner_first_name(skiptrace_owner_name) or skiptrace_owner_name
        )
        greeting_source = skiptrace_owner_name
    else:
        recipient_full_name = parcel_owner_name
        greeting_source = parcel_owner_name

    greeting_name = _extract_owner_first_name(greeting_source) or "Neighbor"
    greeting_name = _normalize_capitalization(greeting_name) or greeting_name

    contact_phone_default = getattr(settings, "MAILER_CONTACT_PHONE", "555-555-5555")
    contact_phone = contact_phone_default
    text_keyword = getattr(settings, "MAILER_TEXT_KEYWORD", "HOME") or "HOME"
    text_keyword_upper = text_keyword.upper()

    agent_name = _normalize_capitalization(getattr(settings, "MAILER_AGENT_NAME", None))
    agent_title = _normalize_capitalization(getattr(settings, "MAILER_AGENT_TITLE", None))
    agent_company = _normalize_capitalization(getattr(settings, "MAILER_AGENT_COMPANY", None))
    agent_tagline = _normalize_capitalization(getattr(settings, "MAILER_AGENT_TAGLINE", None))

    workspace_user = None
    workspace_profile = None
    if request is not None and hasattr(request, "user"):
        try:
            workspace_user = get_workspace_owner(request.user)
        except Exception:  # noqa: BLE001
            workspace_user = None
        if workspace_user is not None:
            workspace_profile = getattr(workspace_user, "profile", None)

    if workspace_user is not None:
        name_candidate = (
            getattr(workspace_user, "get_full_name", lambda: "")() or workspace_user.get_username()
        )
        agent_name = _normalize_capitalization(name_candidate) or agent_name

    if workspace_profile is not None:
        company_candidate = getattr(workspace_profile, "company_name", "") or None
        title_candidate = getattr(workspace_profile, "job_title", "") or None
        phone_candidates = [
            getattr(workspace_profile, "mobile_phone", "") or None,
            getattr(workspace_profile, "work_phone", "") or None,
        ]

        def _clean_text(value: Optional[str]) -> Optional[str]:
            if not value:
                return None
            text = str(value).strip()
            return text or None

        company_candidate = _clean_text(company_candidate) or agent_company
        title_candidate = _clean_text(title_candidate) or agent_title

        cleaned_phone = None
        for candidate in phone_candidates:
            cleaned_candidate = _clean_text(candidate)
            if cleaned_candidate:
                cleaned_phone = cleaned_candidate
                break
        contact_phone = cleaned_phone or contact_phone

        tagline_candidate = _clean_text(getattr(workspace_profile, "bio", "") or None)
        if tagline_candidate:
            first_line = tagline_candidate.splitlines()[0].strip()
            if len(first_line) > 140:
                first_line = f"{first_line[:137].rstrip()}…"
            agent_tagline = first_line

        agent_company = company_candidate
        agent_title = title_candidate

    if not contact_phone:
        contact_phone = contact_phone_default

    agent = MailerAgentProfile(
        name=agent_name,
        title=agent_title,
        company=agent_company,
        tagline=agent_tagline,
    )

    agent_photo_data = (
        agent_photo_url if agent_photo_url and agent_photo_url.startswith("data:") else None
    )
    if not agent_photo_data:
        agent_photo_data = _load_agent_photo_data()

    qr_base = getattr(settings, "MAILER_QR_BASE_URL", None)
    raw_loc = getattr(parcel, "loc_id", None)
    normalized_loc = _normalize_loc_id(raw_loc) or slugify(
        property_address
    ) or ""

    town = getattr(parcel, "town", None)
    town_id = (
        getattr(town, "town_id", None)
        or getattr(parcel, "town_id", None)
        or town_id_override
    )

    schedule_url = None
    slug_loc = raw_loc or normalized_loc
    if town_id and slug_loc:
        try:
            schedule_path = reverse("schedule_call_request", args=[int(town_id), slug_loc])

            # Add user_id parameter to ensure leads are attributed to the correct user
            user_id = getattr(user, 'id', None) if user else None
            if user_id:
                from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
                parsed = urlparse(schedule_path)
                query_params = parse_qs(parsed.query)
                query_params['user_id'] = [str(user_id)]
                new_query = urlencode(query_params, doseq=True)
                schedule_path = urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, new_query, parsed.fragment))

            if request is not None:
                schedule_url = request.build_absolute_uri(schedule_path)
            else:
                base_url = getattr(settings, "SITE_BASE_URL", None)
                if base_url:
                    schedule_url = urljoin(base_url, schedule_path.lstrip("/"))
        except Exception:  # noqa: BLE001
            schedule_url = None

    qr_target = None
    if schedule_url:
        qr_target = schedule_url
    elif qr_base:
        if "__LOC__" in qr_base and normalized_loc:
            qr_target = qr_base.replace("__LOC__", normalized_loc)
        else:
            suffix = normalized_loc if normalized_loc else ""
            if qr_base.endswith(("/", "-", "_", "=")):
                qr_target = f"{qr_base}{suffix}"
            elif suffix:
                qr_target = f"{qr_base}/{suffix}"
            else:
                qr_target = qr_base

    if not qr_target:
        normalized_phone = re.sub(r"[^\d+]", "", contact_phone)
        qr_target = f"tel:{normalized_phone}" if normalized_phone else None

    qr_image_url = _build_qr_code_image_url(qr_target)
    if contact_phone:
        qr_caption = f"Call {contact_phone} or scan the QR code to schedule a call."
    else:
        qr_caption = "Scan the QR code to schedule a call."

    def _format_currency_value(raw) -> Optional[str]:
        if raw in (None, "", " "):
            return None
        try:
            return f"${float(raw):,.0f}"
        except (ValueError, TypeError):
            return None

    def _format_percent_value(raw) -> Optional[str]:
        if raw in (None, "", " "):
            return None
        try:
            return f"{float(raw):.1f}%"
        except (ValueError, TypeError):
            return None

    assessed_value_display = _format_currency_value(getattr(parcel, "total_value", None))
    equity_value_display = _format_currency_value(getattr(parcel, "estimated_equity_value", None))
    equity_percent_display = _format_percent_value(getattr(parcel, "equity_percent", None))

    property_address_display = _normalize_capitalization(
        property_address or full_address or getattr(parcel, "loc_id", "") or ""
    )
    property_facts = collect_property_facts(
        parcel,
        property_address_display,
        full_address,
        zillow_zestimate_display=None,
        assessed_value_display=assessed_value_display,
    )

    fallback_ctx = MailerFallbackContext(
        parcel=parcel,
        property_address=property_address_display,
        full_address=full_address,
        greeting_name=greeting_name,
        recipient_full_name=recipient_full_name,
        contact_phone=contact_phone,
        text_keyword=text_keyword,
        text_keyword_upper=text_keyword_upper,
        agent=agent,
        zillow_url=zillow_url,
        property_facts=property_facts,
        equity_value_display=equity_value_display,
        equity_percent_display=equity_percent_display,
        total_value_display=assessed_value_display,
        zillow_zestimate_value=None,
        zillow_zestimate_display=None,
    )

    base_fields = {
        "qr_image_url": qr_image_url,
        "qr_target": qr_target,
        "qr_caption": qr_caption,
        "contact_phone": contact_phone,
        "text_keyword": text_keyword_upper,
        "agent_name": agent.name,
        "agent_title": agent.title,
        "agent_company": agent.company,
        "agent_tagline": agent.tagline,
        "agent_photo_data": agent_photo_data,
        "hero_image_url": hero_image_url,
        "zillow_url": zillow_url,
        "property_facts": property_facts,
        "greeting_name": greeting_name,
    }

    scripts: Dict[str, dict] = {}
    prompt_options: List[dict] = []
    options = get_mailer_script_options(workspace_user)

    for option in options:
        script_payload = render_mailer_script(option.id, fallback_ctx, owner=workspace_user)
        script_payload.update(base_fields)
        script_schedule_url = schedule_url
        if schedule_url:
            separator = "&" if "?" in schedule_url else "?"
            script_schedule_url = f"{schedule_url}{separator}script={option.id}"
        script_payload["schedule_url"] = script_schedule_url
        script_qr_target = script_schedule_url or qr_target
        if script_qr_target:
            script_payload["qr_target"] = script_qr_target
            script_payload["qr_image_url"] = _build_qr_code_image_url(script_qr_target)
        script_payload.setdefault("generated", True)
        script_payload.setdefault("ai_generated", False)
        script_payload.setdefault("ai_model", None)
        scripts[option.id] = script_payload
        prompt_options.append(
            {
                "id": option.id,
                "label": option.label,
                "sector": option.sector,
                "summary": option.summary,
                "promptText": option.prompt_text or "",
            }
        )

    default_sector = guess_property_sector(parcel)
    default_option = next(
        (option for option in options if option.sector == default_sector),
        options[0],
    )

    return {
        "scripts": scripts,
        "options": prompt_options,
        "default_id": default_option.id,
        "sector": default_sector,
    }


def _build_mailer_context(
    parcel,
    *,
    full_address: Optional[str],
    zillow_url: Optional[str],
    hero_image_url: Optional[str] = None,
    agent_photo_url: Optional[str] = None,
    download_endpoint: Optional[str] = None,
    request=None,
    town_id_override: Optional[int] = None,
    skiptrace_record=None,
    user=None,
) -> dict:
    bundle = _generate_mailer_bundle(
        parcel,
        full_address=full_address,
        zillow_url=zillow_url,
        hero_image_url=hero_image_url,
        agent_photo_url=agent_photo_url,
        request=request,
        town_id_override=town_id_override,
        skiptrace_record=skiptrace_record,
        user=user,
    )

    scripts = bundle["scripts"]
    script_view_map: Dict[str, dict] = {}
    for script_id, script in scripts.items():
        pdf_url = None
        if download_endpoint:
            separator = "&" if "?" in download_endpoint else "?"
            pdf_url = f"{download_endpoint}{separator}script={quote(script_id)}"
        script["pdf_url"] = pdf_url
        rendered_html = render_to_string(
            "leads/partials/_mailer_letter.html",
            {"mailer": script},
        )
        script["rendered_html"] = rendered_html
        script_view_map[script_id] = {
            "html": rendered_html,
            "summary": script.get("summary"),
            "label": script.get("prompt_label"),
            "sector": script.get("sector"),
            "pdfUrl": pdf_url,
            "scheduleUrl": script.get("schedule_url"),
        }

    default_id = bundle["default_id"]
    default_script = scripts[default_id]

    mailer_context = dict(default_script)
    mailer_context.setdefault("generated", True)
    mailer_context.setdefault("ai_generated", False)
    mailer_context.setdefault("ai_model", None)
    mailer_context["prompt_options"] = bundle["options"]
    mailer_context["prompt_selected"] = default_id
    mailer_context["script_map_json"] = json.dumps(script_view_map, ensure_ascii=False)
    mailer_context["summary"] = default_script.get("summary")
    mailer_context["status"] = f'Displaying "{default_script.get("prompt_label")}".'
    mailer_context["preview_html"] = script_view_map[default_id]["html"]
    mailer_context["download_base"] = download_endpoint
    return mailer_context


def _render_mailer_script_for_parcel(
    parcel,
    script_id: str,
    *,
    full_address: Optional[str],
    zillow_url: Optional[str],
    hero_image_url: Optional[str] = None,
    request=None,
    town_id_override: Optional[int] = None,
    skiptrace_record=None,
    user=None,
) -> tuple[str, dict, str]:
    bundle = _generate_mailer_bundle(
        parcel,
        full_address=full_address,
        zillow_url=zillow_url,
        hero_image_url=hero_image_url,
        request=request,
        town_id_override=town_id_override,
        skiptrace_record=skiptrace_record,
        user=user,
    )
    scripts = bundle["scripts"]
    selected_id = script_id if script_id in scripts else bundle["default_id"]
    script = scripts[selected_id]
    html = render_to_string("leads/partials/_mailer_letter.html", {"mailer": script})
    return selected_id, script, html


def _pdf_escape(text: str) -> str:
    return (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def _wrap_lines_for_pdf(lines: Iterable[str], max_width: int = 90) -> List[str]:
    wrapped: List[str] = []
    for raw in lines:
        if raw is None:
            raw = ""
        text = str(raw)
        if not text:
            wrapped.append("")
            continue
        segments = textwrap.wrap(text, width=max_width, replace_whitespace=False)
        if not segments:
            wrapped.append("")
        else:
            wrapped.extend(segments)
    return wrapped


# Minimal PDF writer that lays out text payloads and embeds QR images.
def _build_simple_pdf(pages: List[dict]) -> bytes:
    if not pages:
        pages = [{"lines": [""], "image": None}]

    pdf_objects: List[Optional[bytes]] = []

    def add_object(payload: Optional[bytes]) -> int:
        pdf_objects.append(payload)
        return len(pdf_objects)

    catalog_obj_id = add_object(None)
    pages_obj_id = add_object(None)
    font_obj_id = add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_ids: List[int] = []

    for page_index, page in enumerate(pages, start=1):
        lines = page.get("lines") or [""]
        image_info = page.get("image")

        content_lines = [
            "BT",
            "/F1 12 Tf",
            "72 720 Td",
            "16 TL",
        ]
        for line in lines:
            content_lines.append(f"({_pdf_escape(line)}) Tj")
            content_lines.append("T*")
        content_lines.append("ET")

        image_obj_id = None
        image_name = None
        if image_info:
            display_width = float(image_info.get("display_width") or 144.0)
            display_height = float(image_info.get("display_height") or display_width)
            origin_x = 360.0
            origin_y = 120.0
            image_name = f"Im{page_index}"
            content_lines.extend([
                "q",
                f"{display_width:.2f} 0 0 {display_height:.2f} {origin_x:.2f} {origin_y:.2f} cm",
                f"/{image_name} Do",
                "Q",
            ])

        content_stream = "\n".join(content_lines).encode("latin-1", "ignore")
        stream_bytes = (
            f"<< /Length {len(content_stream)} >>\n".encode("latin-1")
            + b"stream\n"
            + content_stream
            + b"\nendstream"
        )
        content_obj_id = add_object(stream_bytes)

        if image_info:
            image_bytes = image_info["bytes"]
            width = image_info["width"]
            height = image_info["height"]
            image_stream = (
                f"<< /Type /XObject /Subtype /Image /Width {width} /Height {height} /ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /DCTDecode /Length {len(image_bytes)} >>\n".encode("latin-1")
                + b"stream\n"
                + image_bytes
                + b"\nendstream"
            )
            image_obj_id = add_object(image_stream)

        resources_parts = [f"/Font << /F1 {font_obj_id} 0 R >>"]
        if image_obj_id is not None and image_name:
            resources_parts.append(f"/XObject << /{image_name} {image_obj_id} 0 R >>")
        resources_block = " ".join(resources_parts)

        page_obj_bytes = (
            f"<< /Type /Page /Parent {pages_obj_id} 0 R /MediaBox [0 0 612 792] /Resources << {resources_block} >> /Contents [{content_obj_id} 0 R] >>"
        ).encode("latin-1")
        page_obj_id = add_object(page_obj_bytes)
        page_obj_ids.append(page_obj_id)

    kids = " ".join(f"{obj_id} 0 R" for obj_id in page_obj_ids)
    pages_obj = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_obj_ids)} >>".encode("latin-1")
    catalog_obj = f"<< /Type /Catalog /Pages {pages_obj_id} 0 R >>".encode("latin-1")

    pdf_objects[catalog_obj_id - 1] = catalog_obj
    pdf_objects[pages_obj_id - 1] = pages_obj

    buffer = BytesIO()
    buffer.write(b"%PDF-1.4\n")
    offsets: List[int] = []
    for index, obj in enumerate(pdf_objects, start=1):
        offsets.append(buffer.tell())
        buffer.write(f"{index} 0 obj\n".encode("latin-1"))
        buffer.write(obj or b"")
        buffer.write(b"\nendobj\n")
    xref_position = buffer.tell()
    buffer.write(f"xref\n0 {len(pdf_objects)+1}\n".encode("latin-1"))
    buffer.write(b"0000000000 65535 f \n")
    for offset in offsets:
        buffer.write(f"{offset:010} 00000 n \n".encode("latin-1"))
    buffer.write(
        f"trailer\n<< /Size {len(pdf_objects)+1} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF".encode("latin-1")
    )
    return buffer.getvalue()





def _script_to_pdf_lines(script: dict) -> List[str]:
    lines: List[str] = []

    for entry in script.get("letter_lines", []):
        lines.append(entry or "")

    value_props = script.get("value_props") or []
    if value_props:
        lines.append("")
        title = script.get("value_props_title")
        if title:
            lines.append(title)
        for prop in value_props:
            lines.append(f"- {prop}")

    qr_caption = script.get("qr_caption")
    if qr_caption:
        if lines and lines[-1] != "":
            lines.append("")
        lines.append(qr_caption)

    return lines


# Helper to download/resize QR code assets for PDF output.
def _prepare_pdf_image(script: dict) -> Optional[dict]:
    url = script.get("qr_image_url")
    if not url:
        return None

    try:
        if url.startswith('data:') and ';base64,' in url:
            _, encoded = url.split(',', 1)
            raw_bytes = base64.b64decode(encoded)
        else:
            response = requests.get(url, timeout=5)
            response.raise_for_status()
            raw_bytes = response.content
    except Exception as exc:
        logger.debug('Unable to fetch QR image: %s', exc)
        return None

    if Image is None:
        return None

    try:
        with Image.open(BytesIO(raw_bytes)) as img:
            img = img.convert('RGB')
            buffer = BytesIO()
            img.save(buffer, format='JPEG', quality=85)
            jpeg_bytes = buffer.getvalue()
            display_width = 144.0
            aspect_ratio = img.height / img.width if img.width else 1.0
            display_height = display_width * aspect_ratio
            return {
                'bytes': jpeg_bytes,
                'width': img.width,
                'height': img.height,
                'display_width': display_width,
                'display_height': display_height,
            }
    except Exception as exc:
        logger.debug('Unable to process QR image: %s', exc)
        return None


# Convert mailer scripts into page dictionaries that feed the PDF builder.
def _render_mailer_pdf(scripts: Iterable[dict]) -> bytes:
    max_lines_per_page = 46
    pdf_pages: List[dict] = []

    for script in scripts:
        image_info = _prepare_pdf_image(script)
        raw_lines = _script_to_pdf_lines(script)
        wrapped_lines = _wrap_lines_for_pdf(raw_lines)

        page_lines: List[str] = []
        for line in wrapped_lines:
            if len(page_lines) >= max_lines_per_page:
                pdf_pages.append({"lines": page_lines, "image": None})
                page_lines = []
            page_lines.append(line)

        if page_lines:
            pdf_pages.append({"lines": page_lines, "image": image_info})
        elif image_info:
            pdf_pages.append({"lines": [""], "image": image_info})

    return _build_simple_pdf(pdf_pages)


def _render_mailer_docx(scripts: list[dict]) -> bytes:
    """
    Generate a Word document (.docx) containing all mailer letters.
    Each letter includes formatted text, value propositions, and QR code if available.

    Args:
        scripts: List of mailer script dictionaries with letter_lines, value_props, qr_image_url, etc.

    Returns:
        Bytes of the generated .docx file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx library is required for Word document generation")

    # Pre-fetch all QR codes in parallel for better performance
    qr_cache = {}
    qr_urls_to_fetch = []
    for script in scripts:
        qr_url = script.get("qr_image_url")
        if qr_url and not (qr_url.startswith('data:') and ';base64,' in qr_url):
            qr_urls_to_fetch.append(qr_url)

    # Fetch QR codes in parallel using ThreadPoolExecutor
    if qr_urls_to_fetch:
        from concurrent.futures import ThreadPoolExecutor, as_completed

        def fetch_qr(url):
            try:
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                return url, response.content
            except Exception as exc:
                logger.warning(f"Failed to fetch QR code from {url}: {exc}")
                return url, None

        with ThreadPoolExecutor(max_workers=10) as executor:
            futures = {executor.submit(fetch_qr, url): url for url in qr_urls_to_fetch}
            for future in as_completed(futures):
                url, content = future.result()
                if content:
                    qr_cache[url] = content

    doc = Document()

    # Set up document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for idx, script in enumerate(scripts):
        # Add page break before each letter (except the first)
        if idx > 0:
            doc.add_page_break()

        # Add letter content
        letter_lines = script.get("letter_lines", [])
        for line in letter_lines:
            if line:
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                # Set font
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
            else:
                # Empty line for spacing
                doc.add_paragraph()

        # Add value propositions if present
        value_props = script.get("value_props") or []
        if value_props:
            doc.add_paragraph()  # Spacing

            value_props_title = script.get("value_props_title", "Why work with us:")
            title_p = doc.add_paragraph(value_props_title)
            title_p.style = 'Normal'
            for run in title_p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True

            for prop in value_props:
                p = doc.add_paragraph(f"• {prop}")
                p.style = 'Normal'
                p.paragraph_format.left_indent = Inches(0.25)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

        # Add QR code if present
        qr_image_url = script.get("qr_image_url")
        qr_caption = script.get("qr_caption")

        if qr_image_url:
            doc.add_paragraph()  # Spacing before QR code

            # Add caption if present
            if qr_caption:
                caption_p = doc.add_paragraph(qr_caption)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption_p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    run.font.italic = True

            # Add QR code image from cache or decode base64
            try:
                if qr_image_url.startswith('data:') and ';base64,' in qr_image_url:
                    _, encoded = qr_image_url.split(',', 1)
                    image_bytes = base64.b64decode(encoded)
                else:
                    image_bytes = qr_cache.get(qr_image_url)

                if image_bytes:
                    image_stream = BytesIO(image_bytes)
                    # Add centered paragraph for QR code
                    qr_paragraph = doc.add_paragraph()
                    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    qr_run = qr_paragraph.add_run()
                    qr_run.add_picture(image_stream, width=Inches(2))
                else:
                    raise ValueError("QR code not available")

            except Exception as exc:
                logger.warning(f"Failed to add QR code to Word doc: {exc}")
                # Add fallback text
                fallback_p = doc.add_paragraph("[QR code unavailable]")
                fallback_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in fallback_p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document to bytes
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _generate_label_sheet_docx(parcels: list, format_spec: dict) -> bytes:
    """
    Generate a Word document of mailing labels formatted for label sheets (e.g., Avery 5160).

    Args:
        parcels: List of parcel objects with owner_name, mailing_address, etc.
        format_spec: Dictionary with label dimensions and layout
            - cols: number of columns
            - rows: number of rows per page
            - width: label width in inches
            - height: label height in inches
            - margin_top: top margin in inches
            - margin_left: left margin in inches
            - gutter_h: horizontal gutter (space between columns) in inches
            - gutter_v: vertical gutter (space between rows) in inches

    Returns:
        Word document bytes
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Set page margins to match format_spec
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(format_spec["margin_top"])
        section.bottom_margin = Inches(0.5)  # Small bottom margin
        section.left_margin = Inches(format_spec["margin_left"])
        section.right_margin = Inches(0.5)  # Small right margin

    cols = format_spec["cols"]
    rows = format_spec["rows"]
    labels_per_page = cols * rows

    # Process parcels in batches of labels_per_page
    for page_idx, page_start in enumerate(range(0, len(parcels), labels_per_page)):
        if page_idx > 0:
            # Add page break between pages
            doc.add_page_break()

        page_parcels = parcels[page_start:page_start + labels_per_page]

        # Create a table for this page of labels
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths and row heights
        for row in table.rows:
            row.height = Inches(format_spec["height"])
            for cell in row.cells:
                cell.width = Inches(format_spec["width"])
                # Remove cell padding for tighter fit
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    node = OxmlElement(f'w:{margin_name}')
                    node.set(qn('w:w'), '0')
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)

        # Fill in the labels
        for idx, parcel in enumerate(page_parcels):
            row_idx = idx // cols
            col_idx = idx % cols
            cell = table.rows[row_idx].cells[col_idx]

            # Format the mailing address
            owner_name = getattr(parcel, 'owner_name', '') or ''
            mailing_address = getattr(parcel, 'mailing_address', '') or ''
            mailing_city = getattr(parcel, 'mailing_city', '') or ''
            mailing_state = getattr(parcel, 'mailing_state', '') or ''
            mailing_zip = getattr(parcel, 'mailing_zip', '') or ''

            # Build address lines
            address_lines = []
            if owner_name:
                address_lines.append(owner_name[:35])  # Truncate if too long
            if mailing_address:
                address_lines.append(mailing_address[:35])

            city_state_zip = f"{mailing_city}, {mailing_state} {mailing_zip}".strip(", ")
            if city_state_zip:
                address_lines.append(city_state_zip[:35])

            # If no mailing address, use site address as fallback
            if not address_lines or len(address_lines) < 2:
                site_address = getattr(parcel, 'site_address', '') or ''
                site_city = getattr(parcel, 'site_city', '') or ''
                site_zip = getattr(parcel, 'site_zip', '') or ''

                address_lines = []
                if owner_name:
                    address_lines.append(owner_name[:35])
                if site_address:
                    address_lines.append(site_address[:35])

                site_city_zip = f"{site_city}, MA {site_zip}".strip(", ")
                if site_city_zip != ", MA":
                    address_lines.append(site_city_zip[:35])

            # Add address to cell with small font
            for line in address_lines[:4]:  # Max 4 lines per label
                p = cell.add_paragraph(line)
                p.style.font.size = Pt(9)
                p.style.font.name = 'Arial'
                # Remove spacing
                p_format = p.paragraph_format
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                p_format.line_spacing = 1.0

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _generate_label_sheet_pdf(parcels: list, format_spec: dict) -> bytes:
    """
    Generate a PDF of mailing labels formatted for label sheets (e.g., Avery 5160).

    DEPRECATED: Use _generate_label_sheet_docx instead for Word document output.

    Args:
        parcels: List of parcel objects with owner_name, mailing_address, etc.
        format_spec: Dictionary with label dimensions and layout
            - cols: number of columns
            - rows: number of rows per page
            - width: label width in inches
            - height: label height in inches
            - margin_top: top margin in inches
            - margin_left: left margin in inches
            - gutter_h: horizontal gutter (space between columns) in inches
            - gutter_v: vertical gutter (space between rows) in inches

    Returns:
        PDF bytes
    """
    # Convert inches to points (72 points per inch)
    cols = format_spec["cols"]
    rows = format_spec["rows"]
    label_width = format_spec["width"] * 72
    label_height = format_spec["height"] * 72
    margin_top = format_spec["margin_top"] * 72
    margin_left = format_spec["margin_left"] * 72
    gutter_h = format_spec["gutter_h"] * 72
    gutter_v = format_spec["gutter_v"] * 72

    labels_per_page = cols * rows
    pdf_pages = []

    # Process parcels in batches of labels_per_page
    for page_start in range(0, len(parcels), labels_per_page):
        page_parcels = parcels[page_start:page_start + labels_per_page]
        page_content = []

        page_content.append("BT")
        page_content.append("/F1 9 Tf")  # 9pt font for labels

        for idx, parcel in enumerate(page_parcels):
            row = idx // cols
            col = idx % cols

            # Calculate label position
            x = margin_left + col * (label_width + gutter_h) + 4  # 4pt padding
            y = 792 - margin_top - row * (label_height + gutter_v) - 12  # Start from top, 12pt down

            # Format the mailing address
            owner_name = getattr(parcel, 'owner_name', '') or ''
            mailing_address = getattr(parcel, 'mailing_address', '') or ''
            mailing_city = getattr(parcel, 'mailing_city', '') or ''
            mailing_state = getattr(parcel, 'mailing_state', '') or ''
            mailing_zip = getattr(parcel, 'mailing_zip', '') or ''

            # Build address lines
            address_lines = []
            if owner_name:
                address_lines.append(owner_name[:35])  # Truncate if too long
            if mailing_address:
                address_lines.append(mailing_address[:35])

            city_state_zip = f"{mailing_city}, {mailing_state} {mailing_zip}".strip(", ")
            if city_state_zip:
                address_lines.append(city_state_zip[:35])

            # If no mailing address, use site address as fallback
            if not address_lines or len(address_lines) < 2:
                site_address = getattr(parcel, 'site_address', '') or ''
                site_city = getattr(parcel, 'site_city', '') or ''
                site_zip = getattr(parcel, 'site_zip', '') or ''

                address_lines = []
                if owner_name:
                    address_lines.append(owner_name[:35])
                if site_address:
                    address_lines.append(site_address[:35])

                site_city_zip = f"{site_city}, MA {site_zip}".strip(", ")
                if site_city_zip != ", MA":
                    address_lines.append(site_city_zip[:35])

            # Write address lines to PDF
            for line_idx, line in enumerate(address_lines[:4]):  # Max 4 lines per label
                line_y = y - (line_idx * 11)  # 11pt line spacing
                page_content.append(f"{x:.2f} {line_y:.2f} Td")
                page_content.append(f"({_pdf_escape(line)}) Tj")
                page_content.append(f"{-x:.2f} {-line_y:.2f} Td")  # Reset position

        page_content.append("ET")

        content_stream = "\n".join(page_content).encode("latin-1", "ignore")
        stream_bytes = (
            f"<< /Length {len(content_stream)} >>\n".encode("latin-1")
            + b"stream\n"
            + content_stream
            + b"\nendstream"
        )

        pdf_pages.append({"stream": stream_bytes, "lines": [], "image": None})

    # Use custom PDF builder for labels
    return _build_label_pdf(pdf_pages)


def _build_label_pdf(pages: List[dict]) -> bytes:
    """Build a simple PDF for label sheets."""
    if not pages:
        return b""

    pdf_objects: List[Optional[bytes]] = []

    def add_object(payload: Optional[bytes]) -> int:
        pdf_objects.append(payload)
        return len(pdf_objects)

    catalog_obj_id = add_object(None)
    pages_obj_id = add_object(None)
    font_obj_id = add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_obj_ids: List[int] = []

    for page in pages:
        stream_bytes = page.get("stream", b"")
        content_obj_id = add_object(stream_bytes)

        page_obj_bytes = (
            f"<< /Type /Page /Parent {pages_obj_id} 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_obj_id} 0 R >> >> /Contents [{content_obj_id} 0 R] >>"
        ).encode("latin-1")
        page_obj_id = add_object(page_obj_bytes)
        page_obj_ids.append(page_obj_id)

    kids = " ".join(f"{obj_id} 0 R" for obj_id in page_obj_ids)
    pages_obj = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_obj_ids)} >>".encode("latin-1")
    catalog_obj = f"<< /Type /Catalog /Pages {pages_obj_id} 0 R >>".encode("latin-1")

    pdf_objects[catalog_obj_id - 1] = catalog_obj
    pdf_objects[pages_obj_id - 1] = pages_obj

    xref_offsets = [0]
    output = BytesIO()
    output.write(b"%PDF-1.4\n")

    for obj_id, obj_data in enumerate(pdf_objects, start=1):
        xref_offsets.append(output.tell())
        output.write(f"{obj_id} 0 obj\n".encode("latin-1"))
        if obj_data:
            output.write(obj_data)
        output.write(b"\nendobj\n")

    xref_start = output.tell()
    output.write(b"xref\n")
    output.write(f"0 {len(xref_offsets)}\n".encode("latin-1"))
    for offset in xref_offsets:
        output.write(f"{offset:010d} {'65535' if offset == 0 else '00000'} {'f' if offset == 0 else 'n'} \n".encode("latin-1"))

    output.write(b"trailer\n")
    output.write(f"<< /Size {len(xref_offsets)} /Root {catalog_obj_id} 0 R >>\n".encode("latin-1"))
    output.write(b"startxref\n")
    output.write(f"{xref_start}\n".encode("latin-1"))
    output.write(b"%%EOF\n")

    return output.getvalue()


def _get_skiptrace_cost_per_lookup() -> Optional[Decimal]:
    raw = getattr(settings, "SKIPTRACE_COST_PER_LOOKUP", None)
    if raw in (None, ""):
        return None
    try:
        return Decimal(str(raw))
    except (InvalidOperation, TypeError, ValueError):
        logger.warning("Invalid SKIPTRACE_COST_PER_LOOKUP value: %s", raw)
        return None


def _get_skiptrace_vendor_cost_per_lookup() -> Optional[Decimal]:
    raw = getattr(settings, "SKIPTRACE_VENDOR_COST_PER_LOOKUP", None)
    if raw in (None, ""):
        return None
    try:
        return Decimal(str(raw))
    except (InvalidOperation, TypeError, ValueError):
        logger.warning("Invalid SKIPTRACE_VENDOR_COST_PER_LOOKUP value: %s", raw)
        return None


def _get_skiptrace_markup_amount_per_lookup() -> Optional[Decimal]:
    raw = getattr(settings, "SKIPTRACE_MARKUP_AMOUNT_PER_LOOKUP", None)
    if raw in (None, ""):
        return None
    try:
        return Decimal(str(raw))
    except (InvalidOperation, TypeError, ValueError):
        logger.warning("Invalid SKIPTRACE_MARKUP_AMOUNT_PER_LOOKUP value: %s", raw)
        return None


def _get_processing_fee_rate() -> Decimal:
    raw = getattr(settings, "STRIPE_PROCESSING_FEE_RATE", None)
    try:
        return Decimal(str(raw)) if raw not in (None, "") else Decimal("0")
    except (InvalidOperation, TypeError, ValueError):
        logger.warning("Invalid STRIPE_PROCESSING_FEE_RATE value: %s", raw)
        return Decimal("0")


def _get_processing_fee_fixed() -> Decimal:
    raw = getattr(settings, "STRIPE_PROCESSING_FEE_FIXED", None)
    try:
        return Decimal(str(raw)) if raw not in (None, "") else Decimal("0")
    except (InvalidOperation, TypeError, ValueError):
        logger.warning("Invalid STRIPE_PROCESSING_FEE_FIXED value: %s", raw)
        return Decimal("0")


def _quantize_currency(value: Optional[Decimal]) -> Optional[Decimal]:
    if value is None:
        return None
    try:
        decimal_value = value if isinstance(value, Decimal) else Decimal(str(value))
    except (InvalidOperation, TypeError, ValueError):
        return None
    return decimal_value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _format_percentage(rate: Optional[Decimal]) -> Optional[str]:
    if rate in (None, ""):
        return None
    try:
        decimal_rate = Decimal(str(rate))
    except (InvalidOperation, TypeError, ValueError):
        return None
    percentage = (decimal_rate * Decimal("100")).quantize(
        Decimal("0.01"), rounding=ROUND_HALF_UP
    )
    return f"{percentage}%"


def _build_skiptrace_pricing(count: int) -> dict[str, object]:
    if count <= 0:
        return {}

    cost_per_lookup = _get_skiptrace_cost_per_lookup()
    if cost_per_lookup is None:
        return {}

    vendor_cost_per_lookup = _get_skiptrace_vendor_cost_per_lookup()
    markup_amount_per_lookup = _get_skiptrace_markup_amount_per_lookup()
    processing_fee_rate = _get_processing_fee_rate()
    processing_fee_fixed = _get_processing_fee_fixed()

    count_decimal = Decimal(count)

    base_per_lookup_q = _quantize_currency(cost_per_lookup)
    if base_per_lookup_q is None:
        return {}

    vendor_cost_per_lookup_q = _quantize_currency(vendor_cost_per_lookup)
    markup_amount_per_lookup_q = _quantize_currency(markup_amount_per_lookup)

    base_total = base_per_lookup_q * count_decimal
    base_total_q = _quantize_currency(base_total)

    vendor_total_q = (
        _quantize_currency(vendor_cost_per_lookup_q * count_decimal)
        if vendor_cost_per_lookup_q is not None
        else None
    )
    markup_total_q = (
        _quantize_currency(markup_amount_per_lookup_q * count_decimal)
        if markup_amount_per_lookup_q is not None
        else None
    )

    percentage_processing_fee = base_total * processing_fee_rate
    fixed_processing_fee = processing_fee_fixed if base_total > 0 else Decimal("0")
    processing_fee_total = percentage_processing_fee + fixed_processing_fee
    processing_fee_total_q = _quantize_currency(processing_fee_total)

    grand_total = base_total + processing_fee_total

    # Enforce minimum of $0.20 per property on the final total
    minimum_per_property = Decimal("0.20")
    minimum_total_for_count = minimum_per_property * count_decimal

    # Also enforce absolute minimum of $0.50 for any transaction
    absolute_minimum_total = Decimal("0.50")
    minimum_total = max(minimum_total_for_count, absolute_minimum_total)

    minimum_adjustment_total = Decimal("0")
    if grand_total < minimum_total:
        minimum_adjustment_total = minimum_total - grand_total
        grand_total = minimum_total
        processing_fee_total += minimum_adjustment_total

    grand_total_q = _quantize_currency(grand_total)

    per_lookup_total_precise = (grand_total / count_decimal).quantize(
        Decimal("0.0001"), rounding=ROUND_HALF_UP
    )
    processing_fee_per_lookup_precise = (processing_fee_total / count_decimal).quantize(
        Decimal("0.0001"), rounding=ROUND_HALF_UP
    )
    per_lookup_total_q = _quantize_currency(per_lookup_total_precise)
    processing_fee_per_lookup_q = _quantize_currency(processing_fee_per_lookup_precise)
    minimum_adjustment_total_q = _quantize_currency(minimum_adjustment_total)
    minimum_adjustment_per_lookup_precise = (
        (minimum_adjustment_total / count_decimal).quantize(
            Decimal("0.0001"), rounding=ROUND_HALF_UP
        )
        if minimum_adjustment_total
        else None
    )
    minimum_adjustment_per_lookup_q = (
        _quantize_currency(minimum_adjustment_per_lookup_precise)
        if minimum_adjustment_per_lookup_precise
        else None
    )

    markup_rate_display = _get_skiptrace_markup_percentage()
    processing_rate_display = _format_percentage(processing_fee_rate)

    pricing = {
        "count": count,
        "basePerLookup": (
            str(base_per_lookup_q) if base_per_lookup_q is not None else None
        ),
        "basePerLookupDisplay": _format_usd(base_per_lookup_q),
        "vendorCostPerLookup": (
            str(vendor_cost_per_lookup_q)
            if vendor_cost_per_lookup_q is not None
            else None
        ),
        "vendorCostPerLookupDisplay": _format_usd(vendor_cost_per_lookup_q),
        "markupAmountPerLookup": (
            str(markup_amount_per_lookup_q)
            if markup_amount_per_lookup_q is not None
            else None
        ),
        "markupAmountPerLookupDisplay": _format_usd(markup_amount_per_lookup_q),
        "processingFeeRate": str(processing_fee_rate),
        "processingFeeRateDisplay": processing_rate_display,
        "processingFeeFixed": str(processing_fee_fixed),
        "processingFeeFixedDisplay": _format_usd(processing_fee_fixed),
        "processingFeeAmount": (
            str(processing_fee_total_q) if processing_fee_total_q is not None else None
        ),
        "processingFeeAmountDisplay": _format_usd(processing_fee_total_q),
        "processingFeePerLookup": (
            str(processing_fee_per_lookup_precise)
            if processing_fee_per_lookup_precise is not None
            else None
        ),
        "processingFeePerLookupDisplay": _format_usd_precise(
            processing_fee_per_lookup_precise
        ),
        "totalCost": str(grand_total_q) if grand_total_q is not None else None,
        "totalCostDisplay": _format_usd(grand_total_q),
        "perLookupTotal": (
            str(per_lookup_total_precise)
            if per_lookup_total_precise is not None
            else None
        ),
        "perLookupTotalDisplay": _format_usd_precise(per_lookup_total_precise),
        "baseTotalCost": str(base_total_q) if base_total_q is not None else None,
        "baseTotalCostDisplay": _format_usd(base_total_q),
        "vendorTotalCost": str(vendor_total_q) if vendor_total_q is not None else None,
        "vendorTotalCostDisplay": _format_usd(vendor_total_q),
        "markupTotalCost": str(markup_total_q) if markup_total_q is not None else None,
        "markupTotalCostDisplay": _format_usd(markup_total_q),
        "minimumAdjustmentTotal": (
            str(minimum_adjustment_total_q)
            if minimum_adjustment_total_q is not None
            else None
        ),
        "minimumAdjustmentTotalDisplay": _format_usd(minimum_adjustment_total_q),
        "minimumAdjustmentPerLookup": (
            str(minimum_adjustment_per_lookup_precise)
            if minimum_adjustment_per_lookup_precise is not None
            else None
        ),
        "minimumAdjustmentPerLookupDisplay": _format_usd_precise(
            minimum_adjustment_per_lookup_precise
        ),
        "markupRateDisplay": markup_rate_display,
    }
    return pricing


def _format_usd(amount: Optional[Decimal]) -> Optional[str]:
    if amount is None:
        return None
    try:
        value = Decimal(str(amount))
    except (InvalidOperation, TypeError, ValueError):
        return None
    quantized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"${quantized:,.2f}"


def _get_skiptrace_markup_percentage() -> Optional[str]:
    raw = getattr(settings, "SKIPTRACE_MARKUP_RATE", None)
    if raw in (None, ""):
        return None
    try:
        percentage = (Decimal(str(raw)) * Decimal("100")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
    except (InvalidOperation, TypeError, ValueError):
        return None
    return f"{percentage}%"


def _format_usd_precise(amount: Optional[Decimal]) -> Optional[str]:
    if amount is None:
        return None
    try:
        value = amount if isinstance(amount, Decimal) else Decimal(str(amount))
    except (InvalidOperation, TypeError, ValueError):
        return None
    normalized = value.normalize()
    text = format(normalized, "f")
    return f"${text}"


def _stripe_configured() -> bool:
    if stripe is None:
        return False
    return bool(getattr(settings, "STRIPE_SECRET_KEY", "")) and bool(
        getattr(settings, "STRIPE_PUBLISHABLE_KEY", "")
    )


def _create_payment_intent_for_pricing(
    pricing: dict, *, metadata: Optional[dict] = None
):
    total_cost = pricing.get("totalCost")
    if total_cost is None:
        raise ValueError("Pricing total is unavailable.")
    try:
        amount_decimal = Decimal(str(total_cost))
    except (InvalidOperation, TypeError, ValueError) as exc:
        raise ValueError("Invalid pricing total.") from exc
    amount_cents = int(
        (amount_decimal * 100).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    )
    if amount_cents <= 0:
        raise ValueError("Payment amount must be greater than zero.")

    secret_key = getattr(settings, "STRIPE_SECRET_KEY", "")
    if not secret_key or stripe is None:
        raise ValueError("Stripe secret key is not configured.")

    stripe.api_key = secret_key
    intent = stripe.PaymentIntent.create(
        amount=amount_cents,
        currency="usd",
        automatic_payment_methods={"enabled": True},
        metadata=metadata or {},
    )
    return intent


def _compose_full_address(lead) -> Optional[str]:
    parts = [
        lead.site_address,
        lead.site_city,
    ]
    state = getattr(lead, "site_state", None) or getattr(lead, "owner_state", None)
    if state:
        parts.append(state)
    if lead.site_zip:
        parts.append(lead.site_zip)

    cleaned = [str(value).strip() for value in parts if value and str(value).strip()]
    return ", ".join(cleaned) if cleaned else None


def _format_currency(value: Optional[object]) -> Optional[str]:
    if value in (None, "", "0", 0):
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isfinite(number):
        if abs(number - int(number)) < 0.01:
            return f"${int(round(number)):,}"
        return f"${number:,.2f}"
    return None


def _format_timestamp(value: Optional[datetime]) -> Optional[str]:
    if not value:
        return None
    localized = timezone.localtime(value)
    return localized.strftime("%b %d, %Y %I:%M %p")


def _extract_bed_bath_from_attrs(
    attrs: dict,
) -> tuple[Optional[object], Optional[object]]:
    if attrs is None:
        return None, None

    def _first(keys):
        for key in keys:
            if key in attrs:
                value = attrs.get(key)
                if value not in (None, "", " "):
                    return value
        return None

    bed_keys = [
        "BEDS",
        "BEDROOMS",
        "NUM_BEDROOMS",
        "NUMBED",
        "BEDRM",
        "BEDS_CNT",
    ]
    bath_keys = [
        "BATHS",
        "BATHS_FULL",
        "TOTAL_BATHS",
        "FULL_BATH",
        "BATHROOMS",
        "BATHS_TOT",
    ]

    beds = _first(bed_keys)
    baths = _first(bath_keys)
    return beds, baths


# --- Parcel search landing page: handles filters, MassGIS lookups, CSV export.
@login_required
def parcel_search_home(request):
    results = []
    town = None
    total_matches = 0
    criteria_payload = None
    save_form = None
    cleaned = {}
    search_performed = bool(request.GET)
    download_requested = request.GET.get("download") == "csv"

    form = ParcelSearchForm(request.GET or None)

    default_radius_meta = {
        "radius_requested": False,
        "radius_center_found": False,
        "radius_center_source": None,
        "radius_excluded_count": 0,
    }
    context_radius_meta = default_radius_meta.copy()

    boundary_shape_filter = _parse_boundary_shape(request.GET)
    boundary_shape_serialized = _serialize_shape_filter(boundary_shape_filter)

    if request.GET and form.is_valid():
        cleaned = form.cleaned_data
        limit = cleaned.get("limit") or None  # No limit by default
        equity_min = cleaned.get("equity_min")
        if equity_min is not None:
            try:
                equity_min = float(equity_min)
            except (TypeError, ValueError):
                equity_min = None

        min_price = cleaned.get("min_price")
        if min_price is not None:
            try:
                min_price = float(min_price)
            except (TypeError, ValueError):
                min_price = None

        max_price = cleaned.get("max_price")
        if max_price is not None:
            try:
                max_price = float(max_price)
            except (TypeError, ValueError):
                max_price = None

        min_years_owned = cleaned.get("min_years_owned")
        if min_years_owned is not None:
            try:
                min_years_owned = int(min_years_owned)
            except (TypeError, ValueError):
                min_years_owned = None

        max_years_owned = cleaned.get("max_years_owned")
        if max_years_owned is not None:
            try:
                max_years_owned = int(max_years_owned)
            except (TypeError, ValueError):
                max_years_owned = None

        proximity_address = (cleaned.get("proximity_address") or "").strip()
        proximity_radius = cleaned.get("proximity_radius_miles")
        if proximity_radius is not None:
            try:
                proximity_radius = float(proximity_radius)
                if proximity_radius < 0:
                    proximity_radius = None
            except (TypeError, ValueError):
                proximity_radius = None
        if not proximity_address:
            proximity_address = None

        try:
            town_id = int(cleaned["town_id"])
            town, results, total_matches, search_meta = search_massgis_parcels(
                town_id=town_id,
                property_category=cleaned.get("property_category") or "any",
                commercial_subtype=cleaned.get("commercial_subtype") or "any",
                address_contains=cleaned.get("address_contains", ""),
                style_contains=cleaned.get("style", ""),
                property_type=cleaned.get("property_type") or "any",
                equity_min=equity_min,
                absentee=cleaned.get("absentee") or "any",
                min_price=min_price,
                max_price=max_price,
                min_years_owned=min_years_owned,
                max_years_owned=max_years_owned,
                proximity_address=proximity_address,
                proximity_radius_miles=proximity_radius,
                limit=limit,
                shape_filter=boundary_shape_filter,
            )
        except MassGISDataError as exc:
            messages.error(request, str(exc))
            context_radius_meta = default_radius_meta.copy()
        except Exception as exc:  # noqa: BLE001
            logger.exception("MassGIS search failed", exc_info=exc)
            messages.error(
                request,
                "We couldn't complete the search right now. Please try again shortly.",
            )
            context_radius_meta = default_radius_meta.copy()
        else:
            if results:
                criteria_payload = {
                    "town_id": town.town_id,
                    "town_name": town.name,
                    "property_category": cleaned.get("property_category") or "any",
                    "commercial_subtype": cleaned.get("commercial_subtype") or "any",
                    "address_contains": cleaned.get("address_contains", ""),
                    "style": cleaned.get("style", ""),
                    "property_type": cleaned.get("property_type", "any"),
                    "equity_min": equity_min,
                    "absentee": cleaned.get("absentee") or "any",
                    "min_price": min_price,
                    "max_price": max_price,
                    "min_years_owned": min_years_owned,
                    "max_years_owned": max_years_owned,
                    "proximity_address": proximity_address,
                    "proximity_radius_miles": proximity_radius,
                    "limit": limit,
                    "boundary_shape": boundary_shape_serialized,
                    "total_matches": total_matches,
                }
                save_form = ParcelListSaveForm(
                    initial={
                        "town_id": town.town_id,
                        "criteria": json.dumps(criteria_payload),
                        "loc_ids": ",".join(parcel.loc_id for parcel in results),
                    }
                )
            else:
                save_form = None
            context_radius_meta = search_meta
            if results and download_requested:
                loc_ids = [parcel.loc_id for parcel in results if parcel.loc_id]
                skiptrace_map = (
                    _bulk_skiptrace_record_map(
                        loc_ids, town_id=town.town_id, user=request.user
                    )
                    if loc_ids
                    else {}
                )
                filename_parts = [town.name if town else "search-results"]
                address_fragment = cleaned.get("address_contains")
                if address_fragment:
                    filename_parts.append(address_fragment)
                filename = "-".join(part for part in filename_parts if part)
                return _build_parcel_csv_response(
                    results,
                    skiptrace_map,
                    filename=filename or "search-results",
                    default_city=town.name if town else None,
                )
    elif request.GET and not form.is_valid():
        messages.error(request, "Please correct the errors below and try again.")
        context_radius_meta = default_radius_meta.copy()
    else:
        context_radius_meta = default_radius_meta.copy()

    csv_download_query = ""
    if results:
        query_copy = request.GET.copy()
        query_copy["download"] = "csv"
        csv_download_query = query_copy.urlencode()

    # Boston dataset preload disabled - using local file instead
    # Data will be loaded on-demand when a neighborhood is selected

    boundary_polygon_json = ""
    if boundary_shape_serialized and boundary_shape_serialized.get("type") == "polygon":
        boundary_polygon_json = json.dumps(boundary_shape_serialized.get("coordinates") or [])

    context = {
        "form": form,
        "results": results,
        "town": town,
        "total_matches": total_matches,
        "criteria_payload": criteria_payload,
        "save_form": save_form,
        "search_performed": search_performed,
        "display_limit": cleaned.get("limit") or form.fields["limit"].initial,
        "equity_min": cleaned.get("equity_min") if cleaned else None,
        "style_filter": cleaned.get("style") if cleaned else "",
        "property_type": cleaned.get("property_type") if cleaned else "any",
        "min_price": cleaned.get("min_price") if cleaned else None,
        "max_price": cleaned.get("max_price") if cleaned else None,
        "min_years_owned": cleaned.get("min_years_owned") if cleaned else None,
        "max_years_owned": cleaned.get("max_years_owned") if cleaned else None,
        "proximity_address": cleaned.get("proximity_address") if cleaned else None,
        "proximity_radius_miles": (
            cleaned.get("proximity_radius_miles") if cleaned else None
        ),
        "boundary_shape": boundary_shape_serialized or {},
        "boundary_polygon_json": boundary_polygon_json,
        "radius_meta": context_radius_meta,
        "csv_download_query": csv_download_query,
        "initial_auto_lien_search_enabled": bool(results) and len(results) <= LIEN_SEARCH_AUTO_THRESHOLD,
    }
    return render(request, "leads/parcel_search.html", context)


@login_required
def parcel_search_detail(request, town_id, loc_id, list_id=None):
    # Build navigation context if viewing from a saved list
    nav_context = {}
    if list_id:
        try:
            saved_list = get_object_or_404(
                _saved_list_queryset_for_user(request.user), pk=list_id
            )
            # Get all parcel refs from the list
            parcel_refs = list(_iter_saved_list_parcel_refs(saved_list))

            # Find current parcel index
            current_index = None
            for idx, ref in enumerate(parcel_refs):
                if ref.town_id == town_id and _normalize_loc_id(ref.loc_id) == _normalize_loc_id(loc_id):
                    current_index = idx
                    break

            if current_index is not None:
                nav_context = {
                    "from_list": True,
                    "list_id": list_id,
                    "list_name": saved_list.name,
                    "current_index": current_index,
                    "total_parcels": len(parcel_refs),
                }

                # Add previous parcel info
                if current_index > 0:
                    prev_ref = parcel_refs[current_index - 1]
                    nav_context["prev_parcel"] = {
                        "town_id": prev_ref.town_id,
                        "loc_id": prev_ref.loc_id,
                        "url": reverse(
                            "parcel_detail_from_list",
                            args=[prev_ref.town_id, prev_ref.loc_id, list_id],
                        ),
                    }

                # Add next parcel info
                if current_index < len(parcel_refs) - 1:
                    next_ref = parcel_refs[current_index + 1]
                    nav_context["next_parcel"] = {
                        "town_id": next_ref.town_id,
                        "loc_id": next_ref.loc_id,
                        "url": reverse(
                            "parcel_detail_from_list",
                            args=[next_ref.town_id, next_ref.loc_id, list_id],
                        ),
                    }

                # Add list URL
                nav_context["list_url"] = reverse("saved_parcel_list_detail", args=[list_id])
        except Exception:  # noqa: BLE001
            # If we can't load the list, just proceed without navigation
            pass

    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        messages.error(request, str(exc))
        return redirect("parcel_search")

    parcel_shape = get_massgis_parcel_shape(parcel.town, parcel.loc_id)

    attrs = parcel.attributes
    centroid_lon_lat: Optional[Tuple[float, float]] = None
    if parcel_shape and parcel_shape.centroid:
        centroid = parcel_shape.centroid
        converted = massgis_stateplane_to_wgs84(centroid[0], centroid[1])
        if converted:
            centroid_lon_lat = converted

    def _lot_size_display() -> Optional[str]:
        if parcel.lot_size is None:
            return None
        unit = attrs.get("LOT_UNITS")
        if unit:
            return f"{parcel.lot_size:,.2f} {unit}"
        return f"{parcel.lot_size:,.2f}"

    def _book_page_display() -> Optional[str]:
        book = attrs.get("LS_BOOK")
        page = attrs.get("LS_PAGE")
        parts = [part for part in [book, page] if part]
        return " / ".join(parts) if parts else None

    def _format_gis_date(date_value: Optional[object]) -> Optional[str]:
        """Format GIS date to match ATTOM date format (YYYY-MM-DD) for consistency."""
        if not date_value:
            return None
        from datetime import datetime
        date_str = str(date_value).strip()
        # Try multiple formats that GIS dates might be in
        for fmt in ["%Y-%m-%d", "%m/%d/%Y", "%Y%m%d", "%m/%d/%y", "%Y"]:
            try:
                parsed_date = datetime.strptime(date_str, fmt)
                if fmt == "%Y":
                    # If only year, use January 1st
                    parsed_date = parsed_date.replace(month=1, day=1)
                # Return in standardized YYYY-MM-DD format
                return parsed_date.strftime("%Y-%m-%d")
            except ValueError:
                continue
        # If no format matches, return original
        return date_str

    beds, baths = _extract_bed_bath_from_attrs(attrs)
    beds_display = _format_bed_bath(beds)
    baths_display = _format_bed_bath(baths)

    market_value_display = _format_currency(parcel.market_value)
    market_value_psf_display = (
        f"${parcel.market_value_per_sqft:,.0f}"
        if parcel.market_value_per_sqft
        else None
    )
    comparable_psf_display = (
        f"${parcel.market_value_comparable_avg_psf:,.0f}"
        if parcel.market_value_comparable_avg_psf
        else None
    )
    market_value_updated_display = _format_timestamp(parcel.market_value_updated_at)
    market_value_confidence_display = (
        f"{parcel.market_value_confidence * 100:.0f}%"
        if parcel.market_value_confidence is not None
        else None
    )
    market_value_method_display = (
        parcel.market_value_methodology_label
        or parcel.market_value_methodology
    )
    market_value_comp_sample_display = (
        f"{parcel.market_value_comparable_count} comps"
        if parcel.market_value_comparable_count
        else None
    )
    market_value_comp_value_display = _format_currency(parcel.market_value_comparable_value)

    sections = [
        {
            "title": "Property Overview",
            "items": [
                ("Parcel ID", parcel.loc_id),
                ("Town", parcel.town.name),
                ("Category", parcel.property_category),
                ("Use Code", parcel.use_code),
                ("Style", parcel.style),
                ("Units", parcel.units),
                ("Bedrooms", beds_display),
                ("Bathrooms", baths_display),
            ],
        },
        {
            "title": "Location",
            "items": [
                ("Address", parcel.site_address),
                ("City", parcel.site_city),
                ("ZIP", parcel.site_zip),
                ("Town ID", attrs.get("TOWN_ID")),
            ],
        },
        {
            "title": "Valuation",
            "items": [
                ("Est. Market Value", market_value_display),
                ("Market Value $/sf", market_value_psf_display),
                ("Market Value Updated", market_value_updated_display),
                ("Market Value Confidence", market_value_confidence_display),
                ("Valuation Method", market_value_method_display),
                ("Comparable Sample", market_value_comp_sample_display),
                ("Average Comp Value", market_value_comp_value_display),
                ("Comparable Avg $/sf", comparable_psf_display),
                ("Total Value", _format_currency(parcel.total_value)),
                ("Land Value", _format_currency(attrs.get("LAND_VAL"))),
                ("Building Value", _format_currency(attrs.get("BLDG_VAL"))),
                ("Building Area (sqft)", attrs.get("BLD_AREA")),
                (
                    "Est. Mortgage Balance",
                    _format_currency(parcel.estimated_mortgage_balance),
                ),
                (
                    "Est. Equity",
                    _format_currency(parcel.estimated_equity_value),
                ),
                (
                    "Equity %",
                    (
                        f"{parcel.equity_percent:.1f}%"
                        if parcel.equity_percent is not None
                        else None
                    ),
                ),
                (
                    "Est. ROI %",
                    (
                        f"{parcel.estimated_roi_percent:.1f}%"
                        if parcel.estimated_roi_percent is not None
                        else None
                    ),
                ),
            ],
        },
        {
            "title": "Sale History",
            "items": [
                ("Sale Date (GIS)", _format_gis_date(attrs.get("LS_DATE"))),
                ("Sale Price", _format_currency(attrs.get("LS_PRICE"))),
                ("Book / Page", _book_page_display()),
            ],
        },
        {
            "title": "Owner & Mailing",
            "items": [
                ("Unit", None),
                ("Owner", parcel.owner_name),
                ("Mailing Address", parcel.owner_address),
                ("Site Address", parcel.site_address),
                ("Absentee Owner", "Yes" if parcel.absentee else "No"),
            ],
            "corporate_entity": None,  # Will be populated if LLC/Corp detected
        },
        {
            "title": "Skip Trace",
            "items": [],
        },
    ]

    owner_section_items: list[tuple[str, object]] = []
    for section in sections:
        if section["title"] == "Owner & Mailing":
            owner_section_items = list(section["items"])
            break

    owner_labels = [label for label, _ in owner_section_items]
    if "Unit" not in owner_labels:
        owner_section_items = [("Unit", None)] + owner_section_items
        owner_labels = [label for label, _ in owner_section_items]

    units_detail = []
    raw_units = list(parcel.units_detail or [])
    for index, unit in enumerate(raw_units):
        if not isinstance(unit, dict):
            continue
        row_key_source = unit.get("row_key") or unit.get("id")
        if row_key_source in {None, ""}:
            row_key = str(index)
        else:
            row_key = str(row_key_source)
        unit["row_key"] = row_key

        overrides = unit.pop("owner_overrides", {}) or {}
        owner_rows: list[tuple[str, object]] = []
        for label, base_value in owner_section_items:
            if label == "Record Type":
                continue
            value = overrides.get(label)
            if not _is_displayable(value):
                value = base_value
            owner_rows.append((label, value))

        if unit.get("is_master_record"):
            owner_rows.append(("Record Type", "Master"))
        unit["owner_items"] = owner_rows
        units_detail.append(unit)

    def _filter_items(items: Iterable[tuple[str, object]]) -> list[tuple[str, object]]:
        return [
            (label, value)
            for label, value in items
            if _is_displayable(value)
        ]

    selected_unit_detail = next(
        (
            unit
            for unit in units_detail
            if not unit.get("is_master_record")
            and (
                (unit.get("valuation_items") and any(_is_displayable(v) for _, v in unit.get("valuation_items")))
                or (unit.get("sale_history_items") and any(_is_displayable(v) for _, v in unit.get("sale_history_items")))
            )
        ),
        None,
    )
    if not selected_unit_detail and units_detail:
        selected_unit_detail = units_detail[0]

    if selected_unit_detail:
        valuation_items = selected_unit_detail.get("valuation_items") or []
        sale_history_items = selected_unit_detail.get("sale_history_items") or []
        owner_items = selected_unit_detail.get("owner_items") or []
        for section in sections:
            if section["title"] == "Valuation" and valuation_items:
                # Always preserve Total Value, even if it's $0
                total_value_item = next(
                    (item for item in valuation_items if item[0] == "Total Value"),
                    None
                )
                filtered = _filter_items(valuation_items)
                if filtered or total_value_item:
                    # Add Total Value back if it was filtered out
                    if total_value_item and not any(item[0] == "Total Value" for item in filtered):
                        section["items"] = [total_value_item] + filtered
                    else:
                        section["items"] = filtered
            elif section["title"] == "Sale History" and sale_history_items:
                filtered = _filter_items(sale_history_items)
                if filtered:
                    section["items"] = filtered
            elif section["title"] == "Owner & Mailing" and owner_items:
                owner_rows: list[tuple[str, object]] = []
                for label, value in owner_items:
                    display = value if _is_displayable(value) else "—"
                    owner_rows.append((label, display))
                if parcel.absentee is not None:
                    owner_rows = [
                        row for row in owner_rows if row[0] != "Absentee Owner"
                    ]
                    owner_rows.append(
                        (
                            "Absentee Owner",
                            "Yes" if parcel.absentee else "No",
                        )
                    )
                if selected_unit_detail.get("is_master_record") and not any(
                    label == "Record Type" for label, _ in owner_rows
                ):
                    owner_rows.append(("Record Type", "Master"))
                if owner_rows:
                    section["items"] = owner_rows

    # Don't filter sections yet - we need to add ATTOM data first
    # (filtering will happen after ATTOM data is added)

    google_maps = None
    full_address = _compose_full_address(parcel)
    lat = None
    lng = None
    if centroid_lon_lat:
        lng, lat = centroid_lon_lat
    elif full_address:
        coords = geocode_address(full_address)
        if coords:
            lng, lat = coords
    zillow_url = _build_zillow_url(full_address)
    property_embed_url = _build_mapillary_embed_url(lat, lng)
    if not property_embed_url:
        property_embed_url = _build_google_street_view_embed(lat, lng)

    hero_image_url = (
        _build_google_street_view_image_url(lat, lng)
        if lat is not None and lng is not None
        else None
    )

    mailer_download_base = reverse(
        "mailer_download_pdf", args=[parcel.town.town_id, parcel.loc_id]
    )

    record = _get_skiptrace_record_for_loc_id(
        parcel.town.town_id, parcel.loc_id, user=request.user
    )

    mailer_context = _build_mailer_context(
        parcel,
        full_address=full_address,
        zillow_url=zillow_url,
        hero_image_url=hero_image_url,
        download_endpoint=mailer_download_base,
        request=request,
        skiptrace_record=record,
        user=request.user,
    )
    record_is_fresh = _skiptrace_record_is_fresh(record) if record else False
    saved_match = _saved_list_contains_loc_id(
        parcel.town.town_id, parcel.loc_id, user=request.user
    )
    lead_match = _lead_queryset_for_user(request.user).filter(
        loc_id__iexact=parcel.loc_id
    ).exists()
    can_run_skiptrace = _skiptrace_allowed_for_parcel(
        parcel.town.town_id, parcel.loc_id, user=request.user
    )
    skiptrace_endpoint = (
        reverse("parcel_skiptrace", args=[parcel.town.town_id, parcel.loc_id])
        if can_run_skiptrace
        else None
    )

    if lat is not None and lng is not None:
        google_maps = {
            "address": full_address,
            "lat": lat,
            "lng": lng,
        }
    elif full_address:
        google_maps = {
            "address": full_address,
            "lat": None,
            "lng": None,
        }

    if record:
        status_message = f"Last run {timezone.localtime(record.updated_at).strftime('%b %d, %Y %I:%M %p')}"
        if not record_is_fresh:
            ttl = _skiptrace_cache_ttl_days_value()
            if ttl:
                status_message += f" · Cached result older than {ttl} days."
            else:
                status_message += " · Cached result may be out of date."
        if not can_run_skiptrace:
            status_message += " · Save for skip-tracing capabilities."
        initial_payload = _serialize_skiptrace_record(record)
        button_label = (
            "Refresh Skip Trace" if can_run_skiptrace else "Refresh Skip Trace"
        )
    elif not can_run_skiptrace:
        status_message = "Save for skip-tracing capabilities."
        initial_payload = None
        button_label = "Run Skip Trace"
    else:
        status_message = 'Click "Run Skip Trace" to retrieve phone numbers and emails.'
        initial_payload = None
        button_label = "Run Skip Trace"

    skiptrace_items = [("Status", status_message)]

    for section in sections:
        if section["title"] == "Skip Trace":
            section["items"] = skiptrace_items
            break

    pricing = _build_skiptrace_pricing(1)
    stripe_enabled = _stripe_configured()
    payment_config = None
    if stripe_enabled:
        payment_payload = {
            "context": "parcel",
            "townId": parcel.town.town_id,
            "locId": parcel.loc_id,
        }
        payment_config = {
            "enabled": True,
            "mode": "single",
            "createIntentEndpoint": reverse("skiptrace_payment_single"),
            "cardElementId": "skiptrace-card-element",
            "cardErrorId": "skiptrace-card-error",
            "intentPayload": payment_payload,
        }

    skiptrace_config = {
        "allowed": can_run_skiptrace,
        "endpoint": skiptrace_endpoint,
        "buttonLabel": button_label,
        "initial": initial_payload,
        "statusMessage": status_message,
        "savedMatch": saved_match,
        "leadMatch": lead_match,
        "pricing": pricing,
        "payment": payment_config,
    }
    skiptrace_config_json = json.dumps(skiptrace_config, ensure_ascii=False)

    # Check if owner is an LLC/Corporation and look up actual owner
    corporate_entity = None
    if parcel.owner_name:
        # Detect corporate ownership
        corporate_keywords = ['LLC', 'L.L.C.', 'Inc.', 'Inc', 'Corp.', 'Corp', 'LLP', 'L.L.P.', 'Corporation', 'Company', 'Trust']
        owner_name_upper = parcel.owner_name.upper()
        is_corporate = any(keyword.upper() in owner_name_upper for keyword in corporate_keywords)

        if is_corporate:
            from .models import CorporateEntity
            from data_pipeline.jobs.corporate_job import CorporateJob

            # Check cache first (180 days)
            try:
                corporate_entity = CorporateEntity.objects.filter(
                    entity_name__iexact=parcel.owner_name.strip()
                ).order_by('-last_updated').first()

                # Check if cache is fresh (180 days)
                if corporate_entity and corporate_entity.last_updated:
                    cache_age = timezone.now() - corporate_entity.last_updated
                    if cache_age.days > 180:
                        corporate_entity = None  # Cache is stale

                # If no fresh cache, scrape MA Secretary of Commonwealth
                if not corporate_entity:
                    try:
                        corporate_config = {
                            "id": "ma_secretary",
                            "name": "Massachusetts Secretary of Commonwealth",
                            "adapter": "ma_secretary",
                        }
                        job = CorporateJob(corporate_config)
                        result = job.run(
                            entity_name=parcel.owner_name,
                            dry_run=False,
                            force_refresh=False,
                            max_cache_age_days=180
                        )

                        # Reload from database
                        if result:
                            corporate_entity = CorporateEntity.objects.filter(
                                entity_name__iexact=parcel.owner_name.strip()
                            ).order_by('-last_updated').first()
                    except Exception as e:
                        logger.warning(f"Failed to look up LLC owner for '{parcel.owner_name}': {e}")
            except Exception as e:
                logger.warning(f"Error checking corporate entity cache: {e}")

    # Check for existing AttomData (keep old data, but don't fetch new from ATTOM API)
    from datetime import timedelta

    # Check for any existing AttomData (no age limit - keep all existing data)
    attom_data = AttomData.objects.filter(
        town_id=town_id,
        loc_id=loc_id,
    ).order_by('-last_updated').first()

    # Check if we need to trigger scraping:
    # 1. No AttomData at all, OR
    # 2. Has AttomData but no scraped documents yet
    has_scraped_docs = False
    if attom_data and attom_data.raw_response:
        has_scraped_docs = bool(attom_data.raw_response.get('scrape_sources'))

    needs_scraping = not attom_data or not has_scraped_docs

    # Log scraping decision
    if not attom_data:
        logger.info(f"No AttomData found for {loc_id}, will trigger scraping")
    elif not has_scraped_docs:
        logger.info(f"AttomData exists for {loc_id} but no scrape_sources, will trigger scraping")
    else:
        logger.info(f"AttomData with scrape_sources exists for {loc_id}, skipping scraping")

    # Trigger background scraping from Registry of Deeds if needed
    if needs_scraping:
        try:
            from data_pipeline.town_registry_map import get_registry_for_town
            from data_pipeline.jobs.task_queue import run_registry_task
            registry_id = get_registry_for_town(town_id)
            if registry_id:
                # Load registry config
                import os
                from pathlib import Path
                config_path = Path(__file__).resolve().parent.parent / "data_pipeline" / "config" / "sources.json"
                with open(config_path) as f:
                    sources_config = json.load(f)

                # Find the registry config
                registry_config = next(
                    (reg for reg in sources_config.get("registries", []) if reg["id"] == registry_id),
                    None
                )

                if registry_config:
                    # Trigger async scraping in background (won't block page load)
                    # Pass owner name for registry search
                    owner_name = parcel.owner_name if hasattr(parcel, 'owner_name') else None
                    run_registry_task.delay(
                        config=registry_config,
                        owner=owner_name,
                        loc_id=loc_id,
                        force_refresh=False,
                        max_cache_age_days=90,
                    )
                    logger.info(f"Triggered background registry scraping for {loc_id} (owner: {owner_name}) in {registry_id}")
            else:
                logger.warning(f"No registry mapping found for town {town_id}")
        except Exception as e:
            logger.warning(f"Could not trigger background registry scraping: {e}")

    # Calculate mortgage and equity using ATTOM data if available
    mortgage_balance = equity_value = equity_percent = roi_percent = monthly_payment = None
    has_mortgage_data = bool(attom_data and attom_data.mortgage_loan_amount)

    if has_mortgage_data:
        mortgage_balance, equity_value, equity_percent, roi_percent, monthly_payment = _calculate_mortgage_balance_from_attom(
            attom_data, parcel.total_value
        )

    # Update sections with ATTOM data if available (mortgage and/or tax data)
    if attom_data:
        print(f"[DEBUG] ATTOM data available for parcel {loc_id}")
        print(f"[DEBUG] Has mortgage: {has_mortgage_data}")
        print(f"[DEBUG] Has tax data: {bool(attom_data.tax_assessed_value or attom_data.tax_amount_annual)}")
        print(f"[DEBUG] Has foreclosure: {bool(attom_data.pre_foreclosure or attom_data.foreclosure_stage)}")
        print(f"[DEBUG] Has propensity score: {bool(attom_data.propensity_to_default_score)}")
        print(f"[DEBUG] Mortgage amount: {attom_data.mortgage_loan_amount}")
        print(f"[DEBUG] Tax assessed value: {attom_data.tax_assessed_value}")
        print(f"[DEBUG] Tax amount annual: {attom_data.tax_amount_annual}")

        # Update Sale History section with ATTOM mortgage and tax data
        for section in sections:
            if section["title"] == "Sale History":
                attom_items_added = False

                # Check if we have any useful ATTOM data to display
                has_any_attom_data = (
                    has_mortgage_data or
                    attom_data.tax_assessed_value or attom_data.tax_amount_annual or attom_data.tax_assessment_year or
                    attom_data.pre_foreclosure or attom_data.foreclosure_recording_date or attom_data.foreclosure_stage or attom_data.foreclosure_estimated_value or
                    attom_data.propensity_to_default_score or attom_data.propensity_to_default_decile
                )

                # Only show "ATTOM DATA AVAILABLE" if we actually have data to display
                if has_any_attom_data:
                    section["items"].extend([
                        ("─" * 30, ""),
                        ("ATTOM DATA AVAILABLE", "✓"),
                    ])
                    attom_items_added = True

                # Add foreclosure information if available
                if (attom_data.pre_foreclosure or attom_data.foreclosure_recording_date or
                    attom_data.foreclosure_stage or attom_data.foreclosure_estimated_value):
                    section["items"].extend([
                        ("─" * 30, ""),  # Divider
                        ("FORECLOSURE DATA", ""),
                    ])
                    attom_items_added = True

                    if attom_data.foreclosure_stage:
                        section["items"].append(("Foreclosure Stage", attom_data.foreclosure_stage))
                    if attom_data.foreclosure_recording_date:
                        section["items"].append(("Foreclosure Recording Date", attom_data.foreclosure_recording_date))
                    if attom_data.foreclosure_auction_date:
                        section["items"].append(("Foreclosure Auction Date", attom_data.foreclosure_auction_date))
                    if attom_data.foreclosure_estimated_value:
                        section["items"].append(("Foreclosure Est. Value", _format_currency(float(attom_data.foreclosure_estimated_value))))
                    if attom_data.foreclosure_judgment_amount:
                        section["items"].append(("Foreclosure Judgment", _format_currency(float(attom_data.foreclosure_judgment_amount))))
                    if attom_data.foreclosure_default_amount:
                        section["items"].append(("Foreclosure Default Amount", _format_currency(float(attom_data.foreclosure_default_amount))))
                    if attom_data.foreclosure_document_type:
                        section["items"].append(("Foreclosure Doc Type", attom_data.foreclosure_document_type))

                # Add mortgage information if available
                if has_mortgage_data and attom_data.mortgage_loan_amount:
                    if attom_items_added:
                        section["items"].append(("─" * 30, ""))
                    else:
                        section["items"].extend([("─" * 30, "")])
                        attom_items_added = True

                    section["items"].append(("MORTGAGE DATA", ""))
                    section["items"].append(("Mortgage Amount", _format_currency(float(attom_data.mortgage_loan_amount))))

                    # Check if mortgage date matches GIS sale date
                    date_mismatch_warning = None
                    if attom_data.mortgage_recording_date:
                        gis_sale_date = attrs.get("LS_DATE")
                        if gis_sale_date:
                            # Normalize both dates for comparison
                            try:
                                from datetime import datetime
                                # Parse ATTOM mortgage date (YYYY-MM-DD format)
                                mortgage_date = datetime.strptime(attom_data.mortgage_recording_date, "%Y-%m-%d")

                                # Parse GIS sale date (multiple formats possible)
                                gis_date = None
                                gis_date_str = str(gis_sale_date).strip()
                                for fmt in ["%Y-%m-%d", "%m/%d/%Y", "%Y%m%d", "%m/%d/%y", "%Y"]:
                                    try:
                                        gis_date = datetime.strptime(gis_date_str, fmt)
                                        if fmt == "%Y":
                                            gis_date = gis_date.replace(month=1, day=1)
                                        break
                                    except ValueError:
                                        continue

                                # Compare dates (allowing for small differences due to recording delays)
                                if gis_date and mortgage_date:
                                    # If dates are more than 90 days apart, show warning
                                    days_diff = abs((mortgage_date - gis_date).days)
                                    if days_diff > 90:
                                        date_mismatch_warning = f"⚠️ Mortgage date differs from GIS sale date by {days_diff} days. ATTOM mortgage may be from a refinance or different transaction."
                            except (ValueError, TypeError):
                                pass

                        section["items"].append(("Mortgage Date", attom_data.mortgage_recording_date))
                        if date_mismatch_warning:
                            section["items"].append(("⚠️ Date Notice", date_mismatch_warning))
                    if attom_data.mortgage_due_date:
                        section["items"].append(("Mortgage Due Date", attom_data.mortgage_due_date))
                    if attom_data.mortgage_lender_name:
                        section["items"].append(("Lender", attom_data.mortgage_lender_name))
                    if attom_data.mortgage_loan_type:
                        section["items"].append(("Loan Type", attom_data.mortgage_loan_type))
                    if attom_data.mortgage_interest_rate:
                        section["items"].append(("Interest Rate (ATTOM)", f"{attom_data.mortgage_interest_rate}%"))
                    elif mortgage_balance:
                        # Interest rate was estimated from historical data
                        from datetime import datetime
                        from .services import MORTGAGE_RATE_BY_YEAR
                        if attom_data.mortgage_recording_date:
                            try:
                                mortgage_year = datetime.strptime(attom_data.mortgage_recording_date, "%Y-%m-%d").year
                                est_rate = MORTGAGE_RATE_BY_YEAR.get(mortgage_year, MORTGAGE_RATE_BY_YEAR.get(max(MORTGAGE_RATE_BY_YEAR.keys())))
                                section["items"].append(("Interest Rate (Est.)", f"{est_rate}% (historical avg for {mortgage_year})"))
                            except:
                                pass
                    if attom_data.mortgage_term_years:
                        section["items"].append(("Loan Term", f"{attom_data.mortgage_term_years} years"))
                    if attom_data.mortgage_loan_number:
                        section["items"].append(("Loan Number", attom_data.mortgage_loan_number))
                    if monthly_payment:
                        section["items"].append(("Monthly Payment (Est.)", _format_currency(monthly_payment)))

                # Add tax information
                if attom_data.tax_assessed_value or attom_data.tax_amount_annual or attom_data.tax_assessment_year:
                    if attom_items_added:
                        section["items"].append(("─" * 30, ""))
                    else:
                        section["items"].extend([("─" * 30, "")])
                        attom_items_added = True

                    section["items"].append(("TAX DATA", ""))

                    if attom_data.tax_assessment_year:
                        section["items"].append(("Tax Year", str(attom_data.tax_assessment_year)))
                    if attom_data.tax_assessed_value:
                        section["items"].append(("Tax Assessed Value", _format_currency(float(attom_data.tax_assessed_value))))
                    if attom_data.tax_amount_annual:
                        section["items"].append(("Annual Tax", _format_currency(float(attom_data.tax_amount_annual))))
                    if attom_data.tax_delinquent_year:
                        section["items"].append(("Tax Delinquent Year", str(attom_data.tax_delinquent_year)))

                # Add propensity to default information
                if attom_data.propensity_to_default_score or attom_data.propensity_to_default_decile:
                    if attom_items_added:
                        section["items"].append(("─" * 30, ""))
                    else:
                        section["items"].extend([("─" * 30, "")])
                        attom_items_added = True

                    section["items"].append(("DEFAULT RISK SCORE", ""))

                    if attom_data.propensity_to_default_score:
                        section["items"].append(("Propensity Score", f"{attom_data.propensity_to_default_score}/100"))
                    if attom_data.propensity_to_default_decile:
                        section["items"].append(("Risk Decile", f"{attom_data.propensity_to_default_decile}/10 (10=highest risk)"))

            # Update Valuation section with ATTOM-based calculations
            elif section["title"] == "Valuation":
                # Replace estimated values with ATTOM-calculated values
                updated_items = []
                for label, value in section["items"]:
                    if label == "Est. Mortgage Balance" and mortgage_balance is not None:
                        updated_items.append(("Mortgage Balance (ATTOM)", _format_currency(mortgage_balance)))
                    elif label == "Est. Equity" and equity_value is not None:
                        updated_items.append(("Equity (ATTOM)", _format_currency(equity_value)))
                    elif label == "Equity %" and equity_percent is not None:
                        updated_items.append(("Equity % (ATTOM)", f"{equity_percent:.1f}%"))
                    elif label == "Est. ROI %" and roi_percent is not None:
                        updated_items.append(("ROI % (ATTOM)", f"{roi_percent:.1f}%"))
                    else:
                        updated_items.append((label, value))
                section["items"] = updated_items

                # Ensure Total Value is always at the top of the valuation section
                total_value_item = next((item for item in updated_items if item[0] == "Total Value"), None)
                if total_value_item:
                    # Remove it from its current position and add it to the beginning
                    updated_items = [item for item in updated_items if item[0] != "Total Value"]
                    updated_items.insert(0, total_value_item)
                    section["items"] = updated_items
    else:
        # No ATTOM data - Sale History section will show MassGIS data only
        print(f"[DEBUG] No ATTOM data available for parcel {loc_id} - showing MassGIS sale data only")
        for section in sections:
            if section["title"] == "Sale History":
                # Add a note that only MassGIS data is available
                section["items"].extend([
                    ("─" * 30, ""),
                    ("Data Source", "MassGIS (ATTOM data not available)"),
                ])

    # Add corporate entity information to Owner & Mailing section if available
    if corporate_entity:
        for section in sections:
            if section["title"] == "Owner & Mailing":
                section["corporate_entity"] = corporate_entity
                # Add LLC owner details after the owner name
                corporate_items = []
                if corporate_entity.principal_name:
                    title_str = f" ({corporate_entity.principal_title})" if corporate_entity.principal_title else ""
                    corporate_items.append(("Actual Owner (LLC)", f"{corporate_entity.principal_name}{title_str}"))
                if corporate_entity.business_phone:
                    corporate_items.append(("Business Phone", corporate_entity.business_phone))
                if corporate_entity.business_address:
                    corporate_items.append(("Business Address", corporate_entity.business_address))
                if corporate_entity.status:
                    corporate_items.append(("Entity Status", corporate_entity.status))
                if corporate_entity.entity_type:
                    corporate_items.append(("Entity Type", corporate_entity.entity_type))

                # Insert corporate items after "Owner" field
                owner_items = section["items"]
                new_items = []
                for label, value in owner_items:
                    new_items.append((label, value))
                    if label == "Owner":
                        # Add divider and corporate items
                        if corporate_items:
                            new_items.append(("─" * 20, ""))
                            new_items.extend(corporate_items)
                            new_items.append(("─" * 20, ""))
                section["items"] = new_items
                break

    # Now filter out empty items and empty sections (after ATTOM data has been added)
    for section in sections:
        if section["title"] == "Valuation":
            total_value_item = next(
                (item for item in section["items"] if item[0] == "Total Value"),
                None
            )
            market_value_item = next(
                (item for item in section["items"] if item[0] == "Est. Market Value"),
                None
            )

            other_items = [
                item for item in section["items"]
                if item[0] not in {"Total Value", "Est. Market Value"} and _is_displayable(item[1])
            ]

            ordered_items: list[tuple[str, object]] = []
            if market_value_item and _is_displayable(market_value_item[1]):
                ordered_items.append(market_value_item)

            if total_value_item:
                formatted_value = total_value_item[1] or "$0"
                ordered_items.append(("Total Value", formatted_value))

            ordered_items.extend(other_items)
            section["items"] = ordered_items
        else:
            section["items"] = [
                item for item in section["items"] if _is_displayable(item[1])
            ]
    sections = [
        section
        for section in sections
        if section["items"] or section["title"] == "Skip Trace"
    ]

    default_unit_key = selected_unit_detail.get("row_key") if selected_unit_detail else None
    for unit in units_detail:
        unit["has_attom"] = bool(attom_data) and unit.get("row_key") == default_unit_key

    # Fetch liens and legal actions for this parcel
    from .models import LienRecord, LegalAction, LienSearchAttempt

    lien_search_requested = False
    lien_search_queued = False
    last_lien_search = None

    enable_lien_search_param = request.GET.get("enable_lien_search", "").lower() in {"1", "true", "yes"}
    manual_search_triggered = request.GET.get("manual_search", "").lower() in {"1", "true", "yes"}

    try:
        needs_search = enable_lien_search_param and should_search_parcel(request.user, town_id, loc_id)
    except Exception as exc:
        logger.warning(
            "Unable to evaluate lien search cache for %s/%s: %s",
            town_id,
            loc_id,
            exc,
        )
        needs_search = False

    if needs_search:
        lien_search_requested = True
        town_name = parcel.town.name if parcel.town else ""
        county = None
        if town_name:
            town_lower = town_name.lower()
            if town_lower in {"salem", "beverly", "peabody", "lynn", "gloucester", "marblehead", "danvers"}:
                county = "Essex"
            elif town_lower in {"boston", "cambridge", "somerville", "brookline", "chelsea", "revere", "winthrop"}:
                county = "Suffolk"
            elif town_lower in {"worcester", "shrewsbury", "westborough", "auburn", "millbury"}:
                county = "Worcester"
            elif town_lower in {"springfield", "chicopee", "holyoke", "westfield"}:
                county = "Hampden"
            elif town_lower in {"lowell", "newton", "framingham", "waltham"}:
                county = "Middlesex"

        parcel_data = {
            "owner_name": parcel.owner_name or "",
            "address": parcel.site_address or "",
            "town_name": town_name,
            "county": county,
        }

        try:
            lien_search_queued = search_parcel_background(
                request.user,
                town_id,
                loc_id,
                parcel_data,
            )
        except Exception as e:
            # Don't let background search errors crash the page
            logger.warning(f"Failed to queue background search for {town_id}/{loc_id}: {e}")
            lien_search_queued = False
    elif manual_search_triggered:
        # Show info message that search was triggered, but don't show infinite spinner
        # The background search is already running (queued by parcel_refresh_liens)
        lien_search_requested = True
        lien_search_queued = False  # Don't show infinite spinner

    # Now fetch all records to display
    liens = LienRecord.objects.filter(
        created_by=request.user,
        town_id=town_id,
        loc_id=loc_id
    ).order_by('-recording_date', '-created_at')

    legal_actions = LegalAction.objects.filter(
        town_id=town_id,
        loc_id=loc_id
    ).filter(
        Q(created_by=request.user) | Q(source__iexact="CourtListener")
    ).order_by('-filing_date', '-created_at')

    last_lien_search = LienSearchAttempt.objects.filter(
        created_by=request.user,
        town_id=town_id,
        loc_id=loc_id
    ).order_by('-searched_at').first()

    lien_refresh_endpoint = reverse("parcel_refresh_liens", args=[town_id, loc_id])

    market_value_payload = parcel.market_value_payload or {}
    raw_market_comps = market_value_payload.get("comps") or []

    def _comp_price(entry: dict) -> Optional[float]:
        value = entry.get("sale_price")
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    filtered_market_comps = [
        comp for comp in raw_market_comps
        if (_comp_price(comp) or 0) > 100
    ][:5]

    market_value_comps = []
    for comp in filtered_market_comps:
        town_id = comp.get("town_id") or parcel.town.town_id
        loc_id = comp.get("loc_id")
        url = None
        if town_id and loc_id:
            try:
                url = reverse("parcel_detail", args=[int(town_id), loc_id])
            except Exception:  # noqa: BLE001
                url = None

        market_value_comps.append(
            {
                "sale_price": comp.get("sale_price"),
                "sale_date": comp.get("sale_date"),
                "style": comp.get("style"),
                "psf": comp.get("psf"),
                "address": comp.get("address"),
                "living_area": comp.get("living_area"),
                "lot_size": comp.get("lot_size"),
                "url": url,
                "loc_id": loc_id,
            }
        )

    context = {
        "parcel": parcel,
        "sections": sections,
        "parcel_shape": parcel_shape,
        "google_maps": google_maps,
        "zillow_url": zillow_url,
        "property_embed_url": property_embed_url,
        "skiptrace_endpoint": skiptrace_endpoint,
        "skiptrace_allowed": can_run_skiptrace,
        "skiptrace_button_label": button_label,
        "skiptrace_status_message": status_message,
        "skiptrace_initial": initial_payload,
        "skiptrace_config_json": skiptrace_config_json,
        "skiptrace_pricing": pricing,
        "skiptrace_cost_per_lookup_display": (
            pricing.get("perLookupTotalDisplay") if pricing else None
        ),
        "skiptrace_vendor_cost_per_lookup_display": (
            pricing.get("vendorCostPerLookupDisplay") if pricing else None
        ),
        "skiptrace_markup_amount_per_lookup_display": (
            pricing.get("markupAmountPerLookupDisplay") if pricing else None
        ),
        "skiptrace_markup_rate_display": (
            pricing.get("markupRateDisplay") if pricing else None
        ),
        "stripe_payment_required": stripe_enabled,
        "mailer": mailer_context,
        "mailer_download_endpoint": mailer_download_base,
        "attom_data": attom_data, # Add attom_data to the context
        "scraped_documents": attom_data.get_scraped_documents() if attom_data else [],
        "units_detail": units_detail,
        "unit_default_key": default_unit_key or (units_detail[0].get("row_key") if units_detail else ""),
        "unit_attom_endpoint_template": reverse(
            "parcel_unit_attom",
            args=[parcel.town.town_id, "__loc__"],
        ),
        "liens": liens,
        "legal_actions": legal_actions,
        "lien_search_requested": lien_search_requested,
        "lien_search_queued": lien_search_queued,
        "lien_last_search_at": last_lien_search.searched_at if last_lien_search else None,
        "lien_auto_search_enabled": enable_lien_search_param,
        "lien_auto_search_threshold": LIEN_SEARCH_AUTO_THRESHOLD,
        "lien_refresh_endpoint": lien_refresh_endpoint,
        "market_value_comps": market_value_comps,
        "market_value_payload": market_value_payload,
        **nav_context,  # Add navigation context for list browsing
    }

    return render(
        request,
        "leads/parcel_detail.html",
        context,
    )


@login_required
@require_POST
def parcel_search_save_list(request):
    from .attom_service import update_attom_data_for_parcel
    import threading

    form = ParcelListSaveForm(request.POST)
    if not form.is_valid():
        messages.error(request, "List could not be saved. Please try again.")
        return redirect("parcel_search")

    try:
        criteria = json.loads(form.cleaned_data["criteria"])
    except json.JSONDecodeError:
        criteria = {}

    raw_loc_ids = form.cleaned_data["loc_ids"]
    loc_ids = [loc.strip() for loc in raw_loc_ids.split(",") if loc.strip()]

    if not loc_ids:
        messages.warning(request, "No parcels available to save for this search.")
        return redirect("parcel_search")

    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        messages.error(request, "Unable to determine your workspace owner.")
        return redirect("parcel_search")

    saved_list = SavedParcelList.objects.create(
        name=form.cleaned_data["name"],
        town_id=form.cleaned_data["town_id"],
        town_name=criteria.get("town_name", ""),
        criteria=criteria,
        loc_ids=loc_ids,
        created_by=workspace_owner,
    )

    # Run ATTOM enrichment in background thread (non-blocking)
    def enrich_attom_data_background(saved_list_id, town_id, loc_ids):
        """Background task to enrich saved list with ATTOM data."""
        try:
            # Re-fetch the saved list in this thread to avoid threading issues
            from django.db import connection
            connection.close()  # Close any existing connection from parent thread

            saved_list = SavedParcelList.objects.get(pk=saved_list_id)

            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = []
                for loc_id in loc_ids:
                    futures.append(executor.submit(update_attom_data_for_parcel, saved_list, town_id, loc_id))

                for future in as_completed(futures):
                    try:
                        future.result()
                    except Exception as e:
                        logger.error(f"Error updating ATTOM data for parcel: {e}")

            logger.info(f"Successfully enriched {len(loc_ids)} parcels with ATTOM data for list '{saved_list.name}'")
        except Exception as e:
            logger.error(f"Background ATTOM enrichment failed for list {saved_list_id}: {e}")

    # Start background enrichment (non-blocking)
    if saved_list:
        background_thread = threading.Thread(
            target=enrich_attom_data_background,
            args=(saved_list.pk, saved_list.town_id, loc_ids),
            daemon=True  # Daemon thread won't prevent app shutdown
        )
        background_thread.start()
        logger.info(f"Started background ATTOM enrichment for {len(loc_ids)} parcels in list '{saved_list.name}'")

    messages.success(
        request,
        f"Saved {len(loc_ids)} parcels to '{saved_list.name}'. ATTOM data enrichment running in background.",
    )
    return redirect("saved_parcel_list_detail", pk=saved_list.pk)


@login_required
@require_POST
def parcel_save_as_lead(request, town_id, loc_id):
    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        messages.error(request, "Unable to determine your workspace owner.")
        return redirect(request.META.get("HTTP_REFERER", reverse("parcel_search")))

    try:
        lead, created = create_lead_from_parcel(town_id, loc_id, user=workspace_owner)
    except MassGISDataError as exc:
        messages.error(request, str(exc))
        return redirect(request.META.get("HTTP_REFERER", reverse("parcel_search")))
    else:
        if created:
            messages.success(request, "Parcel saved as a new lead.")
        else:
            messages.info(
                request,
                "This parcel already exists as a lead. We've opened the existing record.",
            )
        return redirect("lead_detail", pk=lead.pk)


@login_required
@require_GET
def parcel_get_user_lists(request, town_id):
    """
    Get all active saved lists for the current user (from any town).
    Returns JSON array of {id, name, parcel_count, town_name}
    """
    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        return JsonResponse({"success": False, "error": "Unable to determine your workspace owner."}, status=400)

    lists = SavedParcelList.objects.filter(
        created_by=workspace_owner,
        archived_at__isnull=True
    ).order_by('-created_at')

    data = [{
        "id": lst.pk,
        "name": lst.name,
        "parcel_count": len(lst.loc_ids) if isinstance(lst.loc_ids, list) else 0,
        "town_name": lst.town_name or "Multiple Towns"
    } for lst in lists]

    return JsonResponse({"success": True, "lists": data})


@login_required
@require_POST
def parcel_add_to_list(request, town_id, loc_id):
    """
    Add a single parcel to an existing or new saved list.
    POST params:
    - list_id: ID of existing list (optional)
    - list_name: Name of new list (required if list_id not provided)
    """
    from .attom_service import update_attom_data_for_parcel

    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        return JsonResponse({"success": False, "error": "Unable to determine your workspace owner."}, status=400)

    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        return JsonResponse({"success": False, "error": str(exc)}, status=400)

    list_id = request.POST.get("list_id")
    list_name = request.POST.get("list_name", "").strip()

    if list_id:
        # Add to existing list
        try:
            saved_list = SavedParcelList.objects.get(pk=list_id, created_by=workspace_owner)
        except SavedParcelList.DoesNotExist:
            return JsonResponse({"success": False, "error": "List not found."}, status=404)

        # Get existing parcels - support both old format (strings) and new format (objects)
        loc_ids = saved_list.loc_ids if isinstance(saved_list.loc_ids, list) else []

        # Normalize to new format: list of {"town_id": int, "loc_id": str}
        normalized_parcels = []
        for item in loc_ids:
            if isinstance(item, dict):
                # Already new format
                normalized_parcels.append(item)
            else:
                # Old format (just loc_id string), assume it's from the list's town
                normalized_parcels.append({"town_id": saved_list.town_id, "loc_id": item})

        # Check if parcel is already in the list
        for parcel_ref in normalized_parcels:
            if parcel_ref.get("town_id") == town_id and parcel_ref.get("loc_id") == loc_id:
                return JsonResponse({
                    "success": False,
                    "error": f"This parcel is already in '{saved_list.name}'."
                }, status=400)

        # Add the new parcel
        normalized_parcels.append({"town_id": town_id, "loc_id": loc_id})
        saved_list.loc_ids = normalized_parcels

        # Update town_name to "Multiple Towns" if mixing towns
        if saved_list.town_id != town_id:
            saved_list.town_name = "Multiple Towns"

        saved_list.save()

        message = f"Added parcel to '{saved_list.name}'."
    else:
        # Create new list
        if not list_name:
            return JsonResponse({"success": False, "error": "List name is required."}, status=400)

        saved_list = SavedParcelList.objects.create(
            name=list_name,
            town_id=town_id,
            town_name=parcel.town.name,
            criteria={"source": "single_parcel"},
            loc_ids=[{"town_id": town_id, "loc_id": loc_id}],
            created_by=workspace_owner,
        )

        message = f"Created new list '{list_name}' with this parcel."

    # Queue ATTOM data update in background (non-blocking)
    import threading

    def enrich_single_parcel_background(saved_list_id, town_id, loc_id):
        """Background task to enrich single parcel with ATTOM data."""
        try:
            from django.db import connection
            connection.close()  # Close any existing connection from parent thread

            saved_list = SavedParcelList.objects.get(pk=saved_list_id)
            update_attom_data_for_parcel(saved_list, town_id, loc_id)
            logger.info(f"Successfully enriched parcel {town_id}/{loc_id} with ATTOM data")
        except Exception as e:
            logger.warning(f"Background ATTOM enrichment failed for {town_id}/{loc_id}: {e}")

    background_thread = threading.Thread(
        target=enrich_single_parcel_background,
        args=(saved_list.pk, town_id, loc_id),
        daemon=True
    )
    background_thread.start()

    return JsonResponse({
        "success": True,
        "message": message,
        "list_id": saved_list.pk,
        "list_name": saved_list.name,
    })


# --- Saved lists dashboard.
@login_required
def saved_parcel_lists(request):
    view_mode = request.GET.get("view", "").lower()
    show_archived = view_mode == "archived"

    base_queryset = _saved_list_queryset_for_user(request.user)
    if show_archived:
        lists = base_queryset.filter(archived_at__isnull=False)
    else:
        lists = base_queryset.filter(archived_at__isnull=True)

    active_count = base_queryset.filter(archived_at__isnull=True).count()
    archived_count = base_queryset.filter(archived_at__isnull=False).count()

    return render(
        request,
        "leads/saved_parcel_lists.html",
        {
            "lists": lists,
            "show_archived": show_archived,
            "active_count": active_count,
            "archived_count": archived_count,
        },
    )


@login_required
@require_GET
def preload_town_dataset(request, town_id):
    try:
        preload_massgis_dataset(town_id)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:
        logger.exception(
            "Failed to preload MassGIS dataset for %s", town_id, exc_info=exc
        )
        return JsonResponse({"error": "Unable to preload dataset."}, status=500)
    return JsonResponse({"status": "ok"})


# --- AJAX skip trace endpoint for parcel detail.
@login_required
@require_POST
def parcel_skiptrace(request, town_id, loc_id):
    payload = _parse_skiptrace_payload(request)
    refresh_requested = _is_truthy(payload.get("refresh"))

    workspace_owner = get_workspace_owner(request.user)
    existing_record = _get_skiptrace_record_for_loc_id(
        town_id, loc_id, user=request.user, fresh_only=True
    )
    if existing_record and not refresh_requested:
        serialized = _serialize_skiptrace_record(existing_record)
        if serialized:
            return JsonResponse(serialized)

    if not _skiptrace_allowed_for_parcel(town_id, loc_id, user=request.user):
        return JsonResponse(
            {"error": "Save this parcel to a list or lead before running skip trace."},
            status=403,
        )

    if workspace_owner is None:
        return JsonResponse(
            {"error": "Unable to determine workspace owner."}, status=400
        )

    if not refresh_requested:
        shared_record = _get_shared_skiptrace_record(
            town_id, loc_id, exclude_owner=workspace_owner, fresh_only=True
        )
        if shared_record:
            cloned = _clone_skiptrace_record_for_owner(
                shared_record,
                owner=workspace_owner,
                town_id=town_id,
                loc_id=loc_id,
            )
            serialized = _serialize_skiptrace_record(cloned or shared_record)
            if serialized:
                return JsonResponse(serialized)

    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    street, city, state, zip_code = _extract_skiptrace_address_from_parcel(parcel)

    if not street or not city:
        return JsonResponse(
            {"error": "Owner mailing address is incomplete."}, status=400
        )

    try:
        result = skiptrace_property(street, city, state, zip_code)
    except SkipTraceError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Skip trace lookup failed", exc_info=exc)
        return JsonResponse(
            {"error": "Unable to retrieve skip trace data."}, status=500
        )

    record = _store_skiptrace_result(
        town_id=parcel.town.town_id,
        loc_id=parcel.loc_id,
        result=result,
        user=workspace_owner,
    )

    for lead in _lead_queryset_for_user(request.user).filter(
        loc_id__iexact=parcel.loc_id
    ):
        _update_lead_contact_from_skiptrace(lead, result)

    serialized = _serialize_skiptrace_record(record)
    if not serialized:
        serialized = {
            "ownerName": result.owner_name or "",
            "email": result.email or "",
            "phones": [
                {
                    "number": phone.number,
                    "type": phone.type,
                    "score": phone.score,
                    "dnc": phone.dnc,
                }
                for phone in (result.phones or [])
            ],
            "lastUpdated": timezone.now().isoformat(),
        }

    return JsonResponse(serialized)


@login_required
@require_POST
def parcel_refresh_liens(request, town_id, loc_id):
    next_url = request.POST.get("next") or reverse("parcel_detail", args=[town_id, loc_id])

    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        messages.error(request, str(exc))
        return redirect(next_url)

    town_name = parcel.town.name if parcel.town else ""
    county = None
    if town_name:
        town_lower = town_name.lower()
        if town_lower in {"salem", "beverly", "peabody", "lynn", "gloucester", "marblehead", "danvers"}:
            county = "Essex"
        elif town_lower in {"boston", "cambridge", "somerville", "brookline", "chelsea", "revere", "winthrop"}:
            county = "Suffolk"
        elif town_lower in {"worcester", "shrewsbury", "westborough", "auburn", "millbury"}:
            county = "Worcester"
        elif town_lower in {"springfield", "chicopee", "holyoke", "westfield"}:
            county = "Hampden"
        elif town_lower in {"lowell", "newton", "framingham", "waltham"}:
            county = "Middlesex"

    parcel_data = {
        "owner_name": parcel.owner_name or "",
        "address": parcel.site_address or "",
        "town_name": town_name,
        "county": county,
    }

    try:
        queued = search_parcel_background(
            request.user,
            town_id,
            loc_id,
            parcel_data,
            force=True,
        )
    except Exception as e:
        # Don't let background search errors crash the endpoint
        logger.warning(f"Failed to queue background search for {town_id}/{loc_id}: {e}")
        queued = False

    if queued:
        messages.success(request, "Lien and court searches queued. Refresh the page shortly to see updates.")
    else:
        messages.info(request, "Lien search already running. Please check back in a moment.")

    redirect_url = _append_query_param(next_url, {"manual_search": "1"})
    return redirect(redirect_url)


@login_required
@require_POST
def lead_skiptrace(request, pk):
    lead = get_object_or_404(_lead_queryset_for_user(request.user), pk=pk)

    payload = _parse_skiptrace_payload(request)
    refresh_requested = _is_truthy(payload.get("refresh"))

    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        return JsonResponse(
            {"error": "Unable to determine workspace owner."}, status=400
        )

    lookup_loc_ids = []
    if lead.loc_id:
        lookup_loc_ids.append(lead.loc_id)
    lookup_loc_ids.append(f"LEAD-{lead.pk}")
    existing_record = _get_skiptrace_record_for_loc_ids(
        lookup_loc_ids, user=request.user, fresh_only=True
    )

    if existing_record and not refresh_requested:
        serialized = _serialize_skiptrace_record(existing_record)
        if serialized:
            return JsonResponse(serialized)

    street = lead.owner_street or lead.site_address
    city = lead.owner_city or lead.site_city
    state = lead.owner_state or getattr(lead, "site_state", None) or "MA"
    zip_code = lead.owner_zip or lead.site_zip or ""

    if not street or not city:
        return JsonResponse(
            {"error": "Owner mailing address is incomplete."}, status=400
        )

    try:
        result = skiptrace_property(street, city, state, zip_code)
    except SkipTraceError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:  # noqa: BLE001
        logger.exception("Skip trace lookup failed", exc_info=exc)
        return JsonResponse(
            {"error": "Unable to retrieve skip trace data."}, status=500
        )

    stored_records = []

    if lead.loc_id:
        stored_records.append(
            _store_skiptrace_result(
                town_id=None,
                loc_id=lead.loc_id,
                result=result,
                user=workspace_owner,
            )
        )

        inferred_town_id = _first_saved_list_town_id_for_loc_id(
            lead.loc_id, user=request.user
        )
        if inferred_town_id is not None:
            stored_records.append(
                _store_skiptrace_result(
                    town_id=inferred_town_id,
                    loc_id=lead.loc_id,
                    result=result,
                    user=workspace_owner,
                )
            )

    stored_records.append(
        _store_skiptrace_result(
            town_id=None,
            loc_id=f"LEAD-{lead.pk}",
            result=result,
            user=workspace_owner,
        )
    )

    # Update CRM contact fields for this lead (and any duplicates sharing the same parcel id)
    _update_lead_contact_from_skiptrace(lead, result)
    if lead.loc_id:
        sibling_queryset = _lead_queryset_for_user(request.user).filter(
            loc_id__iexact=lead.loc_id
        ).exclude(pk=lead.pk)
        for sibling in sibling_queryset:
            _update_lead_contact_from_skiptrace(sibling, result)

    record = next((rec for rec in stored_records if rec), None)
    if record is None:
        record = _get_skiptrace_record_for_loc_ids(lookup_loc_ids, user=request.user)

    serialized = _serialize_skiptrace_record(record)
    if not serialized:
        serialized = {
            "ownerName": result.owner_name or "",
            "email": result.email or "",
            "phones": [
                {
                    "number": phone.number,
                    "type": phone.type,
                    "score": phone.score,
                    "dnc": phone.dnc,
                }
                for phone in (result.phones or [])
            ],
            "lastUpdated": timezone.now().isoformat(),
        }

    return JsonResponse(serialized)


# --- Saved list detail: loads parcels, skip-trace status, bulk actions.
@login_required
def saved_parcel_list_detail(request, pk):
    saved_list = get_object_or_404(_saved_list_queryset_for_user(request.user), pk=pk)
    parcels, skiptrace_records, pending_parcels = _pending_parcels_for_saved_list(
        saved_list, user=request.user
    )
    total = len(parcels)

    loc_ids = [parcel.loc_id for parcel in parcels if parcel.loc_id]
    lead_lookup: dict[str, int] = {}
    if loc_ids:
        lead_lookup = {
            _normalize_loc_id(loc_id): lead_pk
            for loc_id, lead_pk in _lead_queryset_for_user(request.user)
            .filter(loc_id__in=loc_ids)
            .values_list("loc_id", "pk")
        }
    parcel_rows = [
        {
            "parcel": parcel,
            "lead_pk": lead_lookup.get(_normalize_loc_id(parcel.loc_id)),
            "skiptraced": bool(skiptrace_records.get(_normalize_loc_id(parcel.loc_id))),
        }
        for parcel in parcels
    ]

    pending_skiptrace_count = len(pending_parcels)
    pricing = _build_skiptrace_pricing(pending_skiptrace_count)

    criteria_qs = None
    if isinstance(saved_list.criteria, dict):
        params = {
            "town_id": saved_list.criteria.get("town_id", saved_list.town_id),
            "property_category": saved_list.criteria.get("property_category", "any"),
            "address_contains": saved_list.criteria.get("address_contains", ""),
            "style": saved_list.criteria.get("style", ""),
            "property_type": saved_list.criteria.get("property_type", "any"),
            "equity_min": saved_list.criteria.get("equity_min", ""),
            "absentee": saved_list.criteria.get("absentee", "any"),
            "min_price": saved_list.criteria.get("min_price") or "",
            "max_price": saved_list.criteria.get("max_price") or "",
            "min_years_owned": saved_list.criteria.get("min_years_owned") or "",
            "max_years_owned": saved_list.criteria.get("max_years_owned") or "",
            "proximity_address": saved_list.criteria.get("proximity_address") or "",
            "proximity_radius_miles": saved_list.criteria.get("proximity_radius_miles")
            or "",
            "limit": saved_list.criteria.get("limit", PARCEL_SEARCH_MAX_RESULTS),
        }
        saved_shape = saved_list.criteria.get("boundary_shape") or {}
        if saved_shape.get("type") == "circle":
            params.update(
                {
                    "boundary_shape_type": "circle",
                    "boundary_circle_lat": saved_shape.get("center_lat") or "",
                    "boundary_circle_lng": saved_shape.get("center_lng") or "",
                    "boundary_circle_radius_miles": saved_shape.get("radius_miles") or "",
                }
            )
        elif saved_shape.get("type") == "polygon":
            params.update(
                {
                    "boundary_shape_type": "polygon",
                    "boundary_polygon_coords": json.dumps(
                        saved_shape.get("coordinates") or []
                    ),
                }
            )
        criteria_qs = urlencode(params)

    stripe_enabled = bool(
        getattr(settings, "STRIPE_SECRET_KEY", "")
        and getattr(settings, "STRIPE_PUBLISHABLE_KEY", "")
    )
    bulk_payment_config = None
    if (
        stripe_enabled
        and pending_skiptrace_count
        and not saved_list.is_archived
    ):
        bulk_payment_config = {
            "enabled": True,
            "mode": "bulk",
            "createIntentEndpoint": reverse(
                "skiptrace_payment_bulk", args=[saved_list.pk]
            ),
            "cardElementId": "bulk-skiptrace-card-element",
            "cardErrorId": "bulk-skiptrace-card-error",
        }

    workspace_owner = get_workspace_owner(request.user)

    script_options = get_mailer_script_options(workspace_owner)
    default_script_option = script_options[0] if script_options else None
    if script_options and parcels:
        inferred_sector = guess_property_sector(parcels[0])
        default_script_option = next(
            (option for option in script_options if option.sector == inferred_sector),
            script_options[0],
        )
    mailer_script_options = [
        {
            "id": option.id,
            "label": option.label,
            "description": option.summary,
            "sector": option.sector,
            "promptText": option.prompt_text or "",
        }
        for option in script_options
    ]
    saved_list_mailer_endpoint = reverse(
        "saved_parcel_list_mailers", args=[saved_list.pk]
    )

    return render(
        request,
        "leads/saved_parcel_list_detail.html",
        {
            "saved_list": saved_list,
            "parcel_rows": parcel_rows,
            "total": total,
            "criteria_qs": criteria_qs,
            "bulk_skiptrace_pending_count": pending_skiptrace_count,
            "bulk_skiptrace_pricing": pricing,
            "bulk_skiptrace_cost_per_lookup_display": (
                pricing.get("perLookupTotalDisplay") if pricing else None
            ),
            "bulk_skiptrace_vendor_cost_per_lookup_display": (
                pricing.get("vendorCostPerLookupDisplay") if pricing else None
            ),
            "bulk_skiptrace_markup_amount_per_lookup_display": (
                pricing.get("markupAmountPerLookupDisplay") if pricing else None
            ),
            "bulk_skiptrace_processing_fee_per_lookup_display": (
                pricing.get("processingFeePerLookupDisplay") if pricing else None
            ),
            "bulk_skiptrace_total_cost_display": (
                pricing.get("totalCostDisplay") if pricing else None
            ),
            "bulk_skiptrace_base_total_cost_display": (
                pricing.get("baseTotalCostDisplay") if pricing else None
            ),
            "bulk_skiptrace_total_markup_display": (
                pricing.get("markupTotalCostDisplay") if pricing else None
            ),
            "bulk_skiptrace_total_processing_display": (
                pricing.get("processingFeeAmountDisplay") if pricing else None
            ),
            "bulk_skiptrace_endpoint": reverse(
                "saved_parcel_list_skiptrace", args=[saved_list.pk]
            ),
            "skiptrace_cost_per_lookup_display": (
                pricing.get("perLookupTotalDisplay") if pricing else None
            ),
            "skiptrace_markup_rate_display": (
                pricing.get("markupRateDisplay") if pricing else None
            ),
            "skiptrace_processing_fee_rate_display": (
                pricing.get("processingFeeRateDisplay") if pricing else None
            ),
            "skiptrace_processing_fee_fixed_display": (
                pricing.get("processingFeeFixedDisplay") if pricing else None
            ),
            "bulk_skiptrace_payment_config_json": (
                json.dumps(bulk_payment_config) if bulk_payment_config else "null"
            ),
            "list_is_archived": saved_list.is_archived,
            "archive_action_endpoint": reverse(
                "saved_parcel_list_restore" if saved_list.is_archived else "saved_parcel_list_archive",
                args=[saved_list.pk],
            ),
            "mailer_script_options": mailer_script_options,
            "mailer_script_selected": (
                default_script_option.id if default_script_option else None
            ),
            "saved_list_mailer_endpoint": saved_list_mailer_endpoint,
        },
    )


@login_required
@require_GET
# --- Export a saved list to CSV for spreadsheets.
def saved_parcel_list_export(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )
    parcels, skiptrace_records, _ = _pending_parcels_for_saved_list(
        saved_list, user=request.user
    )
    filename_basis = f"{saved_list.name}-{saved_list.town_name}".strip("- ")
    return _build_parcel_csv_response(
        parcels,
        skiptrace_records,
        filename=filename_basis or f"saved-list-{saved_list.pk}",
        default_city=saved_list.town_name,
    )


@login_required
@require_POST
# --- Stripe payment intent for single skip trace.
def skiptrace_payment_single(request):
    # Superusers bypass payment
    if request.user.is_superuser:
        pricing = _build_skiptrace_pricing(1)
        return JsonResponse(
            {
                "clientSecret": "superuser_bypass",
                "pricing": pricing,
                "superuser": True,
            }
        )

    if not _stripe_configured():
        return JsonResponse({"error": "Stripe is not configured."}, status=400)

    payload = _parse_skiptrace_payload(request)
    metadata = {
        "mode": "single",
        "refresh": "true" if _is_truthy(payload.get("refresh")) else "false",
        "context": payload.get("context") or "generic",
    }
    if "townId" in payload:
        metadata["town_id"] = str(payload.get("townId"))
    if "locId" in payload:
        metadata["loc_id"] = str(payload.get("locId"))
    if "leadId" in payload:
        metadata["lead_id"] = str(payload.get("leadId"))

    pricing = _build_skiptrace_pricing(1)
    try:
        intent = _create_payment_intent_for_pricing(pricing, metadata=metadata)
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except stripe.error.StripeError as exc:
        logger.exception("Stripe payment intent creation failed", exc_info=exc)
        return JsonResponse({"error": "Unable to initiate payment."}, status=502)

    return JsonResponse(
        {
            "clientSecret": intent["client_secret"],
            "pricing": pricing,
        }
    )


@login_required
@require_POST
# --- Stripe payment intent for bulk skip trace purchases.
def skiptrace_payment_bulk(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )
    parcels, _, pending_parcels = _pending_parcels_for_saved_list(
        saved_list, user=request.user
    )
    pending_count = len(pending_parcels)
    if pending_count <= 0:
        return JsonResponse(
            {"error": "All properties have already been skip traced."}, status=400
        )

    pricing = _build_skiptrace_pricing(pending_count)

    # Superusers bypass payment
    if request.user.is_superuser:
        return JsonResponse(
            {
                "clientSecret": "superuser_bypass",
                "pricing": pricing,
                "superuser": True,
            }
        )

    if not _stripe_configured():
        return JsonResponse({"error": "Stripe is not configured."}, status=400)

    metadata = {
        "mode": "bulk",
        "saved_list_id": str(saved_list.pk),
        "pending_count": str(pending_count),
    }

    try:
        intent = _create_payment_intent_for_pricing(pricing, metadata=metadata)
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except stripe.error.StripeError as exc:
        logger.exception("Stripe payment intent creation failed", exc_info=exc)
        return JsonResponse({"error": "Unable to initiate payment."}, status=502)

    return JsonResponse(
        {
            "clientSecret": intent["client_secret"],
            "pending": pending_count,
            "pricing": pricing,
        }
    )


@login_required
@require_POST
# --- Bulk skip trace executor for a saved list.
def saved_parcel_list_skiptrace(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )

    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        return JsonResponse(
            {"error": "Unable to determine workspace owner."}, status=400
        )

    if saved_list.is_archived:
        return JsonResponse(
            {"error": "This saved list is archived. Restore it before running skip trace."},
            status=403,
        )

    parcels, _, pending = _pending_parcels_for_saved_list(
        saved_list, user=request.user
    )

    if not parcels:
        return JsonResponse(
            {"error": "No parcels available for skip trace."}, status=400
        )

    payload = _parse_skiptrace_payload(request)
    dry_run_requested = _is_truthy(payload.get("dryRun"))

    if not pending:
        return JsonResponse(
            {
                "processed": 0,
                "failed": 0,
                "message": "All properties have already been skip traced.",
                "pending": 0,
            }
        )

    if dry_run_requested:
        pricing = _build_skiptrace_pricing(len(pending))
        return JsonResponse(
            {
                "processed": 0,
                "failed": 0,
                "pending": len(pending),
                "estimatedCost": pricing.get("totalCost"),
                "estimatedCostDisplay": pricing.get("totalCostDisplay"),
                "pricing": pricing,
                "message": "Dry run only.",
            }
        )

    processed = 0
    failures: list[dict[str, object]] = []
    successes: list[tuple[object, object]] = []

    def _run_skiptrace_lookup(target_parcel):
        street, city, state, zip_code = _extract_skiptrace_address_from_parcel(
            target_parcel
        )
        if not street or not city:
            raise ValueError("Owner mailing address is incomplete.")
        return skiptrace_property(street, city, state, zip_code)

    max_workers = min(5, len(pending)) or 1
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_map = {
            executor.submit(_run_skiptrace_lookup, parcel): parcel for parcel in pending
        }
        for future in as_completed(future_map):
            parcel = future_map[future]
            try:
                result = future.result()
            except ValueError as exc:
                failures.append(
                    {
                        "locId": parcel.loc_id,
                        "error": str(exc),
                    }
                )
            except SkipTraceError as exc:
                failures.append(
                    {
                        "locId": parcel.loc_id,
                        "error": str(exc),
                    }
                )
            except Exception as exc:  # noqa: BLE001
                logger.exception(
                    "Bulk skip trace failed for list %s (loc_id=%s)",
                    saved_list.pk,
                    parcel.loc_id,
                    exc_info=exc,
                )
                failures.append(
                    {
                        "locId": parcel.loc_id,
                        "error": "Unable to retrieve skip trace data.",
                    }
                )
            else:
                successes.append((parcel, result))

    for parcel, result in successes:
        _store_skiptrace_result(
            town_id=parcel.town.town_id,
            loc_id=parcel.loc_id,
            result=result,
            user=workspace_owner,
        )

        leads_qs = _lead_queryset_for_user(request.user).filter(
            loc_id__iexact=parcel.loc_id
        )
        for lead in leads_qs:
            _update_lead_contact_from_skiptrace(lead, result)

    processed = len(successes)

    message: Optional[str]
    if processed and not failures:
        message = f"Skip trace completed for {processed} property{'ies' if processed != 1 else ''}."
    elif processed:
        message = (
            f"Skip trace completed for {processed} property{'ies' if processed != 1 else ''} "
            f"with {len(failures)} issue{'s' if len(failures) != 1 else ''}."
        )
    else:
        message = "No properties were skip traced."

    response_payload = {
        "processed": processed,
        "failed": len(failures),
        "pending": len(pending),
        "errors": failures,
        "message": message,
    }

    status_code = 200 if not failures else 207
    return JsonResponse(response_payload, status=status_code)


@login_required
@require_GET
def property_type_choices(request, town_id):
    try:
        choices = get_massgis_property_type_choices(town_id)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "Unable to load property type choices for town %s", town_id, exc_info=exc
        )
        return JsonResponse(
            {"error": "Unable to load property types right now."}, status=500
        )

    payload = {
        "town_id": town_id,
        "choices": [{"code": code, "label": label} for code, label in choices],
    }
    return JsonResponse(payload)


def _is_blank(value):
    return value is None or (isinstance(value, str) and not value.strip())


def _is_displayable(value):
    if value is None:
        return False
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return False
        return text not in {"0", "nan", "None"}
    return str(value) not in {"0", "nan", "None"}


def _collect_categories(queryset, field_name, fallback_label):
    """Group values while normalising whitespace/casing and track record counts."""
    buckets = defaultdict(lambda: {"raw_values": set(), "label": None, "count": 0})

    for raw_value in queryset.values_list(field_name, flat=True):
        cleaned = raw_value.strip() if isinstance(raw_value, str) else raw_value
        if isinstance(cleaned, str) and not cleaned:
            cleaned = None

        bucket_key = cleaned.lower() if isinstance(cleaned, str) else "__empty__"
        bucket = buckets[bucket_key]
        bucket["raw_values"].add(raw_value)
        if bucket["label"] is None and cleaned is not None:
            bucket["label"] = cleaned
        bucket["count"] += 1

    categories = []
    for bucket in buckets.values():
        raw_values = tuple(
            sorted(bucket["raw_values"], key=lambda v: "" if v is None else str(v))
        )
        label = bucket["label"] if bucket["label"] is not None else fallback_label
        digest_source = "|".join(
            "__none__" if rv is None else str(rv) for rv in raw_values
        )
        base_slug = slugify(label) or "uncategorized"
        slug = (
            f"{base_slug}-{hashlib.sha1(digest_source.encode('utf-8')).hexdigest()[:6]}"
        )

        categories.append(
            {
                "slug": slug,
                "label": label,
                "raw_values": raw_values,
                "count": bucket["count"],
                "is_uncategorized": bucket["label"] is None,
            }
        )

    categories.sort(key=lambda item: item["label"].lower())
    return categories


def _values_match(raw, value):
    if _is_blank(raw) and _is_blank(value):
        return True
    if isinstance(raw, str) and isinstance(value, str):
        return raw.strip().lower() == value.strip().lower()
    return raw == value


def _resolve_category(slug, field_name, queryset, fallback_label):
    for category in _collect_categories(queryset, field_name, fallback_label):
        if category["slug"] == slug:
            return category
    raise Http404("Requested category does not exist.")


def _category_for_value(value, field_name, queryset, fallback_label):
    for category in _collect_categories(queryset, field_name, fallback_label):
        if any(_values_match(raw, value) for raw in category["raw_values"]):
            return category
    return None


def _build_filter(field_name, raw_values):
    clause = None

    for raw in raw_values:
        if _is_blank(raw):
            current = Q(**{f"{field_name}__isnull": True}) | Q(
                **{f"{field_name}__exact": ""}
            )
            if isinstance(raw, str) and raw:
                current |= Q(**{f"{field_name}__iexact": raw})
        elif isinstance(raw, str):
            current = Q(**{f"{field_name}__iexact": raw})
            cleaned = raw.strip()
            if cleaned and cleaned != raw:
                current |= Q(**{f"{field_name}__iexact": cleaned})
        else:
            current = Q(**{field_name: raw})

        clause = current if clause is None else clause | current

    return clause if clause is not None else Q(pk__in=[])


@login_required
def city_overview(request):
    queryset = _lead_queryset_for_user(request.user)
    city_categories = _collect_categories(queryset, "site_city", CITY_FALLBACK_LABEL)
    total_leads = sum(category["count"] for category in city_categories)

    return render(
        request,
        "leads/city_list.html",
        {
            "cities": city_categories,
            "total_leads": total_leads,
        },
    )


@login_required
def zoning_list(request, city_slug):
    queryset = _lead_queryset_for_user(request.user)
    city_category = _resolve_category(
        city_slug, "site_city", queryset, CITY_FALLBACK_LABEL
    )
    city_filter = _build_filter("site_city", city_category["raw_values"])
    city_queryset = _lead_queryset_for_user(request.user).filter(city_filter)

    zoning_categories = _collect_categories(
        city_queryset, "zoning", ZONING_FALLBACK_LABEL
    )

    return render(
        request,
        "leads/zoning_list.html",
        {
            "city": city_category,
            "zonings": zoning_categories,
            "total_leads": city_category["count"],
        },
    )


@login_required
def category_lead_list(request, city_slug, zoning_slug):
    queryset = _lead_queryset_for_user(request.user)
    city_category = _resolve_category(
        city_slug, "site_city", queryset, CITY_FALLBACK_LABEL
    )
    city_filter = _build_filter("site_city", city_category["raw_values"])
    city_queryset = _lead_queryset_for_user(request.user).filter(city_filter)

    zoning_category = _resolve_category(
        zoning_slug, "zoning", city_queryset, ZONING_FALLBACK_LABEL
    )
    zoning_filter = _build_filter("zoning", zoning_category["raw_values"])
    lead_queryset = city_queryset.filter(zoning_filter).order_by("-created_at")
    leads = list(lead_queryset)

    return render(
        request,
        "leads/lead_list.html",
        {
            "leads": leads,
            "city": city_category,
            "zoning": zoning_category,
            "lead_count": len(leads),
        },
    )


# --- Full CRM detail view for a single lead.
@login_required
def lead_detail(request, pk):
    lead = get_object_or_404(_lead_queryset_for_user(request.user), pk=pk)

    lead_bedrooms_display = _format_bed_bath(getattr(lead, "bedrooms", None))
    lead_bathrooms_display = _format_bed_bath(getattr(lead, "bathrooms", None))
    (
        lead_equity_percent,
        lead_estimated_balance,
        lead_estimated_equity,
        lead_roi_percent,
        lead_estimated_rate,
        lead_estimated_payment,
    ) = calculate_equity_metrics(
        {
            "TOTAL_VAL": lead.total_value,
            "LS_PRICE": lead.sale_price,
            "LS_DATE": lead.sale_date,
        }
    )

    all_leads = _lead_queryset_for_user(request.user)
    city_category = _category_for_value(
        lead.site_city,
        "site_city",
        all_leads,
        CITY_FALLBACK_LABEL,
    )

    if city_category:
        city_filter = _build_filter("site_city", city_category["raw_values"])
        city_queryset = _lead_queryset_for_user(request.user).filter(city_filter)
        zoning_category = _category_for_value(
            lead.zoning,
            "zoning",
            city_queryset,
            ZONING_FALLBACK_LABEL,
        )
    else:
        zoning_category = _category_for_value(
            lead.zoning,
            "zoning",
            all_leads,
            ZONING_FALLBACK_LABEL,
        )

    table_data = [
        ("Parcel ID", lead.loc_id),
        ("Address", lead.site_address),
        ("City", lead.site_city),
        ("Zip", lead.site_zip),
        ("Zoning", lead.zoning),
        ("Use Code", lead.use_description),
        ("Use Description", lead.use_description),
        ("Units", lead.units),
        ("Bedrooms", lead_bedrooms_display if lead_bedrooms_display != "—" else None),
        (
            "Bathrooms",
            lead_bathrooms_display if lead_bathrooms_display != "—" else None,
        ),
        ("Style", lead.style),
        ("Stories", lead.stories),
        ("Year Built", lead.year_built),
        ("Building Area (sqft)", lead.bld_area),
        ("Lot Size", lead.lot_size),
        ("Lot Units", lead.lot_units),
        ("Est. Mortgage Balance", _format_currency(lead_estimated_balance)),
        ("Est. Equity", _format_currency(lead_estimated_equity)),
        (
            "Est. Equity %",
            f"{lead_equity_percent:.1f}%" if lead_equity_percent is not None else None,
        ),
        (
            "Est. ROI %",
            f"{lead_roi_percent:.1f}%" if lead_roi_percent is not None else None,
        ),
        (
            "Est. Loan Rate",
            f"{lead_estimated_rate:.2f}%" if lead_estimated_rate is not None else None,
        ),
        (
            "Est. Monthly Payment",
            _format_currency(lead_estimated_payment),
        ),
        (
            "Building Value",
            f"${lead.building_value:,}" if lead.building_value else None,
        ),
        ("Land Value", f"${lead.land_value:,}" if lead.land_value else None),
        ("Total Value", f"${lead.total_value:,}" if lead.total_value else None),
        ("Sale Price", f"${lead.sale_price:,}" if lead.sale_price else None),
        ("Sale Date", lead.sale_date),
        ("Registry Book/Page", f"{lead.sale_book or ''}/{lead.sale_page or ''}"),
        ("Mailing Owner", lead.owner_name),
        (
            "Mailing Address",
            f"{lead.owner_street}, {lead.owner_city}, {lead.owner_state} {lead.owner_zip}",
        ),
        (
            "Alternate Owner",
            lead.owner_name_2 if hasattr(lead, "owner_name_2") else None,
        ),
        ("Owner Email", lead.email),
        (
            "Phone #1",
            (
                f"{lead.phone_1} {'(DNC)' if lead.dnc_1 == 'TRUE' else ''}"
                if lead.phone_1
                else None
            ),
        ),
        (
            "Phone #2",
            (
                f"{lead.phone_2} {'(DNC)' if lead.dnc_2 == 'TRUE' else ''}"
                if lead.phone_2
                else None
            ),
        ),
        (
            "Phone #3",
            (
                f"{lead.phone_3} {'(DNC)' if lead.dnc_3 == 'TRUE' else ''}"
                if lead.phone_3
                else None
            ),
        ),
        ("Status", lead.status),
        ("Notes", lead.notes),
    ]

    sections = [
        {
            "title": "Property Overview",
            "items": [
                ("Parcel ID", lead.loc_id),
                ("Zoning", lead.zoning),
                ("Use Description", lead.use_description),
                ("Style", lead.style),
                ("Stories", lead.stories),
                ("Units", lead.units),
                ("Bedrooms", lead_bedrooms_display),
                ("Bathrooms", lead_bathrooms_display),
                (
                    "Bedrooms",
                    getattr(lead, "bedrooms", None),
                ),
                (
                    "Bathrooms",
                    getattr(lead, "bathrooms", None),
                ),
            ],
        },
        {
            "title": "Dimensions & Year",
            "items": [
                ("Year Built", lead.year_built),
                ("Building Area (sqft)", lead.bld_area),
                ("Lot Size", lead.lot_size),
                ("Lot Units", lead.lot_units),
            ],
        },
        {
            "title": "Valuation",
            "items": [
                (
                    "Building Value",
                    f"${lead.building_value:,}" if lead.building_value else None,
                ),
                ("Land Value", f"${lead.land_value:,}" if lead.land_value else None),
                ("Total Value", f"${lead.total_value:,}" if lead.total_value else None),
                ("Est. Mortgage Balance", _format_currency(lead_estimated_balance)),
                ("Est. Equity", _format_currency(lead_estimated_equity)),
                (
                    "Est. Equity %",
                    (
                        f"{lead_equity_percent:.1f}%"
                        if lead_equity_percent is not None
                        else None
                    ),
                ),
                (
                    "Est. ROI %",
                    (
                        f"{lead_roi_percent:.1f}%"
                        if lead_roi_percent is not None
                        else None
                    ),
                ),
                (
                    "Est. Loan Rate",
                    (
                        f"{lead_estimated_rate:.2f}%"
                        if lead_estimated_rate is not None
                        else None
                    ),
                ),
                (
                    "Est. Monthly Payment",
                    _format_currency(lead_estimated_payment),
                ),
            ],
        },
        {
            "title": "Sale History",
            "items": [
                ("Sale Price", f"${lead.sale_price:,}" if lead.sale_price else None),
                ("Sale Date", lead.sale_date),
                (
                    "Registry Book/Page",
                    f"{lead.sale_book or ''}/{lead.sale_page or ''}",
                ),
                (
                    "Est. Loan Rate",
                    (
                        f"{lead_estimated_rate:.2f}%"
                        if lead_estimated_rate is not None
                        else None
                    ),
                ),
                (
                    "Est. Monthly Payment",
                    _format_currency(lead_estimated_payment),
                ),
            ],
        },
        {
            "title": "Owner & Mailing",
            "items": [
                ("Primary Owner", lead.owner_name),
                (
                    "Alternate Owner",
                    lead.owner_name_2 if hasattr(lead, "owner_name_2") else None,
                ),
                (
                    "Mailing Address",
                    f"{lead.owner_street}, {lead.owner_city}, {lead.owner_state} {lead.owner_zip}",
                ),
            ],
        },
        {
            "title": "Skip Trace",
            "items": [],
        },
        {
            "title": "Contacts & CRM",
            "items": [
                ("Owner Email", lead.email),
                (
                    "Phone #1",
                    (
                        f"{lead.phone_1} {'(DNC)' if lead.dnc_1 == 'TRUE' else ''}"
                        if lead.phone_1
                        else None
                    ),
                ),
                (
                    "Phone #2",
                    (
                        f"{lead.phone_2} {'(DNC)' if lead.dnc_2 == 'TRUE' else ''}"
                        if lead.phone_2
                        else None
                    ),
                ),
                (
                    "Phone #3",
                    (
                        f"{lead.phone_3} {'(DNC)' if lead.dnc_3 == 'TRUE' else ''}"
                        if lead.phone_3
                        else None
                    ),
                ),
                ("Status", lead.status),
                ("Notes", lead.notes),
            ],
        },
    ]

    parcel_shape = fetch_parcel_shape_for_lead(lead)
    centroid_lon_lat: Optional[Tuple[float, float]] = None
    if parcel_shape and parcel_shape.centroid:
        centroid = parcel_shape.centroid
        converted = massgis_stateplane_to_wgs84(centroid[0], centroid[1])
        if converted:
            centroid_lon_lat = converted

    for section in sections:
        section["items"] = [
            item for item in section["items"] if _is_displayable(item[1])
        ]
    sections = [
        section
        for section in sections
        if section["items"] or section["title"] == "Skip Trace"
    ]

    full_address = _compose_full_address(lead)
    google_maps = None
    lat = None
    lng = None
    if centroid_lon_lat:
        lng, lat = centroid_lon_lat
    elif full_address:
        coords = geocode_address(full_address)
        if coords:
            lng, lat = coords

    property_embed_url = _build_mapillary_embed_url(lat, lng)
    if not property_embed_url:
        property_embed_url = _build_google_street_view_embed(lat, lng)

    lookup_loc_ids = []
    if lead.loc_id:
        lookup_loc_ids.append(lead.loc_id)
    lookup_loc_ids.append(f"LEAD-{lead.pk}")
    record = _get_skiptrace_record_for_loc_ids(lookup_loc_ids, user=request.user)
    record_is_fresh = _skiptrace_record_is_fresh(record) if record else False
    has_address = bool(
        (lead.owner_street or lead.site_address) and (lead.owner_city or lead.site_city)
    )
    skiptrace_endpoint = (
        reverse("lead_skiptrace", args=[lead.pk]) if has_address else None
    )

    if lat is not None and lng is not None:
        google_maps = {
            "address": full_address,
            "lat": lat,
            "lng": lng,
        }
    elif full_address:
        google_maps = {
            "address": full_address,
            "lat": None,
            "lng": None,
        }

    if record:
        status_message = f"Last run {timezone.localtime(record.updated_at).strftime('%b %d, %Y %I:%M %p')}"
        if not record_is_fresh:
            ttl = _skiptrace_cache_ttl_days_value()
            if ttl:
                status_message += f" · Cached result older than {ttl} days."
            else:
                status_message += " · Cached result may be out of date."
        if not has_address:
            status_message += " · Add a mailing address to refresh skip tracing."
        initial_payload = _serialize_skiptrace_record(record)
        button_label = "Refresh Skip Trace" if has_address else "Refresh Skip Trace"
    elif not has_address:
        status_message = "Add a mailing address to enable skip tracing."
        initial_payload = None
        button_label = "Run Skip Trace"
    else:
        status_message = 'Click "Run Skip Trace" to retrieve phone numbers and emails.'
        initial_payload = None
        button_label = "Run Skip Trace"

    skiptrace_items = [("Status", status_message)]

    for section in sections:
        if section["title"] == "Skip Trace":
            section["items"] = skiptrace_items
            break

    pricing = _build_skiptrace_pricing(1)
    stripe_enabled = _stripe_configured()
    payment_config = None
    if stripe_enabled:
        payment_payload = {
            "context": "lead",
            "leadId": lead.pk,
        }
        payment_config = {
            "enabled": True,
            "mode": "single",
            "createIntentEndpoint": reverse("skiptrace_payment_single"),
            "cardElementId": "skiptrace-card-element",
            "cardErrorId": "skiptrace-card-error",
            "intentPayload": payment_payload,
        }

    skiptrace_config = {
        "allowed": has_address,
        "endpoint": skiptrace_endpoint,
        "buttonLabel": button_label,
        "initial": initial_payload,
        "statusMessage": status_message,
        "pricing": pricing,
        "payment": payment_config,
    }
    skiptrace_config_json = json.dumps(skiptrace_config, ensure_ascii=False)

    return render(
        request,
        "leads/lead_detail.html",
        {
            "lead": lead,
            "table_data": table_data,
            "sections": sections,
            "city": city_category,
            "zoning": zoning_category,
            "zillow_url": _build_zillow_url(full_address),
            "parcel_shape": parcel_shape,
            "google_maps": google_maps,
            "property_embed_url": property_embed_url,
            "skiptrace_endpoint": skiptrace_endpoint,
            "skiptrace_allowed": has_address,
            "skiptrace_button_label": button_label,
            "skiptrace_status_message": status_message,
            "skiptrace_initial": initial_payload,
            "skiptrace_config_json": skiptrace_config_json,
            "skiptrace_pricing": pricing,
            "skiptrace_cost_per_lookup_display": (
                pricing.get("perLookupTotalDisplay") if pricing else None
            ),
            "skiptrace_vendor_cost_per_lookup_display": (
                pricing.get("vendorCostPerLookupDisplay") if pricing else None
            ),
            "skiptrace_markup_amount_per_lookup_display": (
                pricing.get("markupAmountPerLookupDisplay") if pricing else None
            ),
            "skiptrace_markup_rate_display": (
                pricing.get("markupRateDisplay") if pricing else None
            ),
            "stripe_payment_required": stripe_enabled,
        },
    )


# --- Simple form to create a lead manually.
@login_required
def lead_create(request):
    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        messages.error(request, "Unable to determine your workspace owner.")
        return redirect("parcel_search")

    if request.method == "POST":
        form = LeadForm(request.POST)
        if form.is_valid():
            lead = form.save(commit=False)
            if not lead.created_by:
                lead.created_by = workspace_owner
            lead.save()
            return redirect("parcel_search")
    else:
        form = LeadForm()
    return render(request, "leads/lead_form.html", {"form": form})


def _normalize_cell(value):
    if value is None:
        return None
    if isinstance(value, float) and pd.isna(value):
        return None
    if isinstance(value, str):
        cleaned = value.strip()
        if cleaned.lower() in {"", "nan", "none", "null"}:
            return None
        return cleaned
    return value


# --- Excel import workflow for bulk lead ingestion.
@login_required
def lead_upload(request):
    workspace_owner = get_workspace_owner(request.user)
    if workspace_owner is None:
        messages.error(request, "Unable to determine your workspace owner.")
        return redirect("parcel_search")

    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES["file"]
            try:
                df = pd.read_excel(file, dtype=str)

                rename_map = {
                    "LOC_ID": "loc_id",
                    "ASSESS_SITE_ADDR": "site_address",
                    "ASSESS_CITY": "site_city",
                    "ASSESS_ZIP": "site_zip",
                    "ASSESS_ZONING": "zoning",
                    "ASSESS_STYLE": "style",
                    "ASSESS_STORIES": "stories",
                    "ASSESS_YEAR_BUILT": "year_built",
                    "ASSESS_LOT_SIZE": "lot_size",
                    "ASSESS_LOT_UNITS": "lot_units",
                    "ASSESS_BLD_AREA": "bld_area",
                    "ASSESS_UNITS": "units",
                    "ASSESS_BLDG_VAL": "building_value",
                    "ASSESS_LAND_VAL": "land_value",
                    "ASSESS_TOTAL_VAL": "total_value",
                    "ASSESS_LS_DATE": "sale_date",
                    "ASSESS_LS_PRICE": "sale_price",
                    "ASSESS_LS_BOOK": "sale_book",
                    "ASSESS_LS_PAGE": "sale_page",
                    "ASSESS_OWNER1": "owner_name",
                    "ASSESS_OWN_ADDR": "owner_street",
                    "ASSESS_OWN_CITY": "owner_city",
                    "ASSESS_OWN_STATE": "owner_state",
                    "ASSESS_OWN_ZIP": "owner_zip",
                    "Phone #1": "phone_1",
                    "Phone #2": "phone_2",
                    "Phone #3": "phone_3",
                    "DNC #1": "dnc_1",
                    "DNC #2": "dnc_2",
                    "DNC #3": "dnc_3",
                    "Email": "email",
                    "M007UC_LUT_CY24_FY24_USE_DESC": "use_description",
                    "M258UC_LUT_CY24_FY24_USE_DESC": "use_description",
                }

                df = df.rename(
                    columns={
                        source: target
                        for source, target in rename_map.items()
                        if source in df.columns
                    }
                )

                valid_fields = [
                    field.name
                    for field in Lead._meta.fields
                    if field.name not in {"id", "created_at", "created_by"}
                ]

                for _, row in df.iterrows():
                    lead_data = {}

                    for field in valid_fields:
                        if field in row:
                            cleaned_value = _normalize_cell(row[field])
                            lead_data[field] = cleaned_value

                    lead_data["status"] = lead_data.get("status") or "Cold"
                    lead_data["notes"] = lead_data.get("notes") or ""

                    meaningful_data = [
                        value
                        for key, value in lead_data.items()
                        if key not in {"status", "notes"} and not _is_blank(value)
                    ]

                    if not meaningful_data:
                        continue

                    lead_data["created_by"] = workspace_owner

                    Lead.objects.create(**lead_data)

                messages.success(request, "✅ Leads imported successfully.")
                return redirect("parcel_search")

            except Exception as exc:
                messages.error(request, f"Upload failed: {exc}")
    else:
        form = UploadFileForm()

    return render(request, "leads/lead_upload.html", {"form": form})


@require_http_methods(["GET", "POST"])
# --- Public scheduling form hit from QR codes.
def schedule_call_request(request, town_id, loc_id):
    import logging
    logger = logging.getLogger(__name__)

    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        raise Http404(str(exc))

    full_address = _compose_full_address(parcel)
    property_city = _normalize_capitalization(
        getattr(parcel, "site_city", None) or getattr(parcel, "owner_city", None) or ""
    )
    raw_recipient = getattr(parcel, "owner_name", None)
    recipient_name = raw_recipient or _extract_owner_first_name(raw_recipient) or "Neighbor"
    recipient_name = _normalize_capitalization(recipient_name) or recipient_name

    # Send QR scan notification on first GET (when owner scans QR code)
    script_param = request.GET.get("script") or request.POST.get("script")
    if request.method == "GET" and not request.GET.get("submitted"):
        # Get user from QR code parameter
        user_id_param = request.GET.get("user_id")
        owner_user = None
        if user_id_param:
            from django.contrib.auth import get_user_model
            User = get_user_model()
            try:
                owner_user = User.objects.get(id=int(user_id_param))
            except (ValueError, User.DoesNotExist):
                logger.warning(f"Invalid user_id {user_id_param} in QR code for scan notification")

        if owner_user:
            from accounts.emails import send_qr_scan_notification
            from django.utils import timezone
            try:
                crm_url = request.build_absolute_uri(reverse("crm_overview"))
                send_qr_scan_notification(owner_user, parcel, timezone.now(), crm_url)
            except Exception as e:
                # Don't fail the page load if notification fails
                logger.error(f"Failed to send QR scan notification: {e}")

    if request.method == "POST":
        form = ScheduleCallRequestForm(request.POST)
        if form.is_valid():
            call_request = form.save(commit=False)
            call_request.town_id = town_id
            call_request.loc_id = loc_id
            call_request.property_address = full_address or getattr(parcel, "site_address", "") or ""
            call_request.property_city = property_city
            call_request.recipient_name = recipient_name
            if script_param:
                note_prefix = f"Source script: {script_param}"
                if call_request.notes:
                    call_request.notes = f"{note_prefix}\n\n{call_request.notes.strip()}"
                else:
                    call_request.notes = note_prefix

            # Extract user_id from QR code URL and assign lead to that specific user
            user_id_param = request.GET.get("user_id") or request.POST.get("user_id")
            if user_id_param:
                from django.contrib.auth import get_user_model
                User = get_user_model()
                try:
                    owner_user = User.objects.get(id=int(user_id_param))
                    call_request.created_by = owner_user
                    logger.info(f"Assigning ScheduleCallRequest for {loc_id} to user {owner_user.username} (ID: {owner_user.id}) from QR code")
                except (ValueError, User.DoesNotExist) as e:
                    logger.warning(f"Invalid user_id {user_id_param} in QR code, lead will be unassigned: {e}")
                    call_request.created_by = None
            else:
                # Fallback: no user_id in QR code (shouldn't happen with new QR codes)
                logger.warning(f"No user_id in QR code for {loc_id}, lead will be unassigned")
                call_request.created_by = None

            call_request.save()
            logger.info(f"Saved ScheduleCallRequest (ID: {call_request.pk}) for {loc_id} with created_by={getattr(call_request.created_by, 'username', None)}")

            # Send notification to the user who sent the mailer (from QR code)
            if call_request.created_by:
                from accounts.emails import send_call_request_notification
                try:
                    # Create a mock lead object with the call request data
                    class CallRequestLead:
                        def __init__(self, call_req, parcel_obj):
                            self.owner_name = call_req.owner_name
                            self.phone = call_req.phone
                            self.email = call_req.email_address
                            self.interest_level = call_req.interest_level
                            self.preferred_contact_time = call_req.preferred_contact_time
                            self.notes = call_req.notes
                            self.created_at = call_req.created_at
                            self.parcel = parcel_obj

                        def get_interest_level_display(self):
                            levels = {
                                "very_interested": "Very Interested",
                                "somewhat_interested": "Somewhat Interested",
                                "just_browsing": "Just Browsing",
                            }
                            return levels.get(self.interest_level, "Interested")

                    mock_lead = CallRequestLead(call_request, parcel)
                    lead_url = request.build_absolute_uri(
                        reverse("crm_overview")
                    )
                    send_call_request_notification(call_request.created_by, mock_lead, lead_url)
                except Exception as e:
                    # Don't fail the request if notification fails
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.error(f"Failed to send call request notification: {e}")

            messages.success(request, "Thanks! We'll be in touch shortly.")
            redirect_url = reverse("schedule_call_request", args=[town_id, loc_id])
            if script_param:
                redirect_url = f"{redirect_url}?submitted=1&script={script_param}"
            else:
                redirect_url = f"{redirect_url}?submitted=1"
            return redirect(redirect_url)
    else:
        form = ScheduleCallRequestForm()

    submitted = request.GET.get("submitted") == "1"

    context = {
        "form": form,
        "parcel": parcel,
        "full_address": full_address,
        "recipient_name": recipient_name,
        "property_city": property_city,
        "script_param": script_param,
        "submitted": submitted,
    }
    return render(request, "leads/schedule_call_form.html", context)


# --- Download single parcel mailer as PDF (with fallback HTML).
@login_required
def mailer_download_pdf(request, town_id, loc_id):
    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        raise Http404(str(exc))

    full_address = _compose_full_address(parcel)
    zillow_url = _build_zillow_url(full_address)

    # Get skip trace record if available
    skiptrace_record = _get_skiptrace_record_for_loc_id(
        town_id, loc_id, user=request.user
    )

    # Try to build a hero image
    hero_image_url = None
    try:
        shape = get_massgis_parcel_shape(parcel.town, parcel.loc_id)
        if shape and shape.centroid:
            lon, lat = massgis_stateplane_to_wgs84(
                shape.centroid[0], shape.centroid[1]
            ) or (None, None)
            if lat is not None and lon is not None:
                hero_image_url = _build_google_street_view_image_url(lat, lon)
    except Exception:
        hero_image_url = None

    try:
        bundle = _generate_mailer_bundle(
            parcel,
            full_address=full_address,
            zillow_url=zillow_url,
            hero_image_url=hero_image_url,
            request=request,
            town_id_override=town_id,
            skiptrace_record=skiptrace_record,
            user=request.user,
        )
        scripts = bundle["scripts"]
        requested_script = request.GET.get("script") or bundle["default_id"]
        if requested_script not in scripts:
            requested_script = bundle["default_id"]
        mailer_ctx = scripts[requested_script]
        pdf_bytes = _render_mailer_pdf([mailer_ctx])
        filename = f"mailer-{loc_id}-{requested_script}.pdf"
        response = HttpResponse(pdf_bytes, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        response["Content-Length"] = str(len(pdf_bytes))
        return response
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to build PDF mailer for %s/%s", town_id, loc_id, exc_info=exc)
        # Fallback to a simplified PDF so the user still gets the content.
        mailer_ctx = _build_mailer_context(
            parcel,
            full_address=full_address,
            zillow_url=zillow_url,
            hero_image_url=hero_image_url,
            request=request,
            town_id_override=town_id,
            skiptrace_record=skiptrace_record,
            user=request.user,
        )
        try:
            pdf_bytes = _render_mailer_pdf([mailer_ctx])
            fallback_name = f"mailer-{loc_id}-fallback.pdf"
            response = HttpResponse(pdf_bytes, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="{fallback_name}"'
            response["Content-Length"] = str(len(pdf_bytes))
            return response
        except Exception as pdf_exc:  # noqa: BLE001
            logger.exception(
                "Failed to generate fallback PDF mailer for %s/%s",
                town_id,
                loc_id,
                exc_info=pdf_exc,
            )
            return HttpResponse(
                "Failed to generate mailer PDF.",
                status=500,
                content_type="text/plain; charset=utf-8",
            )


@login_required
@require_GET
# --- AJAX endpoint returning pre-rendered mailer JSON for the UI.
def parcel_generate_mailer(request, town_id, loc_id):
    """
    JSON API that returns the mailer context for a parcel.
    """
    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    full_address = _compose_full_address(parcel)
    zillow_url = _build_zillow_url(full_address)

    hero_image_url = None
    try:
        shape = get_massgis_parcel_shape(parcel.town, parcel.loc_id)
        if shape and shape.centroid:
            lon, lat = massgis_stateplane_to_wgs84(
                shape.centroid[0], shape.centroid[1]
            ) or (None, None)
            if lat is not None and lon is not None:
                hero_image_url = _build_google_street_view_image_url(lat, lon)
    except Exception:
        pass

    # Get skip trace record if available
    skiptrace_record = _get_skiptrace_record_for_loc_id(
        town_id, loc_id, user=request.user
    )

    download_base = reverse("mailer_download_pdf", args=[parcel.town.town_id, parcel.loc_id])
    mailer_ctx = _build_mailer_context(
        parcel,
        full_address=full_address,
        zillow_url=zillow_url,
        hero_image_url=hero_image_url,
        download_endpoint=download_base,
        request=request,
        skiptrace_record=skiptrace_record,
        user=request.user,
    )
    payload = {
        "locId": parcel.loc_id,
        "address": full_address,
        "zillowUrl": zillow_url,
        "heroImageUrl": hero_image_url,
        "mailer": mailer_ctx,
        "downloadBase": download_base,
    }
    return JsonResponse(payload)


@login_required
@require_GET
def parcel_unit_attom(request, town_id, loc_id):
    normalized_loc = _normalize_loc_id(loc_id)
    if not normalized_loc:
        return JsonResponse({"error": "Invalid unit identifier."}, status=400)

    max_age_days = getattr(settings, "ATTOM_CACHE_MAX_AGE_DAYS", 60)
    cache_cutoff = timezone.now() - timedelta(days=max_age_days)

    try:
        parcel = get_massgis_parcel_detail(town_id, normalized_loc)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)

    unit_key = request.GET.get("unit_key")
    units = parcel.units_detail or []
    target_unit = None
    if unit_key:
        lowered_key = unit_key.strip().lower()
        for candidate in units:
            for key_name in ("row_key", "source_id", "id", "unit_number"):
                value = candidate.get(key_name)
                if value and str(value).strip().lower() == lowered_key:
                    target_unit = candidate
                    break
            if target_unit:
                break
    if target_unit is None:
        target_unit = next(
            (
                candidate
                for candidate in units
                if candidate.get("loc_id") and _normalize_loc_id(candidate["loc_id"]) == normalized_loc
            ),
            None,
        )
    if target_unit is None:
        owner_hint = request.GET.get("owner_name")
        if owner_hint:
            lowered_owner = owner_hint.strip().lower()
            target_unit = next(
                (
                    candidate
                    for candidate in units
                    if candidate.get("owner") and candidate["owner"].strip().lower() == lowered_owner
                ),
                None,
            )

    cache_key = build_attom_cache_key(normalized_loc, target_unit)

    attom_data = (
        AttomData.objects.filter(
            town_id=town_id,
            loc_id=cache_key,
            last_updated__gte=cache_cutoff,
        )
        .order_by("-last_updated")
        .first()
    )

    if not attom_data:
        attom_payload = get_or_fetch_attom_data(
            town_id,
            normalized_loc,
            max_age_days,
            parcel=parcel,
            unit=target_unit,
            unit_key=unit_key,
        )
        raw_responses = attom_payload.get("raw_response", {}) if isinstance(attom_payload, dict) else {}
        has_data = bool(raw_responses.get("expandedprofile", {}).get("property"))
        if not has_data:
            return JsonResponse({"hasData": False}, status=200)

        attom_data = (
            AttomData.objects.filter(
                town_id=town_id,
                loc_id=cache_key,
            )
            .order_by("-last_updated")
            .first()
        )

    if not attom_data:
        return JsonResponse({"hasData": False}, status=200)

    current_value_param = request.GET.get("current_value")
    current_value = None
    if current_value_param not in {None, "", "null"}:
        try:
            current_value = float(current_value_param)
        except (TypeError, ValueError):
            current_value = None

    valuation_items: list[dict[str, object]] = []
    sale_items: list[dict[str, object]] = []

    mortgage_balance = equity_value = equity_percent = roi_percent = monthly_payment = None
    if attom_data.mortgage_loan_amount:
        (
            mortgage_balance,
            equity_value,
            equity_percent,
            roi_percent,
            monthly_payment,
        ) = _calculate_mortgage_balance_from_attom(attom_data, current_value)

        if mortgage_balance is not None:
            valuation_items.append(
                {
                    "label": "Mortgage Balance (ATTOM)",
                    "value": _format_currency(mortgage_balance),
                    "divider": True,
                }
            )
        if equity_value is not None:
            valuation_items.append(
                {
                    "label": "Equity (ATTOM)",
                    "value": _format_currency(equity_value),
                }
            )
        if equity_percent is not None:
            valuation_items.append(
                {
                    "label": "Equity % (ATTOM)",
                    "value": f"{equity_percent:.1f}%",
                }
            )
        if roi_percent is not None:
            valuation_items.append(
                {
                    "label": "ROI % (ATTOM)",
                    "value": f"{roi_percent:.1f}%",
                }
            )
        if monthly_payment is not None:
            valuation_items.append(
                {
                    "label": "Monthly Payment (ATTOM)",
                    "value": _format_currency(monthly_payment),
                }
            )

        sale_items.extend(
            [
                {
                    "label": "Mortgage Amount",
                    "value": _format_currency(float(attom_data.mortgage_loan_amount)),
                },
                {
                    "label": "Mortgage Date",
                    "value": attom_data.mortgage_recording_date,
                },
                {"label": "Lender", "value": attom_data.mortgage_lender_name},
                {"label": "Loan Type", "value": attom_data.mortgage_loan_type},
            ]
        )
        if attom_data.mortgage_interest_rate:
            sale_items.append(
                {
                    "label": "Interest Rate",
                    "value": f"{attom_data.mortgage_interest_rate}%",
                }
            )
        if attom_data.mortgage_term_years:
            sale_items.append(
                {
                    "label": "Loan Term",
                    "value": f"{attom_data.mortgage_term_years} years",
                }
            )
        if monthly_payment is not None:
            sale_items.append(
                {
                    "label": "Monthly Payment",
                    "value": _format_currency(monthly_payment),
                }
            )

    if attom_data.tax_assessment_year or attom_data.tax_assessed_value or attom_data.tax_amount_annual:
        if sale_items:
            sale_items.append(
                {
                    "label": "────────────────────────",
                    "value": "",
                    "separator": True,
                }
            )
        if attom_data.tax_assessment_year:
            sale_items.append(
                {
                    "label": "Tax Year",
                    "value": str(attom_data.tax_assessment_year),
                }
            )
        if attom_data.tax_assessed_value:
            sale_items.append(
                {
                    "label": "Tax Assessed Value",
                    "value": _format_currency(float(attom_data.tax_assessed_value)),
                }
            )
        if attom_data.tax_amount_annual:
            sale_items.append(
                {
                    "label": "Annual Tax",
                    "value": _format_currency(float(attom_data.tax_amount_annual)),
                }
            )

    valuation_items = [item for item in valuation_items if item.get("value")]
    sale_items = [item for item in sale_items if item.get("value") or item.get("separator")]

    has_data = bool(valuation_items or sale_items)

    return JsonResponse(
        {
            "hasData": has_data,
            "valuation": valuation_items,
            "saleHistory": sale_items,
        }
    )


@login_required
@require_POST
# --- Archive a saved list (soft delete).
def saved_parcel_list_archive(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )
    redirect_target = request.POST.get("next") or reverse(
        "saved_parcel_list_detail", args=[saved_list.pk]
    )

    if saved_list.archived_at:
        messages.info(
            request,
            f"'{saved_list.name}' is already archived.",
        )
        return redirect(redirect_target)

    saved_list.archived_at = timezone.now()
    saved_list.save(update_fields=["archived_at"])
    messages.success(
        request,
        f"Archived '{saved_list.name}'. Restore it anytime from the archived tab.",
    )
    return redirect(redirect_target)


@login_required
@require_POST
# --- Restore an archived list.
def saved_parcel_list_restore(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )
    redirect_target = request.POST.get("next") or reverse(
        "saved_parcel_list_detail", args=[saved_list.pk]
    )

    if not saved_list.archived_at:
        messages.info(
            request,
            f"'{saved_list.name}' is already active.",
        )
        return redirect(redirect_target)

    saved_list.archived_at = None
    saved_list.save(update_fields=["archived_at"])
    messages.success(
        request,
        f"Restored '{saved_list.name}' to the active lists.",
    )
    return redirect(redirect_target)


@login_required
@require_POST
# --- Bulk mailer generator for saved lists (PDF bundle).
def saved_parcel_list_mailers(request, pk):
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )
    script_id = request.POST.get("script_id") or request.POST.get("prompt_id")
    if not script_id:
        return JsonResponse({"error": "Select a template to generate mailers."}, status=400)

    parcels, skiptrace_records, _ = _pending_parcels_for_saved_list(
        saved_list, user=request.user
    )
    if not parcels:
        return JsonResponse(
            {"error": "No parcels available to generate mailers."}, status=400
        )

    scripts_to_render: list[dict] = []
    generated = 0
    skipped = 0

    logger.info(f"Generating mailers for {len(parcels)} parcels in list '{saved_list.name}' (ID: {saved_list.pk})")

    for parcel in parcels:
        try:
            full_address = _compose_full_address(parcel)
            zillow_url = _build_zillow_url(full_address)
            normalized_loc_id = _normalize_loc_id(parcel.loc_id)
            skiptrace_record = skiptrace_records.get(normalized_loc_id) if normalized_loc_id else None
            _, script, _ = _render_mailer_script_for_parcel(
                parcel,
                script_id,
                full_address=full_address,
                zillow_url=zillow_url,
                request=request,
                town_id_override=saved_list.town_id,
                skiptrace_record=skiptrace_record,
                user=request.user,
            )
            scripts_to_render.append(script)
            generated += 1
            logger.debug(f"✓ Generated mailer for parcel {parcel.loc_id}")
        except Exception as exc:  # noqa: BLE001
            skipped += 1
            logger.exception(
                "Bulk mailer generation failed for list %s (loc_id=%s): %s",
                saved_list.pk,
                getattr(parcel, "loc_id", None),
                str(exc),
                exc_info=exc,
            )

    logger.info(f"Mailer generation complete: {generated} generated, {skipped} skipped")

    if not scripts_to_render:
        return JsonResponse(
            {"error": "Unable to generate mailers for this list."}, status=500
        )

    filename_base = (
        slugify(saved_list.name or f"saved-list-{saved_list.pk}") or f"saved-list-{saved_list.pk}"
    )

    # Try Word document generation first
    try:
        docx_bytes = _render_mailer_docx(scripts_to_render)
        filename = f"{filename_base}-{script_id}.docx"
        response = HttpResponse(
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        response["Content-Length"] = str(len(docx_bytes))
        response["X-Mailers-Generated"] = str(generated)
        response["X-Mailers-Skipped"] = str(skipped)
        response["X-Mailers-Reused"] = "0"
        return response
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "Failed to generate bulk mailer Word doc for list %s", saved_list.pk, exc_info=exc
        )

    try:
        pdf_bytes = _render_mailer_pdf(scripts_to_render)
        filename = f"{filename_base}-{script_id}.pdf"
        response = HttpResponse(pdf_bytes, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        response["Content-Length"] = str(len(pdf_bytes))
        response["X-Mailers-Generated"] = str(generated)
        response["X-Mailers-Skipped"] = str(skipped)
        response["X-Mailers-Reused"] = "0"
        return response
    except Exception as exc:  # noqa: BLE001
        logger.exception(
            "Failed to generate bulk mailer PDF fallback for list %s", saved_list.pk, exc_info=exc
        )
        return JsonResponse(
            {"error": "Failed to generate mailers for this list."}, status=500
        )


@login_required
@require_http_methods(["POST"])
def saved_parcel_list_labels(request, pk):
    """
    Generate a Word document of mailing labels for all parcels in a saved list.
    Supports multiple label formats (Avery 5160, 5163, etc.)
    """
    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )

    # Get label format from request (default to Avery 5160)
    label_format = request.POST.get("label_format", "5160")

    # Define label formats with dimensions in inches
    LABEL_FORMATS = {
        "5160": {"name": "Avery 5160", "cols": 3, "rows": 10, "width": 2.625, "height": 1.0, "margin_top": 0.5, "margin_left": 0.1875, "gutter_h": 0.125, "gutter_v": 0},
        "5163": {"name": "Avery 5163", "cols": 2, "rows": 5, "width": 4.0, "height": 2.0, "margin_top": 0.5, "margin_left": 0.15625, "gutter_h": 0.1875, "gutter_v": 0},
        "5167": {"name": "Avery 5167", "cols": 4, "rows": 20, "width": 1.75, "height": 0.5, "margin_top": 0.5, "margin_left": 0.3125, "gutter_h": 0.125, "gutter_v": 0},
    }

    format_spec = LABEL_FORMATS.get(label_format, LABEL_FORMATS["5160"])

    # Get all parcels in the list
    parcel_refs = list(_iter_saved_list_parcel_refs(saved_list))
    if not parcel_refs:
        return JsonResponse({"error": "No parcels in this list."}, status=400)

    # Load parcel data
    parcels_by_key = {}
    grouped_loc_ids = defaultdict(list)
    for ref in parcel_refs:
        grouped_loc_ids[ref.town_id].append(ref.loc_id)

    for town_id, loc_list in grouped_loc_ids.items():
        for parcel in load_massgis_parcels_by_ids(town_id, loc_list, saved_list=saved_list):
            normalized = _normalize_loc_id(parcel.loc_id)
            if normalized:
                parcels_by_key[(town_id, normalized)] = parcel

    parcels = []
    for ref in parcel_refs:
        parcel = parcels_by_key.get((ref.town_id, ref.normalized_loc_id))
        if parcel:
            parcels.append(parcel)

    if not parcels:
        return JsonResponse({"error": "No parcels available for labels."}, status=400)

    # Generate labels as Word document
    try:
        docx_bytes = _generate_label_sheet_docx(parcels, format_spec)
        filename_base = slugify(saved_list.name or f"saved-list-{saved_list.pk}") or f"saved-list-{saved_list.pk}"
        filename = f"{filename_base}-labels-{label_format}.docx"

        response = HttpResponse(
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        response["Content-Length"] = str(len(docx_bytes))
        return response

    except Exception as exc:
        logger.exception("Failed to generate label sheet for list %s", saved_list.pk, exc_info=exc)
        return JsonResponse({"error": "Failed to generate label sheet."}, status=500)


@login_required
@require_http_methods(["POST"])
def bulk_legal_search(request, pk):
    """
    Trigger manual lien and legal action search for all parcels in a saved list.
    This is the new UX - users explicitly request legal info instead of automatic searches.
    """
    from .background_lien_search import search_parcel_background, should_search_parcel

    saved_list = get_object_or_404(
        _saved_list_queryset_for_user(request.user), pk=pk
    )

    # Get all parcels in the list (not filtered by skip trace status)
    parcel_refs = list(_iter_saved_list_parcel_refs(saved_list))
    if not parcel_refs:
        return JsonResponse(
            {"error": "No parcels in this list."}, status=400
        )

    # Load the actual parcel data
    parcels_by_key = {}
    grouped_loc_ids = defaultdict(list)
    for ref in parcel_refs:
        grouped_loc_ids[ref.town_id].append(ref.loc_id)

    for town_id, loc_list in grouped_loc_ids.items():
        for parcel in load_massgis_parcels_by_ids(town_id, loc_list, saved_list=saved_list):
            normalized = _normalize_loc_id(parcel.loc_id)
            if normalized:
                parcels_by_key[(town_id, normalized)] = parcel

    parcels = []
    for ref in parcel_refs:
        parcel = parcels_by_key.get((ref.town_id, ref.normalized_loc_id))
        if parcel:
            parcels.append(parcel)

    if not parcels:
        return JsonResponse(
            {"error": "No parcels available to search."}, status=400
        )

    # Queue background searches for all residential parcels in the list
    searches_queued = 0
    searches_skipped = 0

    for parcel in parcels:
        town_id = parcel.town.town_id if hasattr(parcel, 'town') else None
        if not town_id:
            searches_skipped += 1
            continue
        loc_id = parcel.loc_id
        property_category = getattr(parcel, 'property_category', '')

        # Only search residential parcels (saves API calls)
        if property_category != 'Residential':
            searches_skipped += 1
            continue

        # Determine county from town name
        county = None
        town_name = parcel.town.name if hasattr(parcel, 'town') else ''
        if town_name:
            town_lower = town_name.lower()
            if town_lower in ["salem", "beverly", "peabody", "lynn", "gloucester", "marblehead", "danvers"]:
                county = "Essex"
            elif town_lower in ["boston", "cambridge", "somerville", "brookline", "chelsea", "revere", "winthrop"]:
                county = "Suffolk"
            elif town_lower in ["worcester", "shrewsbury", "westborough", "auburn", "millbury"]:
                county = "Worcester"
            elif town_lower in ["springfield", "chicopee", "holyoke", "westfield"]:
                county = "Hampden"
            elif town_lower in ["lowell", "cambridge", "newton", "framingham", "waltham"]:
                county = "Middlesex"

        # Queue background search (force=False means it won't duplicate existing searches)
        parcel_data = {
            'owner_name': getattr(parcel, 'owner_name', ''),
            'address': getattr(parcel, 'site_address', ''),
            'town_name': town_name,
            'county': county,
        }

        try:
            if search_parcel_background(request.user, town_id, loc_id, parcel_data, force=False):
                searches_queued += 1
            else:
                # Already searching or searched
                searches_skipped += 1
        except Exception as e:
            logger.warning(f"Failed to queue background search for {town_id}/{loc_id}: {e}")
            searches_skipped += 1

    logger.info(f"Bulk legal search for list {pk}: queued {searches_queued}, skipped {searches_skipped}")

    return JsonResponse({
        "message": f"Legal search started for {searches_queued} properties. Results will appear on the map as they complete.",
        "completed": 0,  # Initial progress
        "total": searches_queued,
        "skipped": searches_skipped,
    })


# --- Lightweight CRM dashboard listing leads and call requests.
@login_required
def crm_overview(request):
    lead_queryset = _lead_queryset_for_user(request.user)
    total_leads = lead_queryset.count()
    distinct_cities_qs = (
        lead_queryset.exclude(site_city__isnull=True)
        .exclude(site_city__exact="")
        .values_list("site_city", flat=True)
        .distinct()
        .order_by("site_city")
    )
    city_names = list(distinct_cities_qs)

    workspace_owner = get_workspace_owner(request.user)
    workspace_username = getattr(workspace_owner, 'username', str(workspace_owner))
    logger.info(f"CRM: Loaded CRM for user {request.user.username}, workspace_owner={workspace_username}")

    # Note: Auto-assignment logic removed - leads are now directly assigned via user_id in QR code URL
    # Each QR code is user-specific, so leads are assigned immediately when the form is submitted

    # Get active leads (not archived)
    active_leads = ScheduleCallRequest.objects.filter(
        created_by=workspace_owner,
        is_archived=False
    ).order_by("-created_at")

    logger.info(f"CRM: Querying for active leads with created_by={workspace_username}, found {active_leads.count()} leads")

    # Organize leads by stage
    leads_by_stage = {
        'new': [],
        'contacted': [],
        'appointment': [],
        'listed': [],
        'under_contract': [],
        'closed': []
    }

    for lead in active_leads:
        if lead.stage in leads_by_stage:
            leads_by_stage[lead.stage].append(lead)

    # Get archived leads
    archived_leads = ScheduleCallRequest.objects.filter(
        created_by=workspace_owner,
        is_archived=True
    ).order_by("-archived_at")

    # Calculate stats
    one_week_ago = timezone.now() - timedelta(days=7)
    stats = {
        'total': active_leads.count(),
        'new': active_leads.filter(created_at__gte=one_week_ago).count(),
        'active': active_leads.exclude(stage__in=['closed']).count(),
        'closed': leads_by_stage['closed'].__len__(),
    }

    return render(
        request,
        "leads/crm_overview.html",
        {
            "total_leads": total_leads,
            "city_count": len(city_names),
            "city_names": city_names,
            "leads_by_stage": leads_by_stage,
            "archived_leads": archived_leads,
            "stats": stats,
        },
    )


# --- Simple per-city lead counter.
@login_required
def crm_city_requests(request, city_slug):
    # Slug is label-ish; show simple list to avoid template dependency.
    leads = _lead_queryset_for_user(request.user).filter(
        site_city__iexact=city_slug.replace("-", " ")
    )
    return HttpResponse(
        f"CRM City: {city_slug} — {leads.count()} leads.",
        content_type="text/plain",
    )


# --- Single parcel geometry API endpoint for map visualization.
@require_GET
def parcel_geometry(request, town_id, loc_id):
    """
    Returns GeoJSON geometry for a single parcel.
    Used for rendering individual parcels on search results map.
    """
    try:
        from .services import (
            _get_massgis_town,
            _ensure_massgis_dataset,
            _find_taxpar_shapefile,
            _lookup_parcel_record,
        )
        from pathlib import Path

        town = _get_massgis_town(town_id)
        dataset_dir = _ensure_massgis_dataset(town)
        tax_par_path = _find_taxpar_shapefile(Path(dataset_dir))

        match = _lookup_parcel_record(tax_par_path, loc_id)
        if match is None:
            return JsonResponse({"error": "Parcel not found"}, status=404)

        shape, attributes = match

        # Convert shapefile shape to GeoJSON
        # Shape has .points (list of [x,y] coordinates) and .shapeType
        if not shape or not hasattr(shape, 'points') or not shape.points:
            return JsonResponse({"error": "Invalid geometry"}, status=404)

        # Convert State Plane coordinates to WGS84 (lat/lng)
        # MassGIS data is in NAD83 Massachusetts State Plane (EPSG:26986)
        from .services import massgis_stateplane_to_wgs84

        wgs84_points = []
        for point in shape.points:
            # point is [x, y] in State Plane meters
            lng, lat = massgis_stateplane_to_wgs84(point[0], point[1])
            wgs84_points.append([lng, lat])

        # Create GeoJSON geometry from converted points
        # For polygons, we need to close the ring
        coordinates = [wgs84_points]

        # Ensure the polygon is closed (first point == last point)
        if coordinates[0][0] != coordinates[0][-1]:
            coordinates[0].append(coordinates[0][0])

        # Extract lot map information from attributes
        map_no = attributes.get('MAP_NO', '')
        map_par_id = attributes.get('MAP_PAR_ID', '')

        geojson = {
            "type": "Feature",
            "geometry": {
                "type": "Polygon",
                "coordinates": coordinates
            },
            "properties": {
                "loc_id": loc_id,
                "town_id": town_id,
                "map_no": map_no,
                "map_par_id": map_par_id,
            }
        }

        return JsonResponse(geojson)

    except MassGISDataError as e:
        return JsonResponse({"error": f"Data error: {str(e)}"}, status=404)
    except Exception as e:
        logger.exception("Error fetching parcel geometry")
        return JsonResponse({"error": str(e)}, status=500)


def town_boundaries(request):
    """
    Returns GeoJSON FeatureCollection of all Massachusetts town boundaries.
    Used for rendering town boundaries on the map.
    """
    try:
        from .services import get_massgis_town_boundaries_geojson
        geojson = get_massgis_town_boundaries_geojson()
        return JsonResponse(geojson, safe=False)
    except Exception as e:
        logger.exception("Error fetching town boundaries")
        return JsonResponse({"error": str(e)}, status=500)


def boston_neighborhoods(request):
    """Return the list of Boston neighborhoods for client-side filtering."""
    try:
        from .services import get_boston_neighborhoods, get_boston_neighborhoods_geojson

        neighborhoods = get_boston_neighborhoods()
        geojson = get_boston_neighborhoods_geojson()
        return JsonResponse({"neighborhoods": neighborhoods, "geojson": geojson})
    except Exception as exc:  # noqa: BLE001
        logger.exception("Error fetching Boston neighborhoods")
        return JsonResponse({"error": str(exc)}, status=500)


def parcels_in_viewport(request):
    """
    Returns parcels within a given map viewport (bounding box).
    Query parameters:
    - north, south, east, west: bounding box coordinates in WGS84
    - limit: max results (default 500, max 2000)
    - property_category, min_price, max_price, town_id, town_name: optional filters
    """
    try:
        from .services import get_parcels_in_bbox

        logger.info(f"Viewport parcels request: {request.GET.dict()}")

        # Get bounding box from query params
        try:
            north = float(request.GET.get('north', 0))
            south = float(request.GET.get('south', 0))
            east = float(request.GET.get('east', 0))
            west = float(request.GET.get('west', 0))
        except (ValueError, TypeError) as e:
            logger.warning(f"Invalid bbox parameters: {e}")
            return JsonResponse({"error": "Invalid bounding box coordinates"}, status=400)

        logger.info(f"Bounding box: N={north}, S={south}, E={east}, W={west}")

        # Get limit - default to unlimited (None means no limit)
        limit = int(request.GET.get('limit')) if request.GET.get('limit') else None

        # Get filters from query params
        filters = {}
        address_contains = request.GET.get('address_contains')
        if address_contains and address_contains.strip():
            filters['address_contains'] = address_contains.strip()
        category = request.GET.get('property_category')
        if category and category != 'any':
            filters['property_category'] = category
        commercial_subtype = request.GET.get('commercial_subtype')
        if commercial_subtype and commercial_subtype != 'any':
            filters['commercial_subtype'] = commercial_subtype
        property_type = request.GET.get('property_type')
        if property_type and property_type.strip() and property_type.lower() != 'any':
            filters['property_type'] = property_type.strip()
        if request.GET.get('min_price'):
            try:
                filters['min_price'] = float(request.GET.get('min_price'))
            except (ValueError, TypeError):
                pass
        if request.GET.get('max_price'):
            try:
                filters['max_price'] = float(request.GET.get('max_price'))
            except (ValueError, TypeError):
                pass
        if request.GET.get('equity_min'):
            try:
                filters['equity_min'] = float(request.GET.get('equity_min'))
            except (ValueError, TypeError):
                pass
        absentee = request.GET.get('absentee')
        if absentee and absentee != 'any':
            filters['absentee'] = absentee
        if request.GET.get('min_years_owned'):
            try:
                filters['min_years_owned'] = int(request.GET.get('min_years_owned'))
            except (ValueError, TypeError):
                pass
        if request.GET.get('max_years_owned'):
            try:
                filters['max_years_owned'] = int(request.GET.get('max_years_owned'))
            except (ValueError, TypeError):
                pass
        shape_filter = _parse_boundary_shape(request.GET)
        if shape_filter:
            logger.info("Applying boundary shape filter: %s", shape_filter.get("type"))
        proximity_address = request.GET.get('proximity_address')
        if proximity_address and proximity_address.strip():
            filters['proximity_address'] = proximity_address.strip()
        if request.GET.get('proximity_radius_miles'):
            try:
                filters['proximity_radius_miles'] = float(request.GET.get('proximity_radius_miles'))
            except (ValueError, TypeError):
                pass
        if request.GET.get('max_years_owned'):
            try:
                filters['max_years_owned'] = int(request.GET.get('max_years_owned'))
            except (ValueError, TypeError):
                pass
        neighborhood = request.GET.get('neighborhood')
        if neighborhood and neighborhood.strip():
            filters['neighborhood'] = neighborhood.strip()
        town_id_param = request.GET.get('town_id')
        if town_id_param:
            try:
                filters['town_id'] = int(town_id_param)
            except (ValueError, TypeError):
                logger.warning("Invalid town_id parameter: %s", town_id_param)
        town_name = request.GET.get('town_name')
        if town_name and town_name.strip():
            town_name_clean = town_name.strip()

            # Check if this is a Boston neighborhood (format: "BOSTON - Neighborhood Name")
            if town_name_clean.upper().startswith("BOSTON - "):
                # Extract neighborhood name and convert to slug format
                neighborhood_name = town_name_clean[9:].strip()  # Remove "BOSTON - " prefix
                # Convert name to slug (e.g., "Back Bay" -> "BACK-BAY")
                neighborhood_slug = neighborhood_name.upper().replace(' ', '-')

                logger.info(f"Detected Boston neighborhood from town_name: '{neighborhood_name}' -> slug: '{neighborhood_slug}'")

                # Set filters for Boston with this neighborhood
                filters['town_id'] = 35  # Boston town ID
                filters['neighborhood'] = neighborhood_slug
            else:
                filters['town_name'] = town_name_clean

        # Safety check: Boston without neighborhood + no limit = 98k+ parcels = server crash
        # Require neighborhood filter or return error message
        if filters.get('town_id') == 35 and not filters.get('neighborhood'):
            logger.warning("Boston parcels requested without neighborhood filter - rejecting request")
            return JsonResponse({
                "count": 0,
                "parcels": [],
                "autoLienSearchEnabled": False,
                "queuedBackgroundSearches": 0,
                "error": "Boston requires neighborhood selection",
                "message": "Please select a Boston neighborhood to load parcels. Boston has 98,845 parcels which is too many to load without filtering."
            })

        # Fetch parcels
        logger.info(f"Fetching parcels with filters: {filters}, limit: {limit}")

        # Check if precomputed data is available (requires migration)
        # Temporarily disabled until migrations run on production
        use_precomputed = False

        # try:
        #     from .models import MassGISParcel
        #     # Check if table exists by testing a simple query
        #     MassGISParcel.objects.exists()
        #     use_precomputed = True
        # except Exception:
        #     use_precomputed = False

        if use_precomputed:
            # Try precomputed database first (50-100x faster)
            try:
                parcels = get_precomputed_parcels_in_bbox(
                    north,
                    south,
                    east,
                    west,
                    limit=limit,
                    shape_filter=shape_filter,
                    **filters,
                )
                logger.info(f"Found {len(parcels)} parcels from precomputed database")
            except Exception as exc:
                # Fallback to file-based search if precomputed fails
                logger.warning(f"Precomputed search failed, falling back to file-based: {exc}")
                parcels = get_parcels_in_bbox(
                    north,
                    south,
                    east,
                    west,
                    limit=limit,
                    shape_filter=shape_filter,
                    **filters,
                )
                logger.info(f"Found {len(parcels)} parcels from file-based search")
        else:
            # Use file-based search (migration not yet run)
            parcels = get_parcels_in_bbox(
                north,
                south,
                east,
                west,
                limit=limit,
                shape_filter=shape_filter,
                **filters,
            )
            logger.info(f"Found {len(parcels)} parcels from file-based search")

        # Automatic lien search disabled - users must manually trigger via "Find Legal Info" button
        auto_lien_search_enabled = False

        # Check for liens and legal actions for each parcel (batch query for efficiency)
        # This displays existing search results but does NOT trigger new searches
        from .models import LienRecord, LegalAction

        # Build list of (town_id, loc_id) tuples for batch query
        parcel_keys = [(p['town_id'], p['loc_id']) for p in parcels]

        # Chunk the queries to avoid "too many SQL variables" error
        # SQLite has a limit of 999 variables, so we chunk by parcel_keys
        # Each parcel_key uses 2 variables (town_id, loc_id)
        chunk_size = 400  # 400 * 2 = 800 variables, safely under 999 limit

        liens_set = set()
        actions_set = set()

        # Process in chunks
        for i in range(0, len(parcel_keys), chunk_size):
            chunk = parcel_keys[i:i + chunk_size]
            town_ids = {town for town, _ in chunk}
            loc_ids = {loc for _, loc in chunk}

            # Batch query for liens (this chunk)
            liens_qs = LienRecord.objects.filter(
                created_by=request.user,
                town_id__in=town_ids,
                loc_id__in=loc_ids,
            ).values_list('town_id', 'loc_id')
            liens_set.update(liens_qs)

            # Batch query for legal actions (this chunk)
            actions_qs = LegalAction.objects.filter(
                town_id__in=town_ids,
                loc_id__in=loc_ids,
            ).filter(
                Q(created_by=request.user) | Q(source__iexact="CourtListener")
            ).values_list('town_id', 'loc_id')
            actions_set.update(actions_qs)

        # No automatic background searches - users trigger manually via "Find Legal Info" button
        searches_queued = 0

        # Format as simple JSON array (not GeoJSON, since we fetch geometry separately)
        results = []
        for parcel in parcels:
            parcel_key = (parcel['town_id'], parcel['loc_id'])
            has_lien = parcel_key in liens_set
            has_legal_action = parcel_key in actions_set

            results.append({
                'loc_id': parcel['loc_id'],
                'town_id': parcel['town_id'],
                'town_name': parcel['town_name'],
                'address': parcel['address'],
                'owner': parcel['owner'],
                'owner_address': parcel.get('owner_address'),
                'total_value': parcel['total_value'],
                'land_value': parcel.get('land_value', 0),
                'building_value': parcel.get('building_value', 0),
                'value_display': f"${parcel['total_value']:,.0f}" if parcel['total_value'] else 'N/A',
                'property_type': parcel['property_type'],
                'property_category': parcel.get('property_category', 'Other'),
                'use_description': parcel.get('use_description', 'Unknown'),
                'style': parcel.get('style'),
                'year_built': parcel.get('year_built'),
                'units': parcel.get('units'),
                'lot_size': parcel.get('lot_size', 0),
                'lot_units': parcel.get('lot_units'),
                'bld_area': parcel.get('bld_area'),
                'total_living_area': parcel.get('total_living_area'),
                'zoning': parcel.get('zoning'),
                'absentee': parcel.get('absentee', False),
                'equity_percent': parcel.get('equity_percent'),
                'last_sale_price': parcel.get('last_sale_price', 0),
                'last_sale_date': parcel.get('last_sale_date'),
                'site_city': parcel.get('site_city'),
                'site_zip': parcel.get('site_zip'),
                'centroid': parcel['centroid'],
                'geometry': parcel.get('geometry', []),
                'has_lien': has_lien,
                'has_legal_action': has_legal_action,
            })

        return JsonResponse({
            'parcels': results,
            'count': len(results),
            'limit': limit,
            'autoLienSearchEnabled': auto_lien_search_enabled,
            'queuedBackgroundSearches': searches_queued,
        })

    except Exception as e:
        logger.exception("Error fetching parcels in viewport")
        return JsonResponse({"error": str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def parcel_flags(request):
    """
    Return lien/legal action flags (and trigger background searches) for a list of parcels.

    This endpoint allows the frontend to load static GeoJSON while still receiving per-user
    lien markers and automated search behavior.
    """
    from .models import LienRecord, LegalAction
    from .background_lien_search import search_parcel_background, should_search_parcel

    try:
        payload = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON payload"}, status=400)

    parcels = payload.get("parcels")
    if not isinstance(parcels, list) or not parcels:
        return JsonResponse({
            "flags": {},
            "autoLienSearchEnabled": True,
            "queuedBackgroundSearches": 0,
        })

    normalized = []
    for item in parcels:
        try:
            town_id = int(item.get("town_id"))
        except (TypeError, ValueError):
            continue
        loc_id_raw = item.get("loc_id")
        if not loc_id_raw:
            continue
        loc_id = str(loc_id_raw).strip()
        if not loc_id:
            continue
        normalized.append({
            "town_id": town_id,
            "loc_id": loc_id,
            "owner": item.get("owner"),
            "address": item.get("address"),
            "town_name": item.get("town_name"),
            "property_category": item.get("property_category"),
        })

    if not normalized:
        return JsonResponse({
            "flags": {},
            "autoLienSearchEnabled": True,
            "queuedBackgroundSearches": 0,
        })

    # Disable auto lien search for parcel-flags endpoint to prevent slowdowns
    # when batching large datasets (each batch of 1000 would trigger searches)
    # Lien search is already handled by the main viewport endpoint
    auto_lien_search_enabled = False

    # Build lookup for response flags
    parcel_keys = [(parcel["town_id"], parcel["loc_id"]) for parcel in normalized]
    flags = {
        f"{town_id}|{loc_id}": {"has_lien": False, "has_legal_action": False}
        for town_id, loc_id in parcel_keys
    }

    chunk_size = 400
    liens_set = set()
    actions_set = set()

    for i in range(0, len(parcel_keys), chunk_size):
        chunk = parcel_keys[i:i + chunk_size]
        town_ids = {town for town, _ in chunk}
        loc_ids = {loc for _, loc in chunk}

        liens_qs = LienRecord.objects.filter(
            created_by=request.user,
            town_id__in=town_ids,
            loc_id__in=loc_ids,
        ).values_list('town_id', 'loc_id')
        liens_set.update(liens_qs)

        actions_qs = LegalAction.objects.filter(
            town_id__in=town_ids,
            loc_id__in=loc_ids,
        ).filter(
            Q(created_by=request.user) | Q(source__iexact="CourtListener")
        ).values_list('town_id', 'loc_id')
        actions_set.update(actions_qs)

    for town_id, loc_id in liens_set:
        key = f"{town_id}|{loc_id}"
        flags.setdefault(key, {"has_lien": False, "has_legal_action": False})
        flags[key]["has_lien"] = True

    for town_id, loc_id in actions_set:
        key = f"{town_id}|{loc_id}"
        flags.setdefault(key, {"has_lien": False, "has_legal_action": False})
        flags[key]["has_legal_action"] = True

    def _infer_county(town_name: Optional[str]) -> Optional[str]:
        if not town_name:
            return None
        town_lower = town_name.lower()
        if town_lower in {"salem", "beverly", "peabody", "lynn", "gloucester", "marblehead", "danvers"}:
            return "Essex"
        if town_lower in {"boston", "cambridge", "somerville", "brookline", "chelsea", "revere", "winthrop"}:
            return "Suffolk"
        if town_lower in {"worcester", "shrewsbury", "westborough", "auburn", "millbury"}:
            return "Worcester"
        return None

    searches_queued = 0
    if auto_lien_search_enabled:
        for parcel in normalized:
            if parcel.get("property_category") != "Residential":
                continue
            town_id = parcel["town_id"]
            loc_id = parcel["loc_id"]
            if not should_search_parcel(request.user, town_id, loc_id):
                continue
            parcel_data = {
                "owner_name": parcel.get("owner") or "",
                "address": parcel.get("address") or "",
                "town_name": parcel.get("town_name") or "",
                "county": _infer_county(parcel.get("town_name")),
            }
            try:
                if search_parcel_background(request.user, town_id, loc_id, parcel_data):
                    searches_queued += 1
            except Exception as exc:  # noqa: BLE001
                logger.warning("Failed to queue background search for %s/%s: %s", town_id, loc_id, exc)

    return JsonResponse({
        "flags": flags,
        "autoLienSearchEnabled": auto_lien_search_enabled,
        "queuedBackgroundSearches": searches_queued,
    })


# ============================================================================
# Lien Record Views
# ============================================================================

@login_required
@require_http_methods(["GET", "POST"])
def lien_create(request, town_id, loc_id):
    """Create a new lien record for a parcel"""
    from .forms import LienRecordForm
    from .models import LienRecord
    
    if request.method == "POST":
        form = LienRecordForm(request.POST)
        if form.is_valid():
            lien = form.save(commit=False)
            lien.created_by = request.user
            lien.town_id = town_id
            lien.loc_id = loc_id
            lien.save()
            messages.success(request, "Lien record created successfully.")
            return redirect("parcel_search_detail", town_id=town_id, loc_id=loc_id)
    else:
        form = LienRecordForm()
    
    context = {
        "form": form,
        "town_id": town_id,
        "loc_id": loc_id,
        "action": "Create",
    }
    return render(request, "leads/lien_form.html", context)


@login_required
@require_http_methods(["GET", "POST"])
def lien_edit(request, pk):
    """Edit an existing lien record"""
    from .forms import LienRecordForm
    from .models import LienRecord
    
    lien = get_object_or_404(LienRecord, pk=pk, created_by=request.user)
    
    if request.method == "POST":
        form = LienRecordForm(request.POST, instance=lien)
        if form.is_valid():
            form.save()
            messages.success(request, "Lien record updated successfully.")
            return redirect("parcel_search_detail", town_id=lien.town_id, loc_id=lien.loc_id)
    else:
        form = LienRecordForm(instance=lien)
    
    context = {
        "form": form,
        "lien": lien,
        "town_id": lien.town_id,
        "loc_id": lien.loc_id,
        "action": "Edit",
    }
    return render(request, "leads/lien_form.html", context)


@login_required
@require_http_methods(["POST"])
def lien_delete(request, pk):
    """Delete a lien record"""
    from .models import LienRecord
    
    lien = get_object_or_404(LienRecord, pk=pk, created_by=request.user)
    town_id = lien.town_id
    loc_id = lien.loc_id
    lien.delete()
    messages.success(request, "Lien record deleted successfully.")
    return redirect("parcel_search_detail", town_id=town_id, loc_id=loc_id)


# ============================================================================
# Legal Action Views
# ============================================================================

@login_required
@require_http_methods(["GET", "POST"])
def legal_action_create(request, town_id, loc_id):
    """Create a new legal action record for a parcel"""
    from .forms import LegalActionForm
    from .models import LegalAction
    
    if request.method == "POST":
        form = LegalActionForm(request.POST)
        if form.is_valid():
            action = form.save(commit=False)
            action.created_by = request.user
            action.town_id = town_id
            action.loc_id = loc_id
            action.save()
            messages.success(request, "Legal action record created successfully.")
            return redirect("parcel_search_detail", town_id=town_id, loc_id=loc_id)
    else:
        form = LegalActionForm()
    
    context = {
        "form": form,
        "town_id": town_id,
        "loc_id": loc_id,
        "action": "Create",
    }
    return render(request, "leads/legal_action_form.html", context)


@login_required
@require_http_methods(["GET", "POST"])
def legal_action_edit(request, pk):
    """Edit an existing legal action record"""
    from .forms import LegalActionForm
    from .models import LegalAction
    
    legal_action = get_object_or_404(LegalAction, pk=pk, created_by=request.user)
    
    if request.method == "POST":
        form = LegalActionForm(request.POST, instance=legal_action)
        if form.is_valid():
            form.save()
            messages.success(request, "Legal action record updated successfully.")
            return redirect("parcel_search_detail", town_id=legal_action.town_id, loc_id=legal_action.loc_id)
    else:
        form = LegalActionForm(instance=legal_action)
    
    context = {
        "form": form,
        "legal_action": legal_action,
        "town_id": legal_action.town_id,
        "loc_id": legal_action.loc_id,
        "action": "Edit",
    }
    return render(request, "leads/legal_action_form.html", context)


@login_required
@require_http_methods(["POST"])
def legal_action_delete(request, pk):
    """Delete a legal action record"""
    from .models import LegalAction
    
    legal_action = get_object_or_404(LegalAction, pk=pk, created_by=request.user)
    town_id = legal_action.town_id
    loc_id = legal_action.loc_id
    legal_action.delete()
    messages.success(request, "Legal action record deleted successfully.")
    return redirect("parcel_search_detail", town_id=town_id, loc_id=loc_id)


@login_required
@require_http_methods(["POST"])
def search_liens_legal_actions(request, town_id, loc_id):
    """
    Search public sources for liens and legal actions on a parcel.
    Uses CourtListener API, provides guidance for MA Trial Court, tax liens, etc.
    """
    from .lien_legal_service import search_all_sources
    from .models import LienRecord, LegalAction
    
    try:
        parcel = get_massgis_parcel_detail(town_id, loc_id)
    except MassGISDataError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    
    owner_name = parcel.owner_name
    address = parcel.site_address
    town_name = parcel.town.name if parcel.town else None
    
    # Determine county from town (simplified - you may want a proper mapping)
    county = None
    if town_name:
        # Common MA counties
        if town_name.lower() in ["salem", "beverly", "peabody", "lynn", "gloucester"]:
            county = "Essex"
        elif town_name.lower() in ["boston", "cambridge", "somerville", "brookline"]:
            county = "Suffolk"
        elif town_name.lower() in ["worcester", "shrewsbury", "westborough"]:
            county = "Worcester"
        # Add more county mappings as needed
    
    # Search all sources
    legal_actions_found, liens_found = search_all_sources(
        owner_name=owner_name or "",
        address=address,
        town_name=town_name,
        county=county
    )
    
    # Store automated findings (except manual lookup guidance)
    saved_actions = 0
    saved_liens = 0
    
    updated_actions = 0
    for action_data in legal_actions_found:
        created, updated = ensure_legal_action_record(request.user, town_id, loc_id, action_data)
        if created:
            saved_actions += 1
        elif updated:
            updated_actions += 1
    
    # Note: Liens from automated search are usually guidance, not actual data
    # Manual entry is typically required
    
    return JsonResponse({
        "success": True,
        "legal_actions_found": len(legal_actions_found),
        "liens_found": len(liens_found),
        "saved_actions": saved_actions,
        "updated_actions": updated_actions,
        "saved_liens": saved_liens,
        "message": (
            f"Found {len(legal_actions_found)} legal action(s) and {len(liens_found)} lien source(s). "
            f"Automatically saved {saved_actions} legal action(s)"
            f"{' and updated ' + str(updated_actions) if updated_actions else ''}. "
            "Review the guidance below for manual lookups."
        ),
        "legal_actions": legal_actions_found,
        "liens": liens_found,
    })


def town_geojson(request, town_id):
    """
    Serve pre-generated GeoJSON for a town (fast path) or fall back to dynamic generation.

    This endpoint provides optimal performance by serving static GeoJSON files when available,
    eliminating the need for shapefile processing on every request.

    Performance:
    - Static file (if pre-generated): <100ms (CDN/static file serving)
    - Dynamic fallback: 2-5 seconds (shapefile processing)

    To pre-generate files: python manage.py generate_town_geojson --towns {town_id}
    """
    from django.conf import settings
    from pathlib import Path
    import os

    try:
        from .services import _get_massgis_town, BOSTON_TOWN_ID
        town = _get_massgis_town(town_id)
        town_name_safe = town.name.replace(' ', '_').replace('/', '_')

        # Boston has 98k+ parcels (917MB GeoJSON) - too large for client-side loading
        # Return 404 to force frontend to use legacy API with server-side filtering
        if int(town_id) == BOSTON_TOWN_ID:
            logger.info(f"Refusing to serve GeoJSON for Boston (town_id={town_id}) - too large (917MB)")
            return JsonResponse({
                "error": "GeoJSON too large for this town",
                "town_id": town_id,
                "town_name": town.name,
                "message": "Boston has 98,845 parcels (917MB). Use /api/parcels-in-viewport/ with neighborhood filter instead.",
                "fallback_api": f"/api/parcels-in-viewport/?town_id={town_id}",
            }, status=404)

        # Try to serve pre-generated GeoJSON first (FAST PATH)
        static_file_name = f"town_{town_id}_{town_name_safe}.geojson"

        # Check multiple possible locations
        possible_paths = [
            Path(settings.BASE_DIR) / "static" / "geojson" / "towns" / static_file_name,
            Path(settings.STATIC_ROOT) / "geojson" / "towns" / static_file_name if settings.STATIC_ROOT else None,
        ]

        for geojson_path in possible_paths:
            if geojson_path and geojson_path.exists():
                logger.info(f"Serving pre-generated GeoJSON for {town.name} from {geojson_path}")

                # Serve the static file with optimal caching headers
                with open(geojson_path, 'r') as f:
                    geojson_data = json.load(f)

                response = JsonResponse(geojson_data, safe=False)
                # Cache for 1 year (immutable data)
                response['Cache-Control'] = 'public, max-age=31536000, immutable'
                response['X-Served-From'] = 'static-geojson'
                return response

        # Attempt to serve from S3 (fastest path) if local files are unavailable
        s3_response = _maybe_redirect_geojson_from_s3(static_file_name)
        if s3_response:
            logger.info(f"Redirecting GeoJSON request for {town.name} to S3")
            return s3_response

        # SLOW PATH: GeoJSON not pre-generated, fall back to dynamic generation
        logger.warning(
            f"Pre-generated GeoJSON not found for {town.name} (town_id={town_id}). "
            f"Falling back to slow dynamic generation. "
            f"Run: python manage.py generate_town_geojson --towns {town_id}"
        )

        # For now, return an error encouraging pre-generation
        # In the future, could fall back to old parcels_in_viewport logic
        return JsonResponse({
            "error": "GeoJSON not pre-generated for this town",
            "town_id": town_id,
            "town_name": town.name,
            "message": (
                f"This town's GeoJSON data has not been pre-generated yet. "
                f"Please run: python manage.py generate_town_geojson --towns {town_id}"
            ),
            "fallback_api": f"/api/parcels-in-viewport/?town_id={town_id}",
        }, status=404)

    except Exception as e:
        logger.exception(f"Error serving town GeoJSON for town_id={town_id}")
        return JsonResponse({
            "error": str(e),
            "town_id": town_id,
        }, status=500)


def _maybe_redirect_geojson_from_s3(static_file_name: str) -> Optional[HttpResponseRedirect]:
    from django.conf import settings

    if not getattr(settings, "USE_S3", False):
        return None

    bucket = getattr(settings, "AWS_STORAGE_BUCKET_NAME", "")
    if not bucket:
        return None

    s3_key = f"geojson/towns/{static_file_name}"
    cache_key = f"geojson_s3_exists::{s3_key}"

    cached = cache.get(cache_key)
    if cached is False:
        return None

    domain = getattr(settings, "AWS_S3_CUSTOM_DOMAIN", "")
    if domain:
        base_url = f"https://{domain}"
    else:
        region = getattr(settings, "AWS_S3_REGION_NAME", "us-east-1")
        base_url = f"https://{bucket}.s3.{region}.amazonaws.com"

    s3_url = f"{base_url}/{s3_key}"

    def _record_missing():
        cache.set(cache_key, False, 300)

    if cached is None:
        try:
            import boto3
            from botocore.exceptions import ClientError

            s3_client = boto3.client(
                "s3",
                aws_access_key_id=getattr(settings, "AWS_ACCESS_KEY_ID", None),
                aws_secret_access_key=getattr(settings, "AWS_SECRET_ACCESS_KEY", None),
                region_name=getattr(settings, "AWS_S3_REGION_NAME", "us-east-1"),
            )
            s3_client.head_object(Bucket=bucket, Key=s3_key)
            cache.set(cache_key, True, 3600)
        except ImportError:
            logger.warning("boto3 not available; cannot verify GeoJSON on S3.")
            return None
        except ClientError as exc:
            error_code = exc.response.get("Error", {}).get("Code")
            if error_code in {"404", "NoSuchKey"}:
                logger.info("GeoJSON %s not found in S3.", s3_key)
                _record_missing()
            else:
                logger.warning("Error checking S3 for GeoJSON %s: %s", s3_key, exc)
            return None
        except Exception as exc:  # noqa: BLE001
            logger.warning("Unexpected error checking S3 for GeoJSON %s: %s", s3_key, exc)
            return None

    response = HttpResponseRedirect(s3_url)
    response['Cache-Control'] = 'public, max-age=31536000, immutable'
    response['X-Served-From'] = 's3-geojson'
    return response


# --- CRM Lead Management AJAX Endpoints ---

@login_required
@require_POST
def crm_update_lead_stage(request, lead_id):
    """Update the stage of a lead via AJAX."""
    try:
        workspace_owner = get_workspace_owner(request.user)
        lead = get_object_or_404(
            ScheduleCallRequest,
            id=lead_id,
            created_by=workspace_owner
        )

        data = json.loads(request.body)
        new_stage = data.get('stage')

        # Validate stage
        valid_stages = [choice[0] for choice in ScheduleCallRequest.STAGE_CHOICES]
        if new_stage not in valid_stages:
            return JsonResponse({'success': False, 'error': 'Invalid stage'}, status=400)

        lead.stage = new_stage
        lead.save(update_fields=['stage', 'updated_at'])

        return JsonResponse({'success': True, 'stage': new_stage})

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        logger.error(f"Error updating lead stage: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def crm_archive_lead(request, lead_id):
    """Archive a lead via AJAX."""
    try:
        workspace_owner = get_workspace_owner(request.user)
        lead = get_object_or_404(
            ScheduleCallRequest,
            id=lead_id,
            created_by=workspace_owner
        )

        lead.is_archived = True
        lead.archived_at = timezone.now()
        lead.save(update_fields=['is_archived', 'archived_at', 'updated_at'])

        return JsonResponse({'success': True})

    except Exception as e:
        logger.error(f"Error archiving lead: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def crm_unarchive_lead(request, lead_id):
    """Unarchive a lead via AJAX."""
    try:
        workspace_owner = get_workspace_owner(request.user)
        lead = get_object_or_404(
            ScheduleCallRequest,
            id=lead_id,
            created_by=workspace_owner
        )

        lead.is_archived = False
        lead.archived_at = None
        lead.save(update_fields=['is_archived', 'archived_at', 'updated_at'])

        return JsonResponse({'success': True})

    except Exception as e:
        logger.error(f"Error unarchiving lead: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def crm_delete_lead(request, lead_id):
    """Permanently delete a lead via AJAX."""
    try:
        workspace_owner = get_workspace_owner(request.user)
        lead = get_object_or_404(
            ScheduleCallRequest,
            id=lead_id,
            created_by=workspace_owner
        )

        lead.delete()

        return JsonResponse({'success': True})

    except Exception as e:
        logger.error(f"Error deleting lead: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
